from __future__ import annotations

import csv
import os
from pathlib import Path

import pytest
from docx import Document
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.opc.constants import RELATIONSHIP_TYPE as RT

from procurement import catalog as catalog_module
from procurement.catalog import POOLED_SPEAKER_LABEL, normalise_header, read_catalog


def write_docx_catalog(path: Path, headers: list[str], rows: list[list[str]]) -> None:
    document = Document()
    table = document.add_table(rows=1 + len(rows), cols=len(headers))
    for column, header in enumerate(headers):
        table.rows[0].cells[column].text = header
    for row_index, values in enumerate(rows, start=1):
        for column, value in enumerate(values):
            table.rows[row_index].cells[column].text = value
    document.save(path)


def write_csv_catalog(path: Path, headers: list[str], rows: list[list[str]]) -> None:
    with path.open("w", encoding="utf-8-sig", newline="") as handle:
        writer = csv.writer(handle)
        writer.writerow(headers)
        writer.writerows(rows)


def add_hyperlink(cell, display_text: str, target: str) -> None:
    """Add a real external hyperlink whose visible label is not the target."""

    relationship_id = cell.part.relate_to(target, RT.HYPERLINK, is_external=True)
    hyperlink = OxmlElement("w:hyperlink")
    hyperlink.set(qn("r:id"), relationship_id)
    run = OxmlElement("w:r")
    run_properties = OxmlElement("w:rPr")
    underline = OxmlElement("w:u")
    underline.set(qn("w:val"), "single")
    run_properties.append(underline)
    text = OxmlElement("w:t")
    text.text = display_text
    run.append(run_properties)
    run.append(text)
    hyperlink.append(run)
    cell._tc.append(hyperlink)


def comparable_sources(catalog) -> list[dict[str, object]]:
    return [
        {
            "source_id": source.source_id,
            "link": source.link,
            "resolved_link": source.resolved_link,
            "source_kind": source.source_kind,
            "speaker": source.speaker,
            "speaker_display": source.speaker_display,
            "metadata": source.metadata,
        }
        for source in catalog.sources
    ]


def test_header_normalisation_is_unicode_aware_and_separator_insensitive() -> None:
    assert normalise_header("  Ｌi_n-k  ") == "link"
    assert normalise_header(" DATE - Accessed ") == "dateaccessed"


def test_csv_requires_only_link_and_preserves_arbitrary_nonblank_metadata(tmp_path: Path) -> None:
    media = tmp_path / "media" / "local clip.mp4"
    media.parent.mkdir()
    media.write_bytes(b"synthetic video")
    catalog_path = tmp_path / "sources.csv"
    headers = [
        " Ｌi_n-k ",
        "S_p-e a k e r",
        "Upload_Date",
        "ENGAGEMENT",
        "Date-Accessed",
        "L e n g t h",
        "Country",
        "Research Group",
        "Language",
    ]
    rows = [
        ["media/local clip.mp4", "", "2026-01-01", "10", "2026-02-01", "00:10", "Ireland", "", "English"],
        ["https://youtu.be/abcdefghijk", "Speaker A", "", "", "", "", "Canada", "Cohort 1", "French"],
        ["https://youtu.be/abcdefghijk", "Speaker A", "", "", "", "", "", "Cohort 2", ""],
    ]
    write_csv_catalog(catalog_path, headers, rows)

    catalog = read_catalog(catalog_path)

    assert catalog.original_headers == tuple(header.strip() for header in headers)
    assert catalog.ignored_headers == ("Upload_Date", "ENGAGEMENT", "Date-Accessed", "L e n g t h")
    assert catalog.metadata_headers == ("Country", "Research Group", "Language")
    assert [source.source_id for source in catalog.sources] == ["source-0001", "source-0002", "source-0003"]
    assert [source.source_kind for source in catalog.sources] == ["local", "youtube", "youtube"]
    assert Path(catalog.sources[0].resolved_link) == media.resolve()
    assert catalog.sources[0].speaker == ""
    assert catalog.sources[0].speaker_display == POOLED_SPEAKER_LABEL
    assert catalog.sources[0].metadata == {"Country": "Ireland", "Language": "English"}
    assert catalog.sources[1].metadata == {
        "Country": "Canada",
        "Research Group": "Cohort 1",
        "Language": "French",
    }
    assert catalog.sources[2].metadata == {"Research Group": "Cohort 2"}
    assert catalog.sources[1].link == catalog.sources[2].link
    assert catalog.sources[1].source_id != catalog.sources[2].source_id


def test_csv_and_docx_catalogs_share_one_row_model(tmp_path: Path) -> None:
    local_video = tmp_path / "clip.mp4"
    local_video.write_bytes(b"synthetic video")
    headers = ["Link", "Speaker", "Country", "Upload Date"]
    rows = [
        ["clip.mp4", "", "Ireland", "ignored"],
        ["https://www.youtube.com/watch?v=abcdefghijk", "Speaker B", "Japan", "ignored"],
    ]
    csv_path = tmp_path / "sources.csv"
    docx_path = tmp_path / "sources.docx"
    write_csv_catalog(csv_path, headers, rows)
    write_docx_catalog(docx_path, headers, rows)

    csv_catalog = read_catalog(csv_path)
    docx_catalog = read_catalog(docx_path)

    assert comparable_sources(csv_catalog) == comparable_sources(docx_catalog)
    assert csv_catalog.original_headers == docx_catalog.original_headers
    assert csv_catalog.ignored_headers == docx_catalog.ignored_headers
    assert csv_catalog.metadata_headers == docx_catalog.metadata_headers


def test_docx_link_column_prefers_hidden_hyperlink_target_over_visible_title(tmp_path: Path) -> None:
    target = "https://www.youtube.com/watch?v=abcdefghijk"
    path = tmp_path / "linked-sources.docx"
    document = Document()
    table = document.add_table(rows=2, cols=2)
    table.cell(0, 0).text = "Link"
    table.cell(0, 1).text = "Country"
    table.cell(1, 0).text = "Readable video title"
    add_hyperlink(table.cell(1, 0), "Readable video title", target)
    table.cell(1, 1).text = "Ireland"
    document.save(path)

    source = read_catalog(path).sources[0]

    assert source.link == target
    assert source.resolved_link == target
    assert source.source_kind == "youtube"
    assert source.metadata == {"Country": "Ireland"}


def test_docx_skips_unrelated_tables_and_preserves_valid_catalog_table_order(tmp_path: Path) -> None:
    path = tmp_path / "mixed-tables.docx"
    document = Document()
    notes = document.add_table(rows=2, cols=2)
    notes.cell(0, 0).text = "Study"
    notes.cell(0, 1).text = "Value"
    notes.cell(1, 0).text = "Notes"
    notes.cell(1, 1).text = "Not a source catalog"
    first = document.add_table(rows=2, cols=2)
    first.cell(0, 0).text = "Link"
    first.cell(0, 1).text = "Country"
    first.cell(1, 0).text = "https://www.youtube.com/watch?v=abcdefghijk"
    first.cell(1, 1).text = "Ireland"
    appendix = document.add_table(rows=2, cols=1)
    appendix.cell(0, 0).text = "Appendix"
    appendix.cell(1, 0).text = "Ignore this table too"
    second = document.add_table(rows=2, cols=2)
    second.cell(0, 0).text = "L-i_n k"
    second.cell(0, 1).text = "Language"
    second.cell(1, 0).text = "https://www.youtube.com/watch?v=lmnopqrstuv"
    second.cell(1, 1).text = "Irish"
    document.save(path)

    catalog = read_catalog(path)

    assert [source.source_id for source in catalog.sources] == ["source-0001", "source-0002"]
    assert [source.youtube_id for source in catalog.sources] == ["abcdefghijk", "lmnopqrstuv"]
    assert catalog.original_headers == ("Link", "Country", "L-i_n k", "Language")
    assert catalog.metadata_headers == ("Country", "Language")


def test_docx_requires_at_least_one_link_bearing_table_document_wide(tmp_path: Path) -> None:
    path = tmp_path / "unrelated-tables.docx"
    document = Document()
    table = document.add_table(rows=2, cols=2)
    table.cell(0, 0).text = "Study"
    table.cell(0, 1).text = "Value"
    table.cell(1, 0).text = "Notes"
    table.cell(1, 1).text = "No sources"
    document.save(path)

    with pytest.raises(ValueError, match="required Link header"):
        read_catalog(path)


def test_catalog_rejects_missing_or_blank_link_values(tmp_path: Path) -> None:
    missing_header = tmp_path / "missing.csv"
    write_csv_catalog(missing_header, ["Speaker", "Country"], [["Speaker A", "Ireland"]])
    with pytest.raises(ValueError, match="required Link header"):
        read_catalog(missing_header)

    blank_link = tmp_path / "blank.csv"
    write_csv_catalog(blank_link, ["Link", "Country"], [["", "Ireland"]])
    with pytest.raises(ValueError, match=r"row 2.*Link"):
        read_catalog(blank_link)


def test_catalog_rejects_ambiguous_normalised_headers(tmp_path: Path) -> None:
    path = tmp_path / "ambiguous.csv"
    write_csv_catalog(path, ["Link", "L-i_n k"], [["a.mp4", "b.mp4"]])
    with pytest.raises(ValueError, match="duplicate normalised header.*link"):
        read_catalog(path)


def test_catalog_csv_reader_is_bounded_before_decoding(tmp_path: Path, monkeypatch) -> None:
    path = tmp_path / "oversized.csv"
    path.write_bytes(b"Link\n" + (b"x" * 64))
    monkeypatch.setattr(catalog_module, "MAX_CATALOG_CSV_BYTES", 16)

    with pytest.raises(ValueError, match="Catalog CSV exceeds 16 bytes"):
        read_catalog(path)


@pytest.mark.parametrize(
    "link",
    (r"\\attacker.example\share\clip.mp4", "//attacker.example/share/clip.mp4", r"\\?\UNC\host\share\clip.mp4"),
)
def test_catalog_rejects_unc_and_device_paths_before_filesystem_access(tmp_path: Path, link: str) -> None:
    path = tmp_path / "sources.csv"
    write_csv_catalog(path, ["Link"], [[link]])

    with pytest.raises(ValueError, match="network|device|namespace"):
        read_catalog(path)


def test_catalog_external_local_file_requires_explicit_cli_authority(tmp_path: Path) -> None:
    catalog_dir = tmp_path / "catalog"
    catalog_dir.mkdir()
    outside = tmp_path / "outside.mp4"
    outside.write_bytes(b"media")
    path = catalog_dir / "sources.csv"
    write_csv_catalog(path, ["Link"], [[str(outside)]])

    with pytest.raises(ValueError, match="outside the catalog directory"):
        read_catalog(path)

    source = read_catalog(path, allow_external_local_paths=True).sources[0]
    assert Path(source.resolved_link) == outside.resolve()


def test_catalog_csv_reader_caps_rows_columns_and_cell_size(tmp_path: Path, monkeypatch) -> None:
    rows_path = tmp_path / "rows.csv"
    write_csv_catalog(rows_path, ["Link"], [["one.mp4"], ["two.mp4"]])
    monkeypatch.setattr(catalog_module, "MAX_CATALOG_ROWS", 1)
    with pytest.raises(ValueError, match="more than 1 data rows"):
        read_catalog(rows_path)

    columns_path = tmp_path / "columns.csv"
    write_csv_catalog(columns_path, ["Link", "Country"], [["one.mp4", "Ireland"]])
    monkeypatch.setattr(catalog_module, "MAX_CATALOG_ROWS", 100)
    monkeypatch.setattr(catalog_module, "MAX_CATALOG_COLUMNS", 1)
    with pytest.raises(ValueError, match="more than 1 columns"):
        read_catalog(columns_path)

    cell_path = tmp_path / "cell.csv"
    write_csv_catalog(cell_path, ["Link", "Country"], [["one.mp4", "Ireland"]])
    monkeypatch.setattr(catalog_module, "MAX_CATALOG_COLUMNS", 10)
    monkeypatch.setattr(catalog_module, "MAX_CATALOG_CELL_CHARS", 4)
    with pytest.raises(ValueError, match="cell longer than 4 characters"):
        read_catalog(cell_path)


def test_catalog_docx_uses_the_existing_bounded_package_reader(tmp_path: Path) -> None:
    path = tmp_path / "not-a-package.docx"
    path.write_bytes(b"not a ZIP package")

    with pytest.raises(RuntimeError, match="Could not open DOCX as a Word package"):
        read_catalog(path)


def test_local_source_identity_rejects_file_symlinks(tmp_path: Path) -> None:
    target = tmp_path / "target.mp4"
    target.write_bytes(b"synthetic video")
    alias = tmp_path / "alias.mp4"
    try:
        os.symlink(target, alias)
    except OSError as exc:
        pytest.skip(f"file symlinks unavailable: {exc}")
    path = tmp_path / "sources.csv"
    write_csv_catalog(path, ["Link"], [["alias.mp4"]])

    with pytest.raises(ValueError, match="symbolic link|reparse"):
        read_catalog(path)


def test_catalog_normalizes_headers_but_not_arbitrary_metadata_values(tmp_path: Path) -> None:
    local_video = tmp_path / "local.mp4"
    local_video.write_bytes(b"synthetic video")
    path = tmp_path / "sources.csv"
    value = "ＡＢＣ ㎏"
    write_csv_catalog(path, ["Link", "Research Note"], [["local.mp4", value]])

    source = read_catalog(path).sources[0]

    assert source.metadata == {"Research Note": value}
