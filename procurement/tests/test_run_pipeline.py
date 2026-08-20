import csv
import json
import os
import subprocess
import sys
import unittest
from pathlib import Path
from tempfile import TemporaryDirectory
from unittest.mock import patch

from docx import Document

from procurement import run_pipeline


class RunPipelineStrategyTests(unittest.TestCase):
    def test_direct_pipeline_script_imports_shared_boundaries_from_any_working_directory(self):
        repository = Path(__file__).resolve().parents[2]
        script = repository / "procurement" / "run_pipeline.py"
        with TemporaryDirectory() as temp_dir:
            result = subprocess.run(
                [sys.executable, str(script), "--help"],
                cwd=temp_dir,
                capture_output=True,
                text=True,
                timeout=30,
            )

        self.assertEqual(result.returncode, 0, result.stderr)
        self.assertIn("usage:", result.stdout.casefold())

    def test_manifest_csv_neutralizes_spreadsheet_formulas_without_changing_json(self):
        with TemporaryDirectory() as temp_dir:
            root = Path(temp_dir)
            item = run_pipeline.PipelineItem(
                table_number=1,
                row_number=2,
                video_id="=1+1",
                url="+cmd",
                speaker="@speaker",
                speaker_reason="\tmetadata",
                license_text="ordinary",
                strategy="-strategy",
                output_path="\rpath",
            )

            run_pipeline.write_manifest((item,), root / "manifest.json", root / "manifest.csv")

            with (root / "manifest.csv").open(encoding="utf-8", newline="") as handle:
                csv_row = next(csv.DictReader(handle))
            json_row = json.loads((root / "manifest.json").read_text(encoding="utf-8"))[0]

        self.assertEqual(csv_row["video_id"], "'=1+1")
        self.assertEqual(csv_row["url"], "'+cmd")
        self.assertEqual(csv_row["speaker"], "'@speaker")
        self.assertEqual(csv_row["speaker_reason"], "'\tmetadata")
        self.assertEqual(csv_row["strategy"], "'-strategy")
        self.assertEqual(csv_row["output_path"], "'\rpath")
        self.assertEqual(csv_row["license_text"], "ordinary")
        self.assertEqual(json_row["video_id"], "=1+1")

    def test_creative_commons_routes_to_full_video(self):
        strategy, reason = run_pipeline.classify_download_strategy(
            "Creative Commons Attribution (CC BY)",
            manual_review_strategy="skip",
        )

        self.assertEqual(strategy, run_pipeline.STRATEGY_FULL_CC)
        self.assertIn("Creative Commons", reason)

    def test_standard_license_routes_to_ten_percent_sample(self):
        strategy, _ = run_pipeline.classify_download_strategy(
            "Standard YouTube License",
            manual_review_strategy="skip",
        )

        self.assertEqual(strategy, run_pipeline.STRATEGY_STANDARD_SAMPLE)

    def test_unknown_license_assumes_standard_by_default(self):
        strategy, _ = run_pipeline.classify_download_strategy(
            "UNKNOWN / NOT RETURNED",
            manual_review_strategy="skip",
        )

        self.assertEqual(strategy, run_pipeline.STRATEGY_ASSUMED_STANDARD_SAMPLE)

    def test_standard_with_unknown_note_still_routes_to_standard_sample(self):
        strategy, reason = run_pipeline.classify_download_strategy(
            "Standard YouTube License; UNKNOWN / NOT RETURNED",
            manual_review_strategy="skip",
        )

        self.assertEqual(strategy, run_pipeline.STRATEGY_STANDARD_SAMPLE)
        self.assertIn("review note", reason)

    def test_manual_review_can_be_overridden_to_standard_sample(self):
        strategy, _ = run_pipeline.classify_download_strategy(
            "UNKNOWN / NOT RETURNED",
            manual_review_strategy="standard-sample",
        )

        self.assertEqual(strategy, run_pipeline.STRATEGY_STANDARD_MANUAL_OVERRIDE)

    def test_load_env_file_reads_config_values(self):
        with TemporaryDirectory() as temp_dir:
            config_path = Path(temp_dir) / "config.env"
            config_path.write_text(
                "# comment\nYOUTUBE_API_KEY='abc123'\nSCAN_TAGS=false\n",
                encoding="utf-8",
            )

            values = run_pipeline.load_env_file(config_path)

        self.assertEqual(values["YOUTUBE_API_KEY"], "abc123")
        self.assertEqual(values["SCAN_TAGS"], "false")

    def test_resolve_api_key_ignores_plaintext_config_values(self):
        with patch.dict(os.environ, {"YOUTUBE_API_KEY": ""}):
            value = run_pipeline.resolve_api_key(
                "PASTE_YOUR_KEY_HERE",
                {"YOUTUBE_API_KEY": "plaintext-file-key"},
            )

        self.assertEqual(value, "")

    def test_license_audit_passes_api_key_only_in_child_environment(self):
        captured: dict[str, object] = {}

        def capture(command, cwd=None, env=None):
            captured["command"] = command
            captured["env"] = env

        with TemporaryDirectory() as temp_dir, patch.object(run_pipeline, "run_command", side_effect=capture), patch.object(
            run_pipeline, "validate_docx"
        ):
            root = Path(temp_dir)
            run_pipeline.run_license_audit(
                docx_path=root / "input.docx",
                audit_dir=root / "audit",
                api_key="environment-only-secret",
                terms_json=None,
            )

        self.assertNotIn("environment-only-secret", captured["command"])
        self.assertEqual(captured["env"]["YOUTUBE_API_KEY"], "environment-only-secret")

    def test_command_display_redacts_api_key(self):
        command = [
            sys.executable,
            "audit_docx.py",
            "--api-key",
            "secret-value",
            "--output",
            "out.docx",
        ]

        display = run_pipeline.format_command_for_display(command)

        self.assertIn("--api-key <redacted>", display)
        self.assertNotIn("secret-value", display)

    def test_speaker_filtered_docx_removes_unselected_video_rows_before_audit(self):
        with TemporaryDirectory() as temp_dir:
            root = Path(temp_dir)
            source = root / "source.docx"
            filtered = root / "selected.docx"
            document = Document()
            table = document.add_table(rows=3, cols=2)
            table.rows[0].cells[0].text = "Video"
            table.rows[0].cells[1].text = "Speaker"
            table.rows[1].cells[0].text = "https://www.youtube.com/watch?v=aaaaaaaaaaa"
            table.rows[1].cells[1].text = "Speaker A"
            table.rows[2].cells[0].text = "https://www.youtube.com/watch?v=bbbbbbbbbbb"
            table.rows[2].cells[1].text = "Speaker B"
            document.save(source)

            result = run_pipeline.write_speaker_filtered_docx(source, filtered, [" speaker   a "])
            rows = run_pipeline.docx_extractions.find_video_rows(Document(str(result)))

        self.assertEqual(len(rows), 1)
        self.assertEqual(rows[0].speaker, "Speaker A")

    def test_speaker_filter_opens_source_through_docx_preflight(self):
        with TemporaryDirectory() as temp_dir:
            root = Path(temp_dir)
            source = root / "source.docx"
            filtered = root / "selected.docx"
            document = Document()
            table = document.add_table(rows=2, cols=2)
            table.rows[0].cells[0].text = "Video"
            table.rows[0].cells[1].text = "Speaker"
            table.rows[1].cells[0].text = "https://www.youtube.com/watch?v=aaaaaaaaaaa"
            table.rows[1].cells[1].text = "Speaker A"
            document.save(source)
            with patch.object(
                run_pipeline.docx_extractions,
                "open_docx_document",
                wraps=run_pipeline.docx_extractions.open_docx_document,
            ) as safe_open:
                run_pipeline.write_speaker_filtered_docx(source, filtered, ["Speaker A"])

        safe_open.assert_any_call(source)

    def test_cli_handles_unicode_speaker_names_on_cp1252_windows_console(self):
        with TemporaryDirectory() as temp_dir:
            temp_path = Path(temp_dir)
            docx_path = temp_path / "unicode_speakers.docx"
            output_root = temp_path / "output"
            document = Document()
            table = document.add_table(rows=2, cols=3)
            table.rows[0].cells[0].text = "Video"
            table.rows[0].cells[1].text = "Speaker"
            table.rows[0].cells[2].text = "License"
            table.rows[1].cells[0].text = "https://www.youtube.com/watch?v=abcdefghijk"
            table.rows[1].cells[1].text = "Sławomir Mentzen"
            table.rows[1].cells[2].text = "Standard YouTube License"
            document.save(docx_path)

            env = os.environ.copy()
            env["PYTHONIOENCODING"] = "cp1252:strict"
            env["PYTHONUTF8"] = "0"
            result = subprocess.run(
                [
                    sys.executable,
                    "-m",
                    "procurement.run_pipeline",
                    str(docx_path),
                    "--audited-docx",
                    str(docx_path),
                    "--output-root",
                    str(output_root),
                    "--manual-review-strategy",
                    "standard-sample",
                    "--dry-run",
                ],
                cwd=Path(__file__).resolve().parents[2],
                env=env,
                capture_output=True,
                timeout=30,
            )

        self.assertEqual(result.returncode, 0, (result.stdout + result.stderr).decode("utf-8", errors="replace"))
        self.assertNotIn(b"UnicodeEncodeError", result.stderr)


if __name__ == "__main__":
    unittest.main()
