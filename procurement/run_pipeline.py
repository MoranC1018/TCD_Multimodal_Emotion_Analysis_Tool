#!/usr/bin/env python3
"""Run the complete procurement pipeline from one DOCX file.

The pipeline keeps the individual steps reviewable:

1. Audit YouTube licences in the input DOCX.
2. Read the audited licence column.
3. Route each video to the correct download workflow:
   - Creative Commons: full-video download.
   - Standard YouTube License: 10 percent iMotions-ready sample.
   - Unknown or missing licence: assumed Standard YouTube License, with a note.

Generated documents, logs, media, and manifests are written under
``procurement/output/`` by default, which is intentionally gitignored.
"""

from __future__ import annotations

import argparse
import csv
import io
import json
import os
import re
import subprocess
import sys
import zipfile
from xml.etree import ElementTree as ET
from dataclasses import asdict, dataclass
from datetime import datetime
from pathlib import Path
from typing import Iterable

if __package__ in {None, ""}:
    sys.path.insert(0, str(Path(__file__).resolve().parents[1]))

from spreadsheet_safety import SpreadsheetSafeWriter

try:
    from docx import Document
except ImportError as exc:  # pragma: no cover - exercised by users without deps
    raise SystemExit("Missing dependency: python-docx. Install with `python -m pip install python-docx`.") from exc

from procurement.console import configure_utf8_stdio
from procurement.video_sampling import run_docx_extractions as docx_extractions


LICENSE_HEADERS = [
    "License",
    "YouTube API License",
    "Raw YouTube API License",
]

STRATEGY_FULL_CC = "creative_commons_full_video"
STRATEGY_STANDARD_SAMPLE = "standard_license_10_percent_sample"
STRATEGY_ASSUMED_STANDARD_SAMPLE = "assumed_standard_license_10_percent_sample"
STRATEGY_FULL_MANUAL_OVERRIDE = "manual_review_full_video_override"
STRATEGY_STANDARD_MANUAL_OVERRIDE = "manual_review_10_percent_sample_override"


@dataclass
class PipelineItem:
    """One row-level decision and outcome written to the manifest."""

    table_number: int
    row_number: int
    video_id: str
    url: str
    speaker: str
    speaker_reason: str
    license_text: str
    strategy: str
    status: str = "planned"
    output_path: str = ""
    notes: str = ""


def normalise_header(text: str) -> str:
    """Normalise DOCX table headers for forgiving lookups."""

    return " ".join(str(text or "").strip().lower().split())


def first_row_header_map(table) -> dict[str, int]:
    """Map the first row's header text to cell indexes."""

    if not table.rows:
        return {}
    return {
        normalise_header(cell.text): index
        for index, cell in enumerate(table.rows[0].cells)
        if normalise_header(cell.text)
    }


def get_row_license_text(document, video_row: docx_extractions.VideoRow) -> str:
    """Read the best available licence value for a YouTube row."""

    table = document.tables[video_row.table_index]
    row = table.rows[video_row.row_index]
    header_map = first_row_header_map(table)

    for header in LICENSE_HEADERS:
        index = header_map.get(normalise_header(header))
        if index is not None and index < len(row.cells):
            value = row.cells[index].text.strip()
            if value:
                return value

    return ""


def classify_download_strategy(license_text: str, manual_review_strategy: str) -> tuple[str, str]:
    """Return the download strategy and a short human-readable reason."""

    value = " ".join(str(license_text or "").strip().split())
    lower_value = value.lower()

    # The launcher uses this explicit strategy for its Full video mode. It is
    # a user choice, not merely a fallback for rows with unknown licensing.
    if manual_review_strategy == "full-video":
        return STRATEGY_FULL_MANUAL_OVERRIDE, "Full-video mode selected by the user."

    is_creative_commons = any(
        marker in lower_value
        for marker in [
            "creative commons",
            "creativecommon",
            "cc by",
            "cc-by",
        ]
    )
    is_standard = any(
        marker in lower_value
        for marker in [
            "standard youtube license",
            "standard youtube licence",
            "standard license",
            "standard licence",
        ]
    ) or lower_value == "youtube"
    needs_manual_review = any(
        marker in lower_value
        for marker in [
            "manual review",
            "unknown",
            "not returned",
            "missing",
            "no youtube url found",
        ]
    )

    if is_standard:
        if needs_manual_review:
            return (
                STRATEGY_STANDARD_SAMPLE,
                "Standard YouTube licence detected; audit also reported unknown/manual-review wording, so this row was processed as Standard with a review note.",
            )
        return STRATEGY_STANDARD_SAMPLE, "Standard YouTube licence detected."

    if needs_manual_review or not lower_value:
        if manual_review_strategy == "standard-sample":
            return STRATEGY_STANDARD_MANUAL_OVERRIDE, "Manual-review row forced to 10 percent sampling."
        return (
            STRATEGY_ASSUMED_STANDARD_SAMPLE,
            "Licence was missing, unknown, or not returned in the DOCX/API audit; assumed Standard YouTube License for 10 percent sampling.",
        )

    if is_creative_commons:
        return STRATEGY_FULL_CC, "Creative Commons licence detected."

    return (
        STRATEGY_ASSUMED_STANDARD_SAMPLE,
        "No supported licence value was found in the DOCX/API audit; assumed Standard YouTube License for 10 percent sampling.",
    )


def timestamp_for_path() -> str:
    return datetime.now().strftime("%Y%m%d_%H%M%S")


def load_env_file(path: Path) -> dict[str, str]:
    """Load simple nonsecret KEY=VALUE workflow settings from config.env."""

    if not path.exists():
        return {}

    values: dict[str, str] = {}
    for line in path.read_text(encoding="utf-8", errors="replace").splitlines():
        stripped = line.strip()
        if not stripped or stripped.startswith("#") or "=" not in stripped:
            continue
        key, value = stripped.split("=", 1)
        values[key.strip()] = value.strip().strip('"').strip("'")
    return values


def looks_like_placeholder_secret(value: str | None) -> bool:
    """Return True for blank or example API-key text."""

    if not value:
        return True
    cleaned = value.strip()
    if not cleaned:
        return True
    upper_value = cleaned.upper()
    return any(marker in upper_value for marker in ["YOUR_", "PASTE", "REPLACE", "TODO"])


def resolve_api_key(cli_value: str | None, config_values: dict[str, str]) -> str:
    """Resolve the YouTube API key from an explicit argument or the environment."""

    _ = config_values  # Nonsecret file settings remain supported; secrets do not.
    for value in [cli_value, os.getenv("YOUTUBE_API_KEY")]:
        if not looks_like_placeholder_secret(value):
            return str(value).strip()
    return ""


def validate_docx(path: Path) -> None:
    """Fail fast if a generated DOCX is not a readable Office ZIP package."""

    try:
        snapshot = docx_extractions.read_docx_snapshot(path)
        with zipfile.ZipFile(io.BytesIO(snapshot)) as archive:
            bad_member = archive.testzip()
            if bad_member:
                raise ValueError(f"bad ZIP member: {bad_member}")
            for member in archive.namelist():
                if member.endswith(".xml") or member.endswith(".rels"):
                    ET.fromstring(archive.read(member))
                if member.endswith(".rels"):
                    validate_relationship_targets(member, archive.read(member))
        docx_extractions.open_docx_snapshot(snapshot, path)
    except Exception as exc:
        raise RuntimeError(f"Generated DOCX did not validate cleanly: {path}") from exc


def validate_relationship_targets(member_name: str, content: bytes) -> None:
    """Catch raw filesystem characters in DOCX hyperlink relationships.

    Microsoft Word can reject a document when local folder hyperlinks contain
    raw spaces, accents, non-breaking spaces, or emoji. Python XML parsing does
    not catch that because the relationship file is still well-formed XML, so
    this check validates the stricter URI-like target Word expects.
    """

    relationships = ET.fromstring(content)
    for relationship in relationships:
        relationship_type = relationship.attrib.get("Type", "")
        target = relationship.attrib.get("Target", "")
        mode = relationship.attrib.get("TargetMode", "")
        if "hyperlink" not in relationship_type or mode != "External":
            continue
        has_control_char = any(ord(character) < 32 for character in target)
        has_raw_non_ascii = any(ord(character) > 127 for character in target)
        if has_control_char or has_raw_non_ascii or " " in target:
            raise ValueError(f"unsafe hyperlink relationship in {member_name}: {target!r}")


def run_command(
    command: list[str],
    cwd: Path | None = None,
    env: dict[str, str] | None = None,
) -> None:
    """Print and run one subprocess command."""

    print("Running:", format_command_for_display(command), flush=True)
    child_env = credential_free_environment() if env is None else dict(env)
    subprocess.run(command, cwd=cwd, check=True, env=child_env)


def credential_free_environment() -> dict[str, str]:
    """Copy the environment without forwarding launcher credentials by default."""

    env = os.environ.copy()
    for name in ("YOUTUBE_API_KEY", "HF_TOKEN", "HUGGINGFACE_TOKEN", "HUGGING_FACE_HUB_TOKEN"):
        env.pop(name, None)
    return env


def format_command_for_display(command: list[str]) -> str:
    """Return a shell-readable command string with secrets redacted."""

    return " ".join(
        f'"{part}"' if " " in str(part) else str(part)
        for part in redact_command_for_display(command)
    )


def redact_command_for_display(command: list[str]) -> list[str]:
    """Return a display-only command where secret flag values are hidden."""

    redacted = list(command)
    secret_flags = {"--api-key"}
    for index, part in enumerate(redacted):
        part_text = str(part)
        if part_text in secret_flags and index + 1 < len(redacted):
            redacted[index + 1] = "<redacted>"
        elif any(part_text.startswith(f"{flag}=") for flag in secret_flags):
            flag_name = part_text.split("=", 1)[0]
            redacted[index] = f"{flag_name}=<redacted>"
    return redacted


def run_license_audit(
    *,
    docx_path: Path,
    audit_dir: Path,
    api_key: str | None,
    terms_json: Path | None,
) -> Path:
    """Run the licence audit and return the audited DOCX path."""

    audit_dir.mkdir(parents=True, exist_ok=True)
    safe_stem = docx_path.stem.replace(" ", "_")
    audited_docx = audit_dir / f"{docx_path.stem}_license_audit.docx"

    command = [
        sys.executable,
        str(Path(__file__).resolve().parent / "license_check" / "audit_docx.py"),
        str(docx_path),
        "--output",
        str(audited_docx),
        "--debug-csv",
        str(audit_dir / f"{safe_stem}_audit_debug.csv"),
        "--summary-txt",
        str(audit_dir / f"{safe_stem}_license_audit_summary.txt"),
        "--log-docx",
        str(audit_dir / f"{safe_stem}_license_audit_log.docx"),
    ]

    if terms_json:
        command.extend(["--terms-json", str(terms_json)])

    env = credential_free_environment()
    if api_key:
        env["YOUTUBE_API_KEY"] = api_key
    run_command(command, env=env)
    validate_docx(audited_docx)
    return audited_docx


def build_pipeline_items(document, manual_review_strategy: str) -> list[PipelineItem]:
    """Scan the audited DOCX and classify every YouTube row."""

    items: list[PipelineItem] = []
    for video_row in docx_extractions.find_video_rows(document):
        license_text = get_row_license_text(document, video_row)
        strategy, reason = classify_download_strategy(license_text, manual_review_strategy)
        items.append(
            PipelineItem(
                table_number=video_row.table_index + 1,
                row_number=video_row.row_index + 1,
                video_id=video_row.video_id,
                url=video_row.url,
                speaker=video_row.speaker,
                speaker_reason=video_row.speaker_reason,
                license_text=license_text,
                strategy=strategy,
                notes=reason,
            )
        )
    return items


def filter_pipeline_items_by_speaker(
    items: Iterable[PipelineItem],
    selected_speakers: Iterable[str] | None,
) -> list[PipelineItem]:
    """Keep only exact speaker groups selected in the launcher."""

    requested = {
        docx_extractions.speaker_match_key(value)
        for value in selected_speakers or []
        if docx_extractions.speaker_match_key(value)
    }
    item_list = list(items)
    if not requested:
        return item_list
    return [item for item in item_list if docx_extractions.speaker_match_key(item.speaker) in requested]


def write_speaker_filtered_docx(
    source_docx: Path,
    destination: Path,
    selected_speakers: Iterable[str],
) -> Path:
    """Write a DOCX containing only selected video rows before API auditing."""

    requested = {
        docx_extractions.speaker_match_key(value)
        for value in selected_speakers
        if docx_extractions.speaker_match_key(value)
    }
    if not requested:
        return source_docx
    document = docx_extractions.open_docx_document(source_docx)
    rows = docx_extractions.find_video_rows(document)
    matching = [row for row in rows if docx_extractions.speaker_match_key(row.speaker) in requested]
    if not matching:
        raise ValueError("No DOCX videos matched the selected speaker groups.")
    matching_coordinates = {(row.table_index, row.row_index) for row in matching}
    row_elements = [
        document.tables[row.table_index].rows[row.row_index]._tr
        for row in rows
        if (row.table_index, row.row_index) not in matching_coordinates
    ]
    for row_element in row_elements:
        row_element.getparent().remove(row_element)
    destination.parent.mkdir(parents=True, exist_ok=True)
    document.save(str(destination))
    validate_docx(destination)
    return destination


def write_manifest(items: Iterable[PipelineItem], json_path: Path, csv_path: Path) -> None:
    """Write machine-readable and spreadsheet-friendly pipeline manifests."""

    rows = [asdict(item) for item in items]
    json_path.parent.mkdir(parents=True, exist_ok=True)
    json_path.write_text(json.dumps(rows, indent=2, ensure_ascii=False) + "\n", encoding="utf-8")

    with csv_path.open("w", newline="", encoding="utf-8") as handle:
        fieldnames = list(PipelineItem.__dataclass_fields__.keys())
        writer = SpreadsheetSafeWriter(csv.DictWriter(handle, fieldnames=fieldnames))
        writer.writeheader()
        writer.writerows(rows)


def newest_matching_folder(folder: Path, video_id: str, suffix: str | None = None) -> Path | None:
    """Find the newest output folder for a video ID."""

    if not folder.exists():
        return None

    candidates = []
    for child in folder.iterdir():
        if not child.is_dir() or f"[{video_id}]" not in child.name:
            continue
        if suffix and not child.name.endswith(suffix):
            continue
        candidates.append(child)

    if not candidates:
        return None
    return max(candidates, key=lambda path: path.stat().st_mtime)


def run_standard_sample(
    *,
    item: PipelineItem,
    speaker_folder: Path,
    force: bool,
    no_stitch: bool,
    dry_run: bool,
    extra_args: list[str],
) -> Path | None:
    """Run the existing 10 percent standard-licence extraction flow."""

    if dry_run:
        return None

    script_folder = Path(__file__).resolve().parent / "video_sampling"
    extractor_script = script_folder / "extraction_router.py"
    passthrough_args = list(extra_args)
    if no_stitch:
        passthrough_args.append("--skip-stitch")

    return docx_extractions.extract_or_reuse_folder(
        url=item.url,
        video_id=item.video_id,
        extractor_script=extractor_script,
        working_folder=speaker_folder,
        force=force,
        extra_extractor_args=passthrough_args,
    )


def run_full_video_download(
    *,
    item: PipelineItem,
    speaker_folder: Path,
    dry_run: bool,
    allow_non_cc: bool,
    extra_args: list[str],
) -> Path | None:
    """Run the full Creative Commons video download flow."""

    command = [
        sys.executable,
        str(Path(__file__).resolve().parent / "video_sampling" / "full_video_download.py"),
        item.url,
        "--output-root",
        str(speaker_folder),
    ]
    if dry_run:
        command.append("--dry-run")
    if allow_non_cc:
        command.append("--allow-non-cc")
    command.extend(extra_args)

    run_command(command)
    return newest_matching_folder(speaker_folder, item.video_id, suffix="_full_video")


def apply_numeric_prefix(folder: Path, index: int) -> Path:
    """Rename a download folder to add a global sequential prefix (e.g. 003_Title_[id]).

    Strips any existing three-digit prefix first so re-runs with different
    ordering produce the correct name rather than double-prefixing.
    """
    prefix = f"{index:03d}_"
    if folder.name.startswith(prefix):
        return folder
    stem = re.sub(r"^\d{3}_", "", folder.name)
    new_path = folder.parent / f"{prefix}{stem}"
    if new_path != folder:
        folder.rename(new_path)
    return new_path


def parse_args() -> argparse.Namespace:
    parser = argparse.ArgumentParser(description="Run licence audit, then route videos to CC/full or standard/10 percent procurement.")
    parser.add_argument("docx_path", type=Path, help="Source DOCX containing YouTube video rows.")
    parser.add_argument("--audited-docx", type=Path, default=None, help="Use an existing audited DOCX instead of running the licence audit first.")
    parser.add_argument("--output-root", type=Path, default=Path(__file__).resolve().parent / "output", help="Generated pipeline run folder root.")
    parser.add_argument("--download-root", type=Path, default=None, help="Root folder for speaker download folders. Defaults inside the pipeline run folder.")
    parser.add_argument("--api-key", default=os.getenv("YOUTUBE_API_KEY"), help="YouTube Data API key for the licence audit.")
    parser.add_argument("--config", type=Path, default=Path(__file__).resolve().parent / "license_check" / "config.env", help="Optional file containing nonsecret workflow settings.")
    parser.add_argument("--terms-json", type=Path, default=Path(__file__).resolve().parent / "license_check" / "license_terms_dictionary.json")
    parser.add_argument("--limit", type=int, default=None, help="Only process the first N classified video rows.")
    parser.add_argument(
        "--speaker",
        action="append",
        default=[],
        help="Process only this resolved speaker label. May be repeated.",
    )
    parser.add_argument("--force", action="store_true", help="Re-run standard sampling even when a completed output folder exists.")
    parser.add_argument("--no-stitch", action="store_true", help="For standard-licence videos, keep raw clips instead of creating stitched_imotions.mp4.")
    parser.add_argument("--dry-run", action="store_true", help="Run the audit/classification and print planned download commands without downloading media.")
    parser.add_argument(
        "--manual-review-strategy",
        choices=["assume-standard", "standard-sample", "full-video", "skip"],
        default="assume-standard",
        help=(
            "What to do with manual-review or unknown licence rows. "
            "Default: assume Standard YouTube License and sample 10 percent. "
            "`skip` is kept as a legacy alias for the same assumption."
        ),
    )
    parser.add_argument("--sample-arg", action="append", default=[], help="Extra argument passed to the 10 percent sampler. Repeat as needed.")
    parser.add_argument("--full-video-arg", action="append", default=[], help="Extra argument passed to the full-video downloader. Repeat as needed.")
    return parser.parse_args()


def main() -> int:
    configure_utf8_stdio()
    args = parse_args()
    source_docx = args.docx_path.resolve()
    if not source_docx.exists():
        raise FileNotFoundError(f"DOCX not found: {source_docx}")

    run_dir = (args.output_root.resolve() / f"{source_docx.stem}_procurement_{timestamp_for_path()}")
    audit_dir = run_dir / "license_audit"
    download_root = args.download_root.resolve() if args.download_root else run_dir / "downloads"
    download_root.mkdir(parents=True, exist_ok=True)

    config_values = load_env_file(args.config.resolve())
    api_key = resolve_api_key(args.api_key, config_values)
    terms_json = args.terms_json.resolve() if args.terms_json and args.terms_json.exists() else None
    if args.audited_docx:
        audited_docx = args.audited_docx.resolve()
        validate_docx(audited_docx)
    else:
        if not api_key:
            raise SystemExit(
                "Missing YouTube API key. Set YOUTUBE_API_KEY or pass --api-key."
            )
        audit_source_docx = write_speaker_filtered_docx(
            source_docx,
            run_dir / "selected_input" / f"{source_docx.stem}_selected_speakers.docx",
            args.speaker,
        )
        audited_docx = run_license_audit(
            docx_path=audit_source_docx,
            audit_dir=audit_dir,
            api_key=api_key,
            terms_json=terms_json,
        )

    document = docx_extractions.open_docx_document(audited_docx)
    items = build_pipeline_items(document, args.manual_review_strategy)
    items = filter_pipeline_items_by_speaker(items, args.speaker)
    if args.speaker and not items:
        raise ValueError("No audited DOCX videos matched the selected speaker groups.")
    if args.limit is not None:
        items = items[: args.limit]

    for table_index in sorted({item.table_number - 1 for item in items}):
        docx_extractions.ensure_extraction_column(document.tables[table_index])

    linked_docx = run_dir / f"{audited_docx.stem}_with_download_links.docx"
    manifest_json = run_dir / "procurement_manifest.json"
    manifest_csv = run_dir / "procurement_manifest.csv"

    print(f"Pipeline run folder: {run_dir}", flush=True)
    print(f"Audited DOCX: {audited_docx}", flush=True)
    print(f"Download root: {download_root}", flush=True)

    for index, item in enumerate(items, start=1):
        print("\n" + "=" * 70, flush=True)
        print(f"[{index}/{len(items)}] {item.video_id} | {item.speaker} | {item.strategy}", flush=True)
        print(f"Licence: {item.license_text or '(missing)'}", flush=True)

        row = document.tables[item.table_number - 1].rows[item.row_number - 1]
        speaker_folder = download_root / docx_extractions.make_folder_name_safe(item.speaker)
        speaker_folder.mkdir(parents=True, exist_ok=True)

        try:
            if item.strategy in {STRATEGY_STANDARD_SAMPLE, STRATEGY_STANDARD_MANUAL_OVERRIDE, STRATEGY_ASSUMED_STANDARD_SAMPLE}:
                output_folder = run_standard_sample(
                    item=item,
                    speaker_folder=speaker_folder,
                    force=args.force,
                    no_stitch=args.no_stitch,
                    dry_run=args.dry_run,
                    extra_args=args.sample_arg,
                )
            elif item.strategy in {STRATEGY_FULL_CC, STRATEGY_FULL_MANUAL_OVERRIDE}:
                # The audit step has already checked YouTube's API licence.
                # This lets the full-video downloader proceed even when yt-dlp
                # exposes a different or empty licence string.
                output_folder = run_full_video_download(
                    item=item,
                    speaker_folder=speaker_folder,
                    dry_run=args.dry_run,
                    allow_non_cc=True,
                    extra_args=args.full_video_arg,
                )
            else:
                item.status = "skipped"
                docx_extractions.write_failure_to_row(row, "SKIPPED: manual review or unsupported licence")
                continue

            item.status = "planned" if args.dry_run else "success"
            if output_folder:
                if not args.dry_run:
                    output_folder = apply_numeric_prefix(output_folder, index)
                item.output_path = str(output_folder)
                docx_extractions.add_folder_link_to_row(row, output_folder, linked_docx)
            else:
                docx_extractions.write_failure_to_row(row, f"DRY RUN: {item.strategy}")
        except Exception as exc:
            item.status = "failed"
            item.notes = f"{item.notes} Error: {type(exc).__name__}: {exc}"
            docx_extractions.write_failure_to_row(row, f"FAILED: {type(exc).__name__}")
            print(f"ERROR: {exc}", flush=True)

        write_manifest(items, manifest_json, manifest_csv)

    linked_docx.parent.mkdir(parents=True, exist_ok=True)
    document.save(str(linked_docx))
    validate_docx(linked_docx)
    write_manifest(items, manifest_json, manifest_csv)

    print("\nPipeline complete.", flush=True)
    print(f"Linked DOCX: {linked_docx}", flush=True)
    print(f"Manifest JSON: {manifest_json}", flush=True)
    print(f"Manifest CSV: {manifest_csv}", flush=True)
    return 0 if all(item.status in {"success", "planned", "skipped"} for item in items) else 1


if __name__ == "__main__":
    raise SystemExit(main())
