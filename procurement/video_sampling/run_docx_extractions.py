#!/usr/bin/env python3
"""
run_docx_extractions.py

Purpose
-------
This script reads a DOCX file, finds YouTube links inside every
DOCX table, calls extractor.py once for each unique YouTube video, and
then adds a relative hyperlink back into the table row pointing to the folder
created by extractor.py.

Expected folder layout
----------------------
Put these files in the same folder:

    extractor.py
    extraction_router.py
    run_docx_extractions.py
    source_catalog.docx

The extractor.py script should create one folder per video, using the
format you already chose:

    Video_Title_[youtubeVideoID]

This wrapper script will add links to those folders in a new right-hand table
column called:

    Extraction folder

Why the links are relative
--------------------------
The inserted hyperlinks point to folders using relative paths such as:

    Some_Video_Title_[dQw4w9WgXcQ]/

That means you can share the whole project folder with someone else, and the
links should still work as long as the edited DOCX stays beside the extracted
video folders.

Requirements
------------
Install python-docx if needed:

    pip install python-docx

extractor.py must also work by itself before using this wrapper.

Example
-------
    python -m procurement.video_sampling.run_docx_extractions "source_catalog.docx"

Useful test run
---------------
Process only the first 3 videos, useful before starting the full document:

    python -m procurement.video_sampling.run_docx_extractions "source_catalog.docx" --limit 3
"""

from __future__ import annotations

import argparse
import io
import json
import os
import re
import subprocess
import sys
import time
import urllib.parse
from dataclasses import dataclass
from datetime import datetime
from pathlib import Path
from typing import Callable, Iterable, Sequence
from urllib.parse import parse_qs, urlparse

from docx import Document
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.opc.constants import RELATIONSHIP_TYPE as RT
from docx.opc.exceptions import PackageNotFoundError
from docx.shared import Inches

if __package__ in {None, ""}:  # pragma: no cover - supports direct script execution.
    sys.path.insert(0, str(Path(__file__).resolve().parents[2]))

from procurement.console import configure_utf8_stdio
from procurement.external_tools import credential_free_media_environment
from procurement.input_limits import (
    DocxPackageFormatError,
    read_bounded_prefix,
    read_control_json,
    read_docx_snapshot,
)


# -----------------------------
# Settings
# -----------------------------

DEFAULT_EXTRACTOR_NAME = "extraction_router.py"
DEFAULT_LINK_COLUMN_TITLE = "Extraction folder"
DEFAULT_LINK_TEXT = "Open extraction folder"
DEFAULT_UNKNOWN_SPEAKER = "Unknown_Speaker"

# This catches normal visible YouTube links pasted directly into table cells.
YOUTUBE_URL_RE = re.compile(
    r"https?://(?:www\.|m\.|music\.)?(?:youtube\.com|youtu\.be|youtube-nocookie\.com)/[^\s<>'\"\])}]+",
    re.IGNORECASE,
)

# These folders are made by the scripts and should not be mistaken for video output folders.
NON_VIDEO_FOLDERS = {
    "logs",
    "video_timecodes",
    "video_timecodes_previous_runs",
    "raw_clips",
    "imotions_clips",
    "__pycache__",
    ".git",
    ".venv",
    "venv",
}

DOCX_RETRY_DELAYS_SECONDS = (0.25, 0.75)
MAX_DOCX_TABLES = 512
MAX_DOCX_ROWS = 100_000
MAX_DOCX_CELLS = 1_000_000
MAX_COMPLETION_JSON_BYTES = 64 * 1024
MAX_COMPLETION_JSON_ITEMS = 1024


# -----------------------------
# Small data object for each video row
# -----------------------------

@dataclass
class VideoRow:
    """
    Stores the location of a row in the DOCX file that contains a YouTube link.

    table_index and row_index are zero-based, because that is how python-docx
    accesses tables and rows internally.
    """

    table_index: int
    row_index: int
    url: str
    video_id: str
    speaker: str
    speaker_reason: str


@dataclass(frozen=True)
class SpeakerDecision:
    """The speaker folder decision made for one DOCX table row."""

    speaker: str | None
    reason: str


class BatchLogger:
    """Small wrapper-level log file for speaker routing and DOCX progress."""

    def __init__(self, base_folder: Path) -> None:
        logs_folder = base_folder / "logs"
        logs_folder.mkdir(exist_ok=True)
        timestamp = datetime.now().strftime("%Y-%m-%d_%H-%M-%S")
        self.path = logs_folder / f"docx_procurement_{timestamp}.log"
        self._handle = self.path.open("w", encoding="utf-8")

    def log(self, message: object = "") -> None:
        message_text = str(message)
        self._handle.write(message_text + "\n")
        self._handle.flush()
        try:
            print(message_text, flush=True)
        except (OSError, UnicodeError):
            pass

    def close(self) -> None:
        self._handle.close()


def open_docx_document(path: Path | str, *, logger: Callable[[str], None] | None = None) -> Document:
    """Open a DOCX, retrying once OneDrive/SharePoint has hydrated the file.

    Files from OneDrive for Business can exist on disk as reparse-point backed
    placeholders. The first Office-package open may fail even when the path
    exists. Reading a few bytes asks Windows/OneDrive to materialise the file,
    so this helper does that and retries before surfacing a clear error.
    """

    docx_path = Path(path).expanduser().resolve()
    last_error: Exception | None = None
    attempts = 1 + len(DOCX_RETRY_DELAYS_SECONDS)

    for attempt_index in range(attempts):
        try:
            snapshot = read_docx_snapshot(docx_path)
            return open_docx_snapshot(snapshot, docx_path)
        except (DocxPackageFormatError, PackageNotFoundError, PermissionError, OSError) as exc:
            last_error = exc
            try:
                prefix = read_bounded_prefix(docx_path)
            except OSError:
                prefix = b""

            if prefix and not prefix.startswith(b"PK"):
                break

            if attempt_index >= len(DOCX_RETRY_DELAYS_SECONDS):
                break

            delay = DOCX_RETRY_DELAYS_SECONDS[attempt_index]
            if logger:
                logger(
                    "DOCX was not ready on first open; retrying after OneDrive/SharePoint "
                    f"has had {delay:.2f}s to materialise it."
                )
            time.sleep(delay)

    message = (
        f"Could not open DOCX as a Word package: {docx_path}. "
        "If this file is stored in OneDrive or SharePoint, make sure it is available offline "
        "and not only a cloud placeholder, then try again."
    )
    raise RuntimeError(message) from last_error


def open_docx_snapshot(snapshot: bytes, path: Path | str) -> Document:
    """Parse and semantically validate one already ZIP-validated DOCX snapshot."""

    docx_path = Path(path).expanduser().resolve()
    document = Document(io.BytesIO(snapshot))
    validate_docx_semantic_limits(document, docx_path)
    return document


def validate_docx_semantic_limits(document: Document, path: Path) -> None:
    """Cap table, row, and cell proxy counts after the ZIP preflight succeeds."""

    tables = list(document.tables)
    if len(tables) > MAX_DOCX_TABLES:
        raise ValueError(f"DOCX contains more than {MAX_DOCX_TABLES} tables: {path}")
    row_count = 0
    cell_count = 0
    for table in tables:
        rows = list(table.rows)
        row_count += len(rows)
        if row_count > MAX_DOCX_ROWS:
            raise ValueError(f"DOCX contains more than {MAX_DOCX_ROWS} table rows: {path}")
        for row in rows:
            cell_count += len(row.cells)
            if cell_count > MAX_DOCX_CELLS:
                raise ValueError(f"DOCX contains more than {MAX_DOCX_CELLS} table cells: {path}")


# -----------------------------
# YouTube URL helpers
# -----------------------------

def clean_possible_url(url: str) -> str:
    """
    Removes common trailing punctuation that can accidentally stick to a URL.

    Example:
        https://youtu.be/abc123).  ->  https://youtu.be/abc123
    """
    return url.strip().rstrip(".,;:)]}")


def get_youtube_video_id(url: str) -> str | None:
    """
    Extracts the YouTube video ID from common YouTube URL formats.

    Supported examples:
    - https://www.youtube.com/watch?v=dQw4w9WgXcQ
    - https://youtu.be/dQw4w9WgXcQ
    - https://www.youtube.com/shorts/dQw4w9WgXcQ
    - https://www.youtube.com/embed/dQw4w9WgXcQ
    - https://www.youtube.com/live/dQw4w9WgXcQ
    """
    try:
        parsed = urlparse(clean_possible_url(url))
    except Exception:
        return None

    if parsed.scheme.casefold() not in {"http", "https"}:
        return None
    try:
        if parsed.username or parsed.password or parsed.port is not None:
            return None
    except ValueError:
        return None
    host = (parsed.hostname or "").lower()
    host = host.removeprefix("www.").removeprefix("m.").removeprefix("music.")

    # Short links: https://youtu.be/videoID
    if host == "youtu.be":
        parts = [part for part in parsed.path.split("/") if part]
        video_id = parts[0] if len(parts) == 1 else ""
        return video_id if re.fullmatch(r"[A-Za-z0-9_-]{11}", video_id) else None

    # Normal YouTube domains.
    if host in {"youtube.com", "youtube-nocookie.com"}:
        query_values = parse_qs(parsed.query)

        # Standard watch links: youtube.com/watch?v=videoID
        if "v" in query_values and len(query_values["v"]) == 1:
            video_id = query_values["v"][0]
            return video_id if re.fullmatch(r"[A-Za-z0-9_-]{11}", video_id) else None

        # Other common formats: /shorts/videoID, /embed/videoID, /live/videoID
        path_parts = [part for part in parsed.path.split("/") if part]

        for marker in ("shorts", "embed", "live", "v"):
            if marker in path_parts:
                marker_index = path_parts.index(marker)
                if marker_index + 1 < len(path_parts):
                    video_id = path_parts[marker_index + 1]
                    return video_id if re.fullmatch(r"[A-Za-z0-9_-]{11}", video_id) else None

    return None


def normalise_youtube_url(url: str) -> str | None:
    """
    Converts any supported YouTube URL format into a standard watch URL.

    This makes duplicate detection easier. For example, these should be treated
    as the same video:

        https://youtu.be/dQw4w9WgXcQ
        https://www.youtube.com/watch?v=dQw4w9WgXcQ&t=10s
    """
    video_id = get_youtube_video_id(url)

    if not video_id:
        return None

    return f"https://www.youtube.com/watch?v={video_id}"


# -----------------------------
# Speaker routing helpers
# -----------------------------

def make_folder_name_safe(text: str, max_length: int = 80) -> str:
    """Convert a speaker label into a stable folder name."""
    cleaned = str(text).strip().replace("\u00a0", " ")
    cleaned = re.sub(r'[<>:"/\\|?*]', "", cleaned)
    cleaned = re.sub(r"\s+", "_", cleaned)
    cleaned = re.sub(r"_+", "_", cleaned)
    cleaned = cleaned.strip("._ ")
    return (cleaned or DEFAULT_UNKNOWN_SPEAKER)[:max_length]


def speaker_cell_text(row) -> str:
    """Return the second table column, which is the speaker column in this workbook."""
    if len(row.cells) < 2:
        return ""
    return row.cells[1].text.strip()


def infer_speaker_name(cell_text: str) -> SpeakerDecision:
    """
    Extract the likely speaker name from the second DOCX table column.

    The source document sometimes mixes the name with notes such as
    "voice-over" or "with few extended views of face". This function keeps the
    name portion and leaves unclear descriptor-only cells for the neighbor pass.
    """
    original = str(cell_text or "").replace("\u00a0", " ").strip()
    cleaned = re.sub(r"\s+", " ", original)
    if not cleaned:
        return SpeakerDecision(None, "empty speaker cell")

    lowered = cleaned.lower()
    descriptor_only = [
        "speaker",
        "n/a",
        "na",
        "unknown",
        "voice-over",
        "voice over",
        "voiceover",
        "voice-over with few extended views of face",
    ]
    if lowered in descriptor_only or "few extended views of face" in lowered and " " not in cleaned.replace("-", " "):
        return SpeakerDecision(None, f"descriptor-only speaker cell: {cleaned}")

    # Treat spaced dashes as note separators, but preserve hyphenated names.
    candidate = re.split(r"\s+(?:-|\u2013|\u2014)\s+", cleaned, maxsplit=1)[0].strip()
    candidate = re.split(r"\s+with\s+", candidate, maxsplit=1, flags=re.IGNORECASE)[0].strip()
    candidate = re.sub(r"\([^)]*\)", "", candidate).strip()
    candidate = re.sub(r"\bvoice[- ]?over\b.*$", "", candidate, flags=re.IGNORECASE).strip()
    candidate = re.sub(r"\bnarration\b.*$", "", candidate, flags=re.IGNORECASE).strip()
    candidate = re.sub(r"\s+", " ", candidate)

    if not candidate or candidate.lower() in descriptor_only:
        return SpeakerDecision(None, f"no speaker name identified in: {cleaned}")
    if not re.search(r"[A-Za-z]", candidate):
        return SpeakerDecision(None, f"speaker cell has no alphabetic name: {cleaned}")

    reason = "exact second-column speaker" if candidate == cleaned else f"cleaned second-column speaker from: {cleaned}"
    return SpeakerDecision(candidate, reason)


def resolve_speaker_names(raw_speaker_cells: Sequence[str]) -> list[SpeakerDecision]:
    """
    Resolve speaker names for a table, filling unclear rows from matching neighbors.

    If a row has no clear speaker but the immediately previous and next rows
    resolve to the same speaker, the row is assigned to that speaker.
    """
    decisions = [infer_speaker_name(value) for value in raw_speaker_cells]
    resolved = list(decisions)
    for index, decision in enumerate(decisions):
        if decision.speaker:
            continue
        if index == 0 or index >= len(decisions) - 1:
            continue
        previous_speaker = decisions[index - 1].speaker
        next_speaker = decisions[index + 1].speaker
        if previous_speaker and previous_speaker == next_speaker:
            resolved[index] = SpeakerDecision(
                previous_speaker,
                f"inferred from matching neighboring speaker: {previous_speaker}",
            )
    return resolved


# -----------------------------
# DOCX hyperlink reading helpers
# -----------------------------

def extract_visible_urls_from_text(text: str) -> list[str]:
    """
    Finds YouTube URLs that are visibly written as text inside a cell.
    """
    return [clean_possible_url(match.group(0)) for match in YOUTUBE_URL_RE.finditer(text or "")]


def extract_word_hyperlink_targets_from_cell(cell, document_part) -> list[str]:
    """
    Finds actual DOCX hyperlink targets inside a cell.

    This matters because many DOCX files display a readable title such as:

        My video title

    while the real YouTube URL is hidden inside the hyperlink relationship.
    cell.text only gives us the visible title, so we need to inspect the DOCX XML.
    """
    targets = []

    # Normal Word hyperlinks are stored as <w:hyperlink r:id="...">.
    for hyperlink in cell._tc.xpath(".//w:hyperlink"):
        relationship_id = hyperlink.get(qn("r:id"))

        if relationship_id and relationship_id in document_part.rels:
            targets.append(document_part.rels[relationship_id].target_ref)

    # Some documents store hyperlinks as field codes instead of normal hyperlink tags.
    # This catches field-code links like: HYPERLINK "https://www.youtube.com/..."
    for instr_text in cell._tc.xpath(".//w:instrText"):
        field_text = instr_text.text or ""
        for match in re.finditer(r'HYPERLINK\s+"([^"]+)"', field_text, flags=re.IGNORECASE):
            targets.append(match.group(1))

    return targets


def get_youtube_links_from_row(row, document_part) -> list[str]:
    """
    Returns all YouTube links found anywhere in a table row.

    It checks:
    1. Visible pasted URLs in the text.
    2. Hidden Word hyperlink targets.

    It does not care which column the link is in.
    """
    found_urls = []
    seen_video_ids = set()

    for cell in row.cells:
        candidates = []
        candidates.extend(extract_visible_urls_from_text(cell.text))
        candidates.extend(extract_word_hyperlink_targets_from_cell(cell, document_part))

        for candidate in candidates:
            normalised_url = normalise_youtube_url(candidate)
            video_id = get_youtube_video_id(candidate)

            if normalised_url and video_id and video_id not in seen_video_ids:
                found_urls.append(normalised_url)
                seen_video_ids.add(video_id)

    return found_urls


# -----------------------------
# DOCX hyperlink writing helpers
# -----------------------------

def clear_cell(cell) -> None:
    """
    Clears a table cell so we can insert a clean hyperlink.
    """
    cell.text = ""


def add_hyperlink(paragraph, link_text: str, link_target: str) -> None:
    """
    Adds a clickable hyperlink to a Word paragraph.

    link_target can be a normal URL, a file path, or a folder path.
    For this script, we use a relative folder path such as:

        Some_Video_Title_[dQw4w9WgXcQ]/
    """
    relationship_id = paragraph.part.relate_to(
        link_target,
        RT.HYPERLINK,
        is_external=True,
    )

    hyperlink = OxmlElement("w:hyperlink")
    hyperlink.set(qn("r:id"), relationship_id)

    run = OxmlElement("w:r")

    # Apply Word's built-in Hyperlink style if it exists.
    run_properties = OxmlElement("w:rPr")
    run_style = OxmlElement("w:rStyle")
    run_style.set(qn("w:val"), "Hyperlink")
    run_properties.append(run_style)
    run.append(run_properties)

    text = OxmlElement("w:t")
    text.text = link_text
    run.append(text)

    hyperlink.append(run)
    paragraph._p.append(hyperlink)


def make_docx_relationship_target_safe(target: str) -> str:
    """
    Encode a relative hyperlink target so Microsoft Word treats it as a URI.

    The extraction folders are often based on video titles, which can contain
    accents, apostrophes, non-breaking spaces, or emoji. Those are legal on the
    filesystem but can make Word reject the whole DOCX if they are written raw
    into a hyperlink relationship. Percent-encoding keeps the folder link
    relative while making the OOXML relationship Word-safe.
    """
    return urllib.parse.quote(str(target), safe="/._-[]")


def row_looks_like_header(row) -> bool:
    """
    Heuristically decides whether a row is a table header row.

    This avoids assuming a fixed table structure. It simply looks for common
    source-table headings like Link, Speaker, Date, Engagement, Length, etc.
    """
    row_text = " | ".join(cell.text.strip().lower() for cell in row.cells)

    header_words = [
        "link",
        "speaker",
        "upload date",
        "date uploaded",
        "date accessed",
        "engagement",
        "time stamps",
        "timestamps",
        "length",
    ]

    return sum(1 for word in header_words if word in row_text) >= 2


def table_already_has_extraction_column(table) -> bool:
    """
    Avoids adding duplicate link columns if the script is run on an already-edited DOCX.
    """
    if not table.rows:
        return False

    for row in table.rows[: min(5, len(table.rows))]:
        if not row.cells:
            continue

        last_cell_text = row.cells[-1].text.strip().lower()

        if (
            last_cell_text == DEFAULT_LINK_COLUMN_TITLE.lower()
            or DEFAULT_LINK_TEXT.lower() in last_cell_text
            or "extraction folder" in last_cell_text
            or "failed" in last_cell_text
        ):
            return True

    return False


def ensure_extraction_column(table) -> int:
    """
    Makes sure a table has a final column for extraction-folder links.

    Returns the index of the link column.
    """
    if table_already_has_extraction_column(table):
        return len(table.rows[0].cells) - 1

    # Add a new column to the right-hand side of the table.
    table.add_column(Inches(1.5))
    link_column_index = len(table.rows[0].cells) - 1

    # If the first row is clearly a header, label the new column.
    # Some tables in messy documents may not have a header row, so we avoid forcing one.
    if table.rows and row_looks_like_header(table.rows[0]):
        table.rows[0].cells[link_column_index].text = DEFAULT_LINK_COLUMN_TITLE

    return link_column_index


def apply_numeric_prefix(folder: Path, index: int) -> Path:
    """Rename a download folder to add a global sequential prefix (e.g. 003_Title_[id]).

    Strips any existing three-digit prefix first so re-runs with different
    ordering produce the correct name rather than double-prefixing.
    """
    prefix = f"{index:03d}_"
    if folder.name.startswith(prefix):
        return folder
    stem = re.sub(r"^\d{3}_", "", folder.name)
    new_path = folder.parent / f"{prefix}{stem}"
    if new_path != folder:
        folder.rename(new_path)
    return new_path


def add_folder_link_to_row(row, folder_path: Path, output_docx_path: Path) -> None:
    """
    Adds a relative hyperlink to the final cell of a table row.
    """
    link_cell = row.cells[-1]
    clear_cell(link_cell)

    # Work out the folder path relative to where the edited DOCX will be saved.
    relative_target = os.path.relpath(folder_path, start=output_docx_path.parent)
    relative_target = relative_target.replace(os.sep, "/")

    # A trailing slash helps Word understand that this points to a folder.
    if not relative_target.endswith("/"):
        relative_target += "/"

    paragraph = link_cell.paragraphs[0]
    add_hyperlink(paragraph, DEFAULT_LINK_TEXT, make_docx_relationship_target_safe(relative_target))


def write_failure_to_row(row, message: str) -> None:
    """
    Writes a clear failure message into the final cell of a row.
    """
    link_cell = row.cells[-1]
    clear_cell(link_cell)
    link_cell.text = message


# -----------------------------
# Video extraction helpers
# -----------------------------

def folder_contains_completed_extraction(
    folder: Path,
    expected_request: dict[str, object] | None = None,
) -> bool:
    """
    Returns True when a folder looks like a finished current-format output.

    A folder is reusable only when the extractor wrote its completion marker.
    Raw clips alone are not enough, because they might come from a failed or
    interrupted run.
    """
    completion_path = folder / "_extraction_complete.json"
    if not completion_path.exists() or not completion_path.is_file():
        return False

    try:
        completion = read_control_json(
            completion_path,
            label="completion marker",
            max_bytes=MAX_COMPLETION_JSON_BYTES,
            max_items=MAX_COMPLETION_JSON_ITEMS,
        )
    except (OSError, json.JSONDecodeError):
        return False

    if not isinstance(completion, dict):
        return False

    if completion.get("status") != "success":
        return False

    if expected_request is not None:
        request_path = folder / "_docx_extraction_request.json"
        try:
            prior_request = read_control_json(
                request_path,
                label="extraction request",
                max_bytes=MAX_COMPLETION_JSON_BYTES,
                max_items=MAX_COMPLETION_JSON_ITEMS,
            )
        except (OSError, json.JSONDecodeError):
            return False
        if prior_request != expected_request:
            return False

    skip_stitch = bool(completion.get("skip_stitch") or completion.get("skip_imotions_conversion"))
    stitched_video = folder / "stitched_imotions.mp4"
    if not skip_stitch:
        return stitched_video.exists() and stitched_video.is_file() and stitched_video.stat().st_size > 0

    raw_clip_folder = folder / "raw_clips"
    if raw_clip_folder.exists():
        for file_path in raw_clip_folder.rglob("*.mp4"):
            if file_path.is_file() and file_path.stat().st_size > 0:
                return True

    return False


def find_existing_output_folder(
    base_folder: Path,
    video_id: str,
    expected_request: dict[str, object] | None = None,
) -> Path | None:
    """
    Looks for an existing completed folder created by extractor.py for this video ID.

    The expected folder naming pattern is:

        Video_Title_[videoID]

    The exact title can vary slightly, so we search for [videoID] instead of
    rebuilding the title ourselves.
    """
    candidates = []

    for child in base_folder.iterdir():
        if not child.is_dir():
            continue

        if child.name in NON_VIDEO_FOLDERS:
            continue

        if f"[{video_id}]" in child.name and folder_contains_completed_extraction(
            child,
            expected_request,
        ):
            candidates.append(child)

    if not candidates:
        return None

    # If more than one completed folder exists, use the newest one.
    return max(candidates, key=lambda path: path.stat().st_mtime)


def run_video_extractor(
    extractor_script: Path,
    url: str,
    working_folder: Path,
    extra_extractor_args: Iterable[str],
) -> None:
    """
    Calls the extractor or router for one YouTube URL.

    Using sys.executable makes it run with the same Python interpreter that is
    running this wrapper script.
    """
    command = [
        sys.executable,
        str(extractor_script),
        url,
    ]
    if extractor_script.name.lower() == DEFAULT_EXTRACTOR_NAME.lower():
        command.extend(["--output-root", str(working_folder)])
    command.extend(extra_extractor_args)

    print("\n" + "-" * 70, flush=True)
    print(f"Running extractor for: {url}", flush=True)
    print(f"Output root: {working_folder}", flush=True)
    print("Command:", " ".join(f'"{part}"' if " " in str(part) else str(part) for part in command), flush=True)
    print("-" * 70, flush=True)

    subprocess.run(
        command,
        cwd=working_folder,
        check=True,
        env=credential_free_media_environment(),
    )


def extract_or_reuse_folder(
    url: str,
    video_id: str,
    extractor_script: Path,
    working_folder: Path,
    force: bool,
    extra_extractor_args: Iterable[str],
) -> Path:
    """
    Either reuses an existing extraction folder or runs extractor.py.
    """
    extractor_args = [str(value) for value in extra_extractor_args]
    request_fingerprint = {
        "video_id": video_id,
        "url": url,
        "extractor": str(extractor_script.expanduser().resolve()),
        "arguments": extractor_args,
    }
    existing_folder = find_existing_output_folder(
        working_folder,
        video_id,
        request_fingerprint,
    )

    if existing_folder and not force:
        print(f"Reusing existing folder for {video_id}: {existing_folder.name}", flush=True)
        return existing_folder

    run_video_extractor(
        extractor_script=extractor_script,
        url=url,
        working_folder=working_folder,
        extra_extractor_args=extractor_args,
    )

    output_folder = find_existing_output_folder(working_folder, video_id)

    if not output_folder:
        raise FileNotFoundError(
            f"The video extractor finished, but no completed output folder containing [{video_id}] was found."
        )

    (output_folder / "_docx_extraction_request.json").write_text(
        json.dumps(request_fingerprint, indent=2, ensure_ascii=False) + "\n",
        encoding="utf-8",
    )

    return output_folder


# -----------------------------
# Scanning the DOCX
# -----------------------------

def find_video_rows(document: Document) -> list[VideoRow]:
    """
    Scans every table and every row in the DOCX for YouTube links.

    This does not assume fixed table numbers, fixed row counts, or fixed columns.
    """
    found_rows = []

    for table_index, table in enumerate(document.tables):
        # python-docx rebuilds its row proxy list whenever ``table.rows`` is
        # accessed. Cache it once so large source documents remain linear.
        table_rows = list(table.rows)
        speaker_decisions = resolve_speaker_names([speaker_cell_text(row) for row in table_rows])
        for row_index, row in enumerate(table_rows):
            youtube_links = get_youtube_links_from_row(row, document.part)

            if not youtube_links:
                continue

            # If a row somehow has more than one YouTube link, use the first one.
            # This matches the source-table structure where each row is one video.
            url = youtube_links[0]
            video_id = get_youtube_video_id(url)

            if video_id:
                speaker_decision = speaker_decisions[row_index]
                speaker = speaker_decision.speaker or DEFAULT_UNKNOWN_SPEAKER
                found_rows.append(
                    VideoRow(
                        table_index=table_index,
                        row_index=row_index,
                        url=url,
                        video_id=video_id,
                        speaker=speaker,
                        speaker_reason=speaker_decision.reason,
                    )
                )

    return found_rows


def filter_video_rows_by_speaker(
    video_rows: Sequence[VideoRow],
    selected_speakers: Sequence[str] | None,
) -> list[VideoRow]:
    """Keep only exact, case-insensitive speaker labels selected by the user."""

    requested = {speaker_match_key(value) for value in selected_speakers or [] if speaker_match_key(value)}
    if not requested:
        return list(video_rows)
    return [row for row in video_rows if speaker_match_key(row.speaker) in requested]


def speaker_match_key(value: object) -> str:
    """Return a stable comparison key for human-entered speaker labels."""

    return " ".join(str(value or "").split()).casefold()


# -----------------------------
# Main script
# -----------------------------

def main() -> int:
    configure_utf8_stdio()
    parser = argparse.ArgumentParser(
        description=(
            "Read every table in a DOCX, run the video extractor for each YouTube video, "
            "and add relative links to the extracted-video folders."
        )
    )

    parser.add_argument(
        "docx_path",
        help="Path to the DOCX file containing the video tables.",
    )

    parser.add_argument(
        "--output",
        default=None,
        help=(
            "Optional output DOCX path. By default, saves beside the original as "
            "<original_name>_with_extraction_links.docx."
        ),
    )

    parser.add_argument(
        "--extractor",
        default=DEFAULT_EXTRACTOR_NAME,
        help="Name/path of the extractor script. Default: extraction_router.py",
    )

    parser.add_argument(
        "--force",
        action="store_true",
        help="Run the extractor even if a completed output folder for that video ID already exists.",
    )

    parser.add_argument(
        "--limit",
        type=int,
        default=None,
        help="Only process the first N found videos. Useful for testing.",
    )

    parser.add_argument(
        "--speaker",
        action="append",
        default=[],
        help="Process only this resolved speaker label. May be repeated.",
    )

    parser.add_argument(
        "--dry-run",
        action="store_true",
        help="Only scan the document and print what would be processed. Does not download or save a DOCX.",
    )

    parser.add_argument(
        "--autosave-every",
        type=int,
        default=10,
        help="Save progress every N processed videos. Default: 10. Use 0 to disable autosave.",
    )

    parser.add_argument(
        "--speaker-output-root",
        default=None,
        help=(
            "Folder where speaker-named output folders are created. "
            "Defaults to the folder containing this script."
        ),
    )

    parser.add_argument(
        "--no-stitch",
        action="store_true",
        help="Download the 10 percent sample as raw clips only by passing --skip-stitch to the extractor.",
    )

    parser.add_argument(
        "--extractor-arg",
        action="append",
        default=[],
        help=(
            "Extra argument to pass through to the video extractor. "
            "Can be used multiple times, e.g. --extractor-arg --seed --extractor-arg 123"
        ),
    )

    args = parser.parse_args()

    script_folder = Path(__file__).resolve().parent
    batch_logger = BatchLogger(script_folder)
    docx_path = Path(args.docx_path).resolve()
    speaker_output_root = Path(args.speaker_output_root).resolve() if args.speaker_output_root else script_folder
    speaker_output_root.mkdir(parents=True, exist_ok=True)
    extra_extractor_args = list(args.extractor_arg)
    if args.no_stitch:
        extra_extractor_args.append("--skip-stitch")

    if not docx_path.exists():
        raise FileNotFoundError(f"DOCX not found: {docx_path}")

    extractor_script = Path(args.extractor)
    if not extractor_script.is_absolute():
        extractor_script = script_folder / extractor_script

    if not args.dry_run and not extractor_script.exists():
        raise FileNotFoundError(
            f"Extractor script not found: {extractor_script}\n"
            "Make sure extraction_router.py and extractor.py are in the same folder as this script."
        )

    if args.output:
        output_docx_path = Path(args.output).resolve()
    else:
        output_docx_path = docx_path.with_name(f"{docx_path.stem}_with_extraction_links.docx")
    output_docx_path.parent.mkdir(parents=True, exist_ok=True)

    # For relative folder links to work cleanly, the edited DOCX should usually
    # live beside the extracted folders. This warning does not stop the script.
    if output_docx_path.parent != script_folder:
        batch_logger.log(
            "WARNING: The output DOCX is not being saved in the same folder as this script.\n"
            "Relative folder links will be calculated relative to the output DOCX location.\n"
        )

    batch_logger.log(f"Reading DOCX: {docx_path}")
    batch_logger.log(f"Speaker output root: {speaker_output_root}")
    batch_logger.log(f"Stitching enabled: {not args.no_stitch}")
    document = open_docx_document(docx_path, logger=batch_logger.log)

    batch_logger.log(f"Tables found: {len(document.tables)}")
    video_rows = find_video_rows(document)
    video_rows = filter_video_rows_by_speaker(video_rows, args.speaker)
    if args.speaker and not video_rows:
        raise ValueError("No DOCX videos matched the selected speaker groups.")

    if args.limit is not None:
        video_rows = video_rows[: args.limit]

    batch_logger.log(f"YouTube video rows found: {len(video_rows)}")

    if not video_rows:
        batch_logger.log("No YouTube links were found in the document tables.")
        batch_logger.log(f"Wrapper log saved as: {batch_logger.path}")
        batch_logger.close()
        return 1

    batch_logger.log("\nFirst few detected videos:")
    for item in video_rows[:10]:
        batch_logger.log(
            f"  Table {item.table_index + 1}, row {item.row_index + 1}: "
            f"{item.video_id}, speaker {item.speaker} -> {item.url}"
        )

    if args.dry_run:
        batch_logger.log("\nDry run complete. No videos were downloaded and no DOCX was saved.")
        batch_logger.log(f"Wrapper log saved as: {batch_logger.path}")
        batch_logger.close()
        return 0

    # Add a final link column to every table that actually contains a YouTube video row.
    tables_to_update = sorted({item.table_index for item in video_rows})
    for table_index in tables_to_update:
        ensure_extraction_column(document.tables[table_index])

    # Cache means duplicate YouTube videos only get extracted once.
    # Multiple table rows can then link to the same output folder.
    folder_cache_by_video_id: dict[tuple[str, str], Path] = {}

    processed_count = 0
    failed_count = 0

    for item_number, item in enumerate(video_rows, start=1):
        row = document.tables[item.table_index].rows[item.row_index]
        speaker_folder = speaker_output_root / make_folder_name_safe(item.speaker)
        speaker_folder.mkdir(parents=True, exist_ok=True)
        cache_key = (item.video_id, speaker_folder.name)

        batch_logger.log("\n" + "=" * 70)
        batch_logger.log(
            f"[{item_number}/{len(video_rows)}] "
            f"Table {item.table_index + 1}, row {item.row_index + 1}, video ID {item.video_id}"
        )
        batch_logger.log(f"URL: {item.url}")
        batch_logger.log(f"Speaker: {item.speaker} ({item.speaker_reason})")
        batch_logger.log(f"Speaker folder: {speaker_folder}")
        batch_logger.log("=" * 70)

        try:
            if cache_key in folder_cache_by_video_id and not args.force:
                output_folder = folder_cache_by_video_id[cache_key]
                batch_logger.log(f"Duplicate video ID/speaker found. Reusing: {output_folder.name}")
            else:
                output_folder = extract_or_reuse_folder(
                    url=item.url,
                    video_id=item.video_id,
                    extractor_script=extractor_script,
                    working_folder=speaker_folder,
                    force=args.force,
                    extra_extractor_args=extra_extractor_args,
                )
                folder_cache_by_video_id[cache_key] = output_folder

            output_folder = apply_numeric_prefix(output_folder, item_number)
            folder_cache_by_video_id[cache_key] = output_folder
            add_folder_link_to_row(
                row=row,
                folder_path=output_folder,
                output_docx_path=output_docx_path,
            )

            processed_count += 1
            batch_logger.log(f"Linked row to folder: {output_folder}")

        except Exception as error:
            failed_count += 1
            error_message = f"FAILED: {type(error).__name__}"
            write_failure_to_row(row, error_message)
            batch_logger.log(f"ERROR while processing video {item.video_id}: {error}")
            batch_logger.log("The script will continue with the next video.")

        # Save progress occasionally so a long run does not lose all DOCX edits.
        if args.autosave_every and processed_count > 0 and processed_count % args.autosave_every == 0:
            document.save(str(output_docx_path))
            batch_logger.log(f"Autosaved progress to: {output_docx_path}")

    document.save(str(output_docx_path))

    batch_logger.log("\n" + "=" * 70)
    batch_logger.log("Finished.")
    batch_logger.log(f"Successful rows linked: {processed_count}")
    batch_logger.log(f"Failed rows marked: {failed_count}")
    batch_logger.log(f"Edited DOCX saved as: {output_docx_path}")
    batch_logger.log(f"Wrapper log saved as: {batch_logger.path}")
    batch_logger.log("=" * 70)
    batch_logger.close()
    return 1 if failed_count else 0


if __name__ == "__main__":
    raise SystemExit(main())
