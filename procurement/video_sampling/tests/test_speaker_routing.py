from pathlib import Path
import sys
from tempfile import TemporaryDirectory
import unittest
from unittest.mock import patch
from zipfile import ZIP_DEFLATED, ZipFile

from docx import Document
from docx.opc.exceptions import PackageNotFoundError

from procurement.video_sampling import run_docx_extractions
from procurement.video_sampling.run_docx_extractions import (
    VideoRow,
    filter_video_rows_by_speaker,
    infer_speaker_name,
    resolve_speaker_names,
    speaker_match_key,
)


class SpeakerRoutingTests(unittest.TestCase):
    def test_standard_extractor_child_does_not_inherit_media_credentials(self):
        captured: dict[str, object] = {}
        secret_names = ("YOUTUBE_API_KEY", "HF_TOKEN", "HUGGINGFACE_TOKEN", "HUGGING_FACE_HUB_TOKEN")

        def fake_run(*_args, **kwargs):
            captured.update(kwargs)

        with TemporaryDirectory() as temp_dir, patch.dict(
            run_docx_extractions.os.environ,
            {name: f"secret-{name}" for name in secret_names},
            clear=False,
        ), patch.object(run_docx_extractions.subprocess, "run", side_effect=fake_run):
            run_docx_extractions.run_video_extractor(
                Path(temp_dir) / "extraction_router.py",
                "https://example.invalid/watch?v=aaaaaaaaaaa",
                Path(temp_dir),
                [],
            )

        environment = captured.get("env")
        self.assertIsInstance(environment, dict)
        self.assertTrue(all(name not in environment for name in secret_names))

    def test_speaker_match_key_normalizes_case_and_repeated_whitespace(self):
        self.assertEqual(speaker_match_key("  Speaker   A "), "speaker a")

    def test_filter_video_rows_uses_normalized_exact_speaker_keys(self):
        rows = [
            VideoRow(0, 1, "https://youtu.be/aaaaaaaaaaa", "aaaaaaaaaaa", "Speaker A", "cell"),
            VideoRow(0, 2, "https://youtu.be/bbbbbbbbbbb", "bbbbbbbbbbb", "Speaker B", "cell"),
        ]

        selected = filter_video_rows_by_speaker(rows, [" speaker   a "])

        self.assertEqual(selected, [rows[0]])

    def test_infer_speaker_name_removes_descriptive_suffixes(self):
        decision = infer_speaker_name("Speaker A - voice-over")

        self.assertEqual(decision.speaker, "Speaker A")
        self.assertIn("cleaned", decision.reason)

    def test_infer_speaker_name_keeps_multiple_named_speakers(self):
        decision = infer_speaker_name("Speaker A and Speaker B - voice-over")

        self.assertEqual(decision.speaker, "Speaker A and Speaker B")

    def test_infer_speaker_name_keeps_hyphenated_given_names(self):
        decision = infer_speaker_name("Speaker-Alpha Example")

        self.assertEqual(decision.speaker, "Speaker-Alpha Example")

    def test_resolve_speaker_names_uses_matching_neighbors_for_unclear_cells(self):
        decisions = resolve_speaker_names(
            [
                "Speaker B",
                "Voice-over with few extended views of face",
                "Speaker B",
            ]
        )

        self.assertEqual(decisions[1].speaker, "Speaker B")
        self.assertIn("matching neighboring speaker", decisions[1].reason)

    def test_open_docx_document_retries_after_placeholder_open_failure(self):
        fake_document = object()
        messages: list[str] = []

        with TemporaryDirectory() as temp_dir:
            docx_path = Path(temp_dir) / "input.docx"
            docx_path.write_bytes(b"PK\x03\x04placeholder")

            with (
                patch.object(
                    run_docx_extractions,
                    "Document",
                    side_effect=[PackageNotFoundError("package not found"), fake_document],
                ) as document_mock,
                patch.object(run_docx_extractions, "read_docx_snapshot", return_value=b"snapshot"),
                patch.object(run_docx_extractions, "validate_docx_semantic_limits"),
                patch.object(run_docx_extractions.time, "sleep") as sleep_mock,
            ):
                document = run_docx_extractions.open_docx_document(docx_path, logger=messages.append)

        self.assertIs(document, fake_document)
        self.assertEqual(document_mock.call_count, 2)
        sleep_mock.assert_called_once()
        self.assertTrue(any("retrying" in message for message in messages))

    def test_open_docx_document_rejects_excessive_archive_entries(self):
        with TemporaryDirectory() as temp_dir:
            docx_path = Path(temp_dir) / "entry-bomb.docx"
            document = Document()
            document.add_paragraph("safe")
            document.save(docx_path)
            with ZipFile(docx_path, "a", compression=ZIP_DEFLATED) as archive:
                for index in range(2050):
                    archive.writestr(f"custom/item-{index}.txt", "x")

            with self.assertRaisesRegex(ValueError, "entry count"):
                run_docx_extractions.open_docx_document(docx_path)

    def test_open_docx_document_parses_the_same_snapshot_that_was_validated(self):
        with TemporaryDirectory() as temp_dir:
            root = Path(temp_dir)
            docx_path = root / "input.docx"
            replacement_path = root / "replacement.docx"
            benign = Document()
            benign.add_paragraph("benign snapshot")
            benign.save(docx_path)
            replacement = Document()
            replacement.add_paragraph("swapped package")
            replacement.save(replacement_path)
            replacement_bytes = replacement_path.read_bytes()
            original_document = Document

            def swap_before_parse(source):
                docx_path.write_bytes(replacement_bytes)
                return original_document(source)

            with patch.object(run_docx_extractions, "Document", side_effect=swap_before_parse):
                opened = run_docx_extractions.open_docx_document(docx_path)

        self.assertEqual([paragraph.text for paragraph in opened.paragraphs], ["benign snapshot"])

    def test_open_docx_document_rejects_excessive_compression_ratio(self):
        with TemporaryDirectory() as temp_dir:
            docx_path = Path(temp_dir) / "ratio-bomb.docx"
            document = Document()
            document.add_paragraph("safe")
            document.save(docx_path)
            with ZipFile(docx_path, "a", compression=ZIP_DEFLATED) as archive:
                archive.writestr("custom/repeated.bin", b"A" * (1024 * 1024))

            with self.assertRaisesRegex(ValueError, "compression ratio"):
                run_docx_extractions.open_docx_document(docx_path)

    def test_open_docx_retry_uses_bounded_prefix_read(self):
        fake_document = object()
        with TemporaryDirectory() as temp_dir:
            docx_path = Path(temp_dir) / "input.docx"
            docx_path.write_bytes(b"PK\x03\x04placeholder")
            with (
                patch.object(
                    run_docx_extractions,
                    "Document",
                    side_effect=[PackageNotFoundError("package not found"), fake_document],
                ),
                patch.object(run_docx_extractions, "read_docx_snapshot", return_value=b"snapshot"),
                patch.object(run_docx_extractions, "validate_docx_semantic_limits"),
                patch.object(Path, "read_bytes", side_effect=AssertionError("unbounded read")),
                patch.object(run_docx_extractions.time, "sleep"),
            ):
                result = run_docx_extractions.open_docx_document(docx_path)

        self.assertIs(result, fake_document)

    def test_main_creates_output_parent_before_saving_linked_docx(self):
        with TemporaryDirectory() as temp_dir:
            temp_path = Path(temp_dir)
            docx_path = temp_path / "input.docx"
            document = Document()
            table = document.add_table(rows=2, cols=2)
            table.rows[0].cells[0].text = "Link"
            table.rows[0].cells[1].text = "Speaker"
            table.rows[1].cells[0].text = "https://www.youtube.com/watch?v=abcdefghijk"
            table.rows[1].cells[1].text = "Speaker Name"
            document.save(docx_path)

            fake_output_folder = temp_path / "downloads" / "Speaker_Name" / "Video_[abcdefghijk]"
            fake_output_folder.mkdir(parents=True)
            output_docx = temp_path / "missing" / "nested" / "linked.docx"
            argv = [
                "run_docx_extractions.py",
                str(docx_path),
                "--limit",
                "1",
                "--speaker-output-root",
                str(temp_path / "downloads"),
                "--output",
                str(output_docx),
            ]

            with (
                patch.object(sys, "argv", argv),
                patch.object(run_docx_extractions, "extract_or_reuse_folder", return_value=fake_output_folder),
            ):
                run_docx_extractions.main()

            self.assertTrue(output_docx.exists())


if __name__ == "__main__":
    unittest.main()
