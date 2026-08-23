#!/usr/bin/env python3
"""
complete_youtube_license_audit_docx.py

Team-friendly YouTube licence audit for Word (.docx) source tables.

This script is designed for non-technical research workflows where future team
members need to run a repeatable licence audit without editing Python code.

What it does
------------
1. Opens a DOCX file containing video-source tables.
2. Finds YouTube video links in the first column of each video table.
   - Handles visible URLs.
   - Handles hidden Word hyperlinks where the cell displays only the video title.
   - Handles headerless tables by optionally inserting a clean header row.
3. Queries YouTube's official Data API field:
       status.license
   Known values:
       creativeCommon -> Creative Commons Attribution (CC BY)
       youtube         -> Standard YouTube License
4. Separately scans the YouTube title/description/tags for licence-related words.
   This does NOT override the YouTube platform licence. It only creates manual
   review notes when the text appears to conflict or is ambiguous.
5. Adds or updates only two clean columns in the DOCX file:
       License
       Date Checked
6. Writes the detailed evidence to a debug CSV, summary TXT, and optional log DOCX.

This script is not legal advice. It records YouTube metadata and text signals at
the time of checking so a research team can make consistent decisions.
"""

from __future__ import annotations

import argparse
import csv
import datetime as dt
import json
import os
import re
import sys
import time
import urllib.parse
import urllib.request
from dataclasses import dataclass, asdict
from pathlib import Path
from typing import Any, Iterable

if __package__ in {None, ""}:  # pragma: no cover - direct script execution.
    sys.path.insert(0, str(Path(__file__).resolve().parents[2]))

from procurement.video_sampling.run_docx_extractions import open_docx_document
from procurement.input_limits import read_term_dictionary
from spreadsheet_safety import SpreadsheetSafeWriter

try:
    from docx import Document
    from docx.oxml.ns import qn
    from docx.shared import Cm
except ImportError as exc:
    raise SystemExit(
        "ERROR: python-docx is not installed. Run: python -m pip install python-docx"
    ) from exc

SCRIPT_VERSION = "2026.05.18-team-setup-v2-minimal-doc-columns"
YOUTUBE_VIDEOS_URL = "https://www.googleapis.com/youtube/v3/videos"

# Columns written into the user's main DOCX file.
# Keep the document itself clean: detailed evidence goes to the CSV/log DOCX.
FINAL_LICENSE_HEADER = "License"
CHECKED_DATE_HEADER = "Date Checked"

# Detailed fields used in the log/debug outputs, not appended to the main DOCX.
API_LICENSE_HEADER = "YouTube API License"
RAW_API_LICENSE_HEADER = "Raw YouTube API License"
DESCRIPTION_SIGNALS_HEADER = "Description License Signals"
LICENSE_NOTES_HEADER = "License Notes"
MANUAL_REVIEW_HEADER = "Needs manual review"
SCRIPT_VERSION_HEADER = "Script version"

OUTPUT_COLUMNS = [
    (FINAL_LICENSE_HEADER, 4.6),
    (CHECKED_DATE_HEADER, 3.2),
]

# If someone accidentally runs this on an older audit output, remove the old
# detailed audit columns so the final DOCX still only contains License + Date Checked.
DETAILED_AUDIT_HEADERS_TO_REMOVE = {
    "youtube api license",
    "raw youtube api license",
    "description license signals",
    "license notes",
    "license checked date",
    "needs manual review",
    "script version",
}

DEFAULT_INPUT_HEADERS = [
    "Link",
    "Speaker",
    "Upload date",
    "Engagement",
    "Time stamps",
    "Date accessed",
    "Length",
]

DEFAULT_TERM_DICTIONARY = {
    # Strong signals that the description/title is asserting a CC licence.
    # These are deliberately regexes so future teams can extend the list.
    "strong_cc_signals": [
        r"\bcreative commons\b",
        r"\bcc[-\s]?by\b",
        r"\bcc[-\s]?by[-\s]?sa\b",
        r"\bcc[-\s]?by[-\s]?nc\b",
        r"\bcc[-\s]?by[-\s]?nd\b",
        r"\bcc[-\s]?0\b",
        r"\bcc0\b",
        r"creativecommons\.org/licenses/",
        r"creativecommons\.org/publicdomain/",
        r"\blicensed under (?:a )?creative commons\b",
        r"\bunder (?:the )?creative commons licence\b",
        r"\bunder (?:the )?creative commons license\b",
    ],
    # Public-domain-like statements are useful but still need review if the API
    # says Standard, because the YouTube platform flag is separate evidence.
    "public_domain_signals": [
        r"\bpublic domain\b",
        r"\bno known copyright restrictions\b",
        r"\bfree cultural work\b",
    ],
    # Restrictive terms can conflict with a CC platform flag.
    "restrictive_signals": [
        r"\ball rights reserved\b",
        r"\bdo not re[-\s]?upload\b",
        r"\bdo not redistribute\b",
        r"\bunauthori[sz]ed use\b",
        r"\bno part of this\b.*\breproduced\b",
        r"\bcopyright(?:ed)? material\b",
        r"\bused with permission\b",
        r"\blicensed to youtube by\b",
    ],
    # Ambiguous words that are often misunderstood. They do not mean CC by themselves.
    "ambiguous_signals": [
        r"\broyalty[-\s]?free\b",
        r"\bcopyright[-\s]?free\b",
        r"\bfair use\b",
        r"\bfor educational purposes\b",
        r"\bcredit to\b",
        r"\bsource:\b",
    ],
}

URL_RE = re.compile(r"https?://[^\s<>\"\]\)]+", re.IGNORECASE)
YOUTUBE_ID_PATTERNS = [
    re.compile(r"(?:youtube\.com/watch\?(?:[^#\s]*&)?v=)([A-Za-z0-9_-]{11})", re.IGNORECASE),
    re.compile(r"(?:youtu\.be/)([A-Za-z0-9_-]{11})", re.IGNORECASE),
    re.compile(r"(?:youtube\.com/shorts/)([A-Za-z0-9_-]{11})", re.IGNORECASE),
    re.compile(r"(?:youtube\.com/embed/)([A-Za-z0-9_-]{11})", re.IGNORECASE),
    re.compile(r"(?:youtube\.com/live/)([A-Za-z0-9_-]{11})", re.IGNORECASE),
]
BARE_VIDEO_ID_RE = re.compile(r"\b[A-Za-z0-9_-]{11}\b")


@dataclass
class WorkItem:
    """A Word table row that should be checked."""

    table_number: int
    row_number: int
    row: Any
    urls: list[str]
    first_col_text: str
    col_final: int
    col_date: int


@dataclass
class Assessment:
    """Licence assessment for one Word row, usually one YouTube URL."""

    final_license: str
    api_license: str
    raw_api_license: str
    description_signals: str
    notes: str
    needs_manual_review: str


@dataclass
class DebugRow:
    """One row in the debug CSV."""

    table_number: int
    row_number: int
    word_first_column_text: str
    video_id: str
    url: str
    youtube_title: str
    youtube_channel: str
    youtube_published_at: str
    raw_api_license: str
    api_license_label: str
    final_license: str
    description_signals: str
    notes: str
    needs_manual_review: str
    checked_date: str
    script_version: str


def clean_text(text: str) -> str:
    """Collapse whitespace for readable output and comparisons."""
    return re.sub(r"\s+", " ", (text or "").strip())


def normalise_header(text: str) -> str:
    """Normalise a table header for forgiving comparisons."""
    return clean_text(text).lower()


def unique_preserving_order(values: Iterable[str]) -> list[str]:
    """De-duplicate a list without changing first-seen order."""
    seen: set[str] = set()
    output: list[str] = []
    for value in values:
        if value and value not in seen:
            seen.add(value)
            output.append(value)
    return output


def extract_hyperlink_targets(cell: Any) -> list[str]:
    """
    Extract hidden Word hyperlink targets from a cell.

    Word often stores the URL separately from visible text. A cell may display a
    video title but internally link to YouTube. This function recovers those URLs.
    """
    links: list[str] = []
    part = cell.part

    for hyperlink in cell._tc.xpath(".//w:hyperlink"):
        rid = hyperlink.get(qn("r:id"))
        if rid and rid in part.rels:
            links.append(part.rels[rid].target_ref)

    # Older Word field-code style: HYPERLINK "https://..."
    for instr in cell._tc.xpath(".//w:instrText"):
        text = instr.text or ""
        match = re.search(r'HYPERLINK\s+"([^"]+)"', text, flags=re.IGNORECASE)
        if match:
            links.append(match.group(1))

    return unique_preserving_order(links)


def extract_youtube_video_id(text: str) -> str:
    """Extract a YouTube video ID from a URL, hidden hyperlink target, or bare ID."""
    if not text:
        return ""

    text = text.strip()
    for pattern in YOUTUBE_ID_PATTERNS:
        match = pattern.search(text)
        if match:
            return match.group(1)

    # Robust URL parsing fallback for watch URLs with unusual query ordering.
    try:
        parsed = urllib.parse.urlparse(text)
        host = parsed.netloc.lower().replace("www.", "")
        if host in {"youtube.com", "m.youtube.com", "music.youtube.com"}:
            qs = urllib.parse.parse_qs(parsed.query)
            video_id = (qs.get("v") or [""])[0]
            if re.fullmatch(r"[A-Za-z0-9_-]{11}", video_id):
                return video_id
        if host == "youtu.be":
            candidate = parsed.path.strip("/").split("/")[0]
            if re.fullmatch(r"[A-Za-z0-9_-]{11}", candidate):
                return candidate
    except Exception:
        pass

    # Bare IDs are accepted only when the surrounding text is short. This avoids
    # treating random prose as a video ID.
    compact = clean_text(text)
    if len(compact) <= 80:
        match = BARE_VIDEO_ID_RE.search(compact)
        if match:
            return match.group(0)

    return ""


def extract_urls_from_cell(cell: Any) -> list[str]:
    """Extract candidate YouTube URLs/IDs from a Word cell."""
    candidates: list[str] = []

    visible_text = cell.text or ""
    candidates.extend(URL_RE.findall(visible_text))
    candidates.extend(extract_hyperlink_targets(cell))

    # In rare cases the visible cell may contain only the 11-character ID.
    bare_id = extract_youtube_video_id(visible_text)
    if bare_id:
        candidates.append(f"https://www.youtube.com/watch?v={bare_id}")

    urls: list[str] = []
    for candidate in candidates:
        video_id = extract_youtube_video_id(candidate)
        if video_id:
            urls.append(f"https://www.youtube.com/watch?v={video_id}")
    return unique_preserving_order(urls)


def table_has_youtube_rows(table: Any) -> bool:
    """Return True if any row has a YouTube link/ID in the first column."""
    if not table.rows:
        return False
    for row in table.rows:
        if row.cells and extract_urls_from_cell(row.cells[0]):
            return True
    return False


def table_has_header_row(table: Any, allowed_link_headers: set[str]) -> bool:
    """Decide whether the first row is already a header row."""
    if not table.rows or not table.rows[0].cells:
        return False
    first_cell = normalise_header(table.rows[0].cells[0].text)
    return first_cell in allowed_link_headers


def set_cell_text(cell: Any, text: str, max_chars: int = 650) -> None:
    """Set cell text, truncating very long notes so Word tables stay readable."""
    text = text or ""
    if max_chars and len(text) > max_chars:
        text = text[: max_chars - 20].rstrip() + " … [see debug CSV]"
    cell.text = text


def insert_header_row(table: Any) -> None:
    """
    Insert a header row at the top of a headerless table.

    python-docx has no public insert-row-at-top API. A tempting shortcut is to
    deep-copy the first row and replace its text, but that can duplicate hidden
    Word internals such as bookmarks, hyperlinks, and revision markers. Some
    versions of Microsoft Word then refuse to open the generated document. The
    safer route is to let python-docx create a clean row, move that clean row
    before the first real table row, and then write generic headers into it. It
    is important that the row is not inserted at raw XML index 0, because that
    would put it before the table properties/grid elements and Word will reject
    the file.
    """
    first_row_xml = table.rows[0]._tr
    new_row = table.add_row()
    first_row_xml.addprevious(new_row._tr)

    header_cells = table.rows[0].cells
    for i, cell in enumerate(header_cells):
        header = DEFAULT_INPUT_HEADERS[i] if i < len(DEFAULT_INPUT_HEADERS) else f"Existing column {i + 1}"
        cell.text = header


def ensure_column(table: Any, header_text: str, width_cm: float) -> int:
    """Return an existing output column index, or append a new output column."""
    wanted = normalise_header(header_text)
    headers = [normalise_header(cell.text) for cell in table.rows[0].cells]
    for index, header in enumerate(headers):
        if header == wanted:
            return index

    table.add_column(Cm(width_cm))
    new_index = len(table.rows[0].cells) - 1
    table.rows[0].cells[new_index].text = header_text
    return new_index

def remove_column(table: Any, index: int) -> None:
    """Remove a table column by index using the underlying Word XML.

    This is used only for legacy audit columns from earlier versions of this
    workflow. Normal input documents are not expected to need this.
    """
    # Remove the table-grid column, if Word stored one.
    tbl_grid = table._tbl.tblGrid
    if tbl_grid is not None and index < len(tbl_grid.gridCol_lst):
        tbl_grid.remove(tbl_grid.gridCol_lst[index])

    # Remove the cell at the same position from every row.
    for row in table.rows:
        if index < len(row.cells):
            tc = row.cells[index]._tc
            tc.getparent().remove(tc)


def remove_legacy_detail_columns(table: Any) -> int:
    """Remove old detailed audit columns from a table.

    The current team-facing DOCX output should contain only:
      - License
      - Date Checked

    Detailed fields still exist in the debug CSV and log DOCX.
    """
    if not table.rows or not table.rows[0].cells:
        return 0

    header_values = [normalise_header(cell.text) for cell in table.rows[0].cells]
    indexes_to_remove = [
        index
        for index, header in enumerate(header_values)
        if header in DETAILED_AUDIT_HEADERS_TO_REMOVE
    ]

    # Remove right-to-left so earlier indexes stay valid.
    for index in sorted(indexes_to_remove, reverse=True):
        remove_column(table, index)
    return len(indexes_to_remove)


def looks_like_summary_or_blank(text: str) -> bool:
    """Skip rows such as '27 total' or blank separator rows."""
    cleaned = normalise_header(text)
    if not cleaned:
        return True
    if re.fullmatch(r"\d+\s*(in\s*)?total", cleaned):
        return True
    if cleaned in {"total", "totals"}:
        return True
    return False


def collect_work_items(
    document: Any,
    process_all_tables: bool,
    insert_headers_for_headerless_tables: bool,
    link_header_names: list[str],
    max_cell_chars: int,
) -> tuple[list[WorkItem], dict[str, int]]:
    """Find all table rows that need YouTube checks and prepare output columns."""
    allowed_link_headers = {normalise_header(name) for name in link_header_names}
    work_items: list[WorkItem] = []
    stats = {
        "tables_seen": 0,
        "tables_with_youtube_rows": 0,
        "tables_processed": 0,
        "headerless_tables_fixed": 0,
        "detail_columns_removed": 0,
        "rows_seen_in_processed_tables": 0,
        "rows_with_youtube_urls": 0,
        "rows_without_urls_but_not_blank": 0,
    }

    for table_number, table in enumerate(document.tables, start=1):
        stats["tables_seen"] += 1
        if not table.rows or not table.rows[0].cells:
            continue

        has_youtube = table_has_youtube_rows(table)
        if has_youtube:
            stats["tables_with_youtube_rows"] += 1

        has_header = table_has_header_row(table, allowed_link_headers)
        should_process = has_youtube if process_all_tables else has_header
        if not should_process:
            continue

        if not has_header:
            if insert_headers_for_headerless_tables:
                insert_header_row(table)
                stats["headerless_tables_fixed"] += 1
                data_start_index = 1
            else:
                data_start_index = 0
        else:
            data_start_index = 1

        stats["detail_columns_removed"] += remove_legacy_detail_columns(table)

        col_indexes = {}
        for header, width in OUTPUT_COLUMNS:
            col_indexes[header] = ensure_column(table, header, width)
        stats["tables_processed"] += 1

        for row_number, row in enumerate(table.rows[data_start_index:], start=data_start_index + 1):
            if not row.cells:
                continue
            stats["rows_seen_in_processed_tables"] += 1
            first_col_text = clean_text(row.cells[0].text)
            urls = extract_urls_from_cell(row.cells[0])

            if urls:
                stats["rows_with_youtube_urls"] += 1
                work_items.append(
                    WorkItem(
                        table_number=table_number,
                        row_number=row_number,
                        row=row,
                        urls=urls,
                        first_col_text=first_col_text,
                        col_final=col_indexes[FINAL_LICENSE_HEADER],
                        col_date=col_indexes[CHECKED_DATE_HEADER],
                    )
                )
            elif not looks_like_summary_or_blank(first_col_text):
                # Leave non-URL rows alone; the summary flags them instead of
                # filling tables with false alarms.
                stats["rows_without_urls_but_not_blank"] += 1

    return work_items, stats


def api_get_json(url: str, params: dict[str, Any], timeout: int) -> dict[str, Any]:
    """GET JSON from the YouTube Data API with helpful errors."""
    query = urllib.parse.urlencode(params)
    request = urllib.request.Request(
        f"{url}?{query}",
        headers={"User-Agent": f"youtube-license-audit/{SCRIPT_VERSION}"},
    )
    try:
        with urllib.request.urlopen(request, timeout=timeout) as response:
            data = json.loads(response.read().decode("utf-8"))
    except Exception as exc:
        raise RuntimeError(f"YouTube API request failed: {exc}") from exc

    if "error" in data:
        raise RuntimeError(json.dumps(data["error"], indent=2))
    return data


def batched(values: list[str], size: int) -> Iterable[list[str]]:
    """Yield a list in fixed-size batches."""
    for i in range(0, len(values), size):
        yield values[i : i + size]


def fetch_youtube_metadata(api_key: str, video_ids: list[str], timeout: int, sleep_seconds: float) -> dict[str, dict[str, Any]]:
    """Fetch YouTube snippet/status metadata for up to 50 IDs per API request."""
    metadata: dict[str, dict[str, Any]] = {}
    if not video_ids:
        return metadata

    for batch in batched(video_ids, 50):
        data = api_get_json(
            YOUTUBE_VIDEOS_URL,
            {
                "part": "snippet,status",
                "id": ",".join(batch),
                "maxResults": 50,
                "key": api_key,
            },
            timeout=timeout,
        )
        for item in data.get("items", []):
            video_id = item.get("id", "")
            if video_id:
                metadata[video_id] = item
        if sleep_seconds > 0:
            time.sleep(sleep_seconds)
    return metadata


def load_term_dictionary(path: Path | None) -> dict[str, list[str]]:
    """Load a JSON dictionary of licence terms, or use the built-in defaults."""
    if not path:
        return DEFAULT_TERM_DICTIONARY
    return read_term_dictionary(path)


def compile_terms(term_dictionary: dict[str, list[str]]) -> dict[str, list[re.Pattern[str]]]:
    """Compile regex dictionary once for speed and validation."""
    compiled: dict[str, list[re.Pattern[str]]] = {}
    for category, patterns in term_dictionary.items():
        compiled[category] = [re.compile(pattern, flags=re.IGNORECASE) for pattern in patterns]
    return compiled


def find_description_signals(text: str, compiled_terms: dict[str, list[re.Pattern[str]]]) -> dict[str, list[str]]:
    """Find licence-related text patterns in title/description/tags."""
    found: dict[str, list[str]] = {}
    for category, patterns in compiled_terms.items():
        for pattern in patterns:
            if pattern.search(text or ""):
                found.setdefault(category, []).append(pattern.pattern)
    return found


def format_signals(signals: dict[str, list[str]]) -> str:
    """Make matched term signals readable in Word/CSV."""
    if not signals:
        return "No description licence signals found"
    parts: list[str] = []
    for category, patterns in signals.items():
        parts.append(f"{category}: " + ", ".join(patterns))
    return " | ".join(parts)


def raw_license_to_label(raw: str) -> str:
    """Convert YouTube raw licence value to a human-readable label."""
    if raw == "creativeCommon":
        return "Creative Commons Attribution (CC BY)"
    if raw == "youtube":
        return "Standard YouTube License"
    if not raw:
        return "UNKNOWN / NOT RETURNED"
    return f"UNKNOWN RAW LICENSE: {raw}"


def assess_video(
    video_id: str,
    url: str,
    metadata: dict[str, dict[str, Any]],
    compiled_terms: dict[str, list[re.Pattern[str]]],
    scan_title: bool,
    scan_tags: bool,
) -> tuple[Assessment, dict[str, str]]:
    """Assess one YouTube video and return Word-cell values plus debug metadata."""
    item = metadata.get(video_id)
    if not item:
        assessment = Assessment(
            final_license="UNKNOWN / NOT RETURNED",
            api_license="UNKNOWN / NOT RETURNED",
            raw_api_license="",
            description_signals="No YouTube API record returned",
            notes="YouTube API did not return this video ID. It may be private, deleted, invalid, unavailable, or blocked.",
            needs_manual_review="YES",
        )
        debug = {
            "youtube_title": "",
            "youtube_channel": "",
            "youtube_published_at": "",
        }
        return assessment, debug

    snippet = item.get("snippet", {})
    status = item.get("status", {})
    raw_license = status.get("license", "")
    api_label = raw_license_to_label(raw_license)

    title = clean_text(snippet.get("title", ""))
    channel = clean_text(snippet.get("channelTitle", ""))
    published_at = clean_text(snippet.get("publishedAt", ""))
    description = snippet.get("description", "") or ""
    tags = "\n".join(snippet.get("tags", []) or []) if scan_tags else ""

    scan_text_parts = [description]
    if scan_title:
        scan_text_parts.insert(0, title)
    if tags:
        scan_text_parts.append(tags)
    signals = find_description_signals("\n\n".join(scan_text_parts), compiled_terms)
    signals_text = format_signals(signals)

    has_cc_claim = bool(signals.get("strong_cc_signals") or signals.get("public_domain_signals"))
    has_restrictive = bool(signals.get("restrictive_signals"))
    has_ambiguous = bool(signals.get("ambiguous_signals"))

    notes: list[str] = [f"YouTube API status.license says {raw_license or 'missing'}."]
    manual_reasons: list[str] = []

    final_license = api_label
    if raw_license == "youtube" and has_cc_claim:
        final_license = "Standard YouTube License - description claims CC/public-domain; manual review"
        manual_reasons.append("Description/title contains CC or public-domain wording but YouTube API says Standard.")
    if raw_license == "creativeCommon" and has_restrictive:
        final_license = "Creative Commons Attribution (CC BY) - conflicting restrictive wording; manual review"
        manual_reasons.append("YouTube API says Creative Commons but description/title contains restrictive wording.")
    if has_ambiguous:
        manual_reasons.append("Description/title contains ambiguous reuse wording such as fair use, royalty-free, or copyright-free.")
    if not raw_license:
        manual_reasons.append("YouTube API did not return a licence value.")

    if manual_reasons:
        notes.append("Manual review recommended: " + "; ".join(manual_reasons))

    assessment = Assessment(
        final_license=final_license,
        api_license=api_label,
        raw_api_license=raw_license,
        description_signals=signals_text,
        notes=" ".join(notes),
        needs_manual_review="YES" if manual_reasons else "NO",
    )
    debug = {
        "youtube_title": title,
        "youtube_channel": channel,
        "youtube_published_at": published_at,
    }
    return assessment, debug


def combine_assessments(assessments: list[Assessment]) -> Assessment:
    """Combine results when a single Word row contains multiple YouTube URLs."""
    if not assessments:
        return Assessment(
            final_license="NO YOUTUBE URL FOUND",
            api_license="NO YOUTUBE URL FOUND",
            raw_api_license="",
            description_signals="No YouTube URL found in first-column cell",
            notes="No YouTube URL was found in the first-column cell.",
            needs_manual_review="YES",
        )
    return Assessment(
        final_license="; ".join(unique_preserving_order(a.final_license for a in assessments)),
        api_license="; ".join(unique_preserving_order(a.api_license for a in assessments)),
        raw_api_license="; ".join(unique_preserving_order(a.raw_api_license for a in assessments)),
        description_signals="; ".join(unique_preserving_order(a.description_signals for a in assessments)),
        notes="; ".join(unique_preserving_order(a.notes for a in assessments)),
        needs_manual_review="YES" if any(a.needs_manual_review == "YES" for a in assessments) else "NO",
    )


def write_debug_csv(debug_rows: list[DebugRow], path: Path) -> None:
    """Write detailed raw output for auditability."""
    path.parent.mkdir(parents=True, exist_ok=True)
    with path.open("w", newline="", encoding="utf-8-sig") as f:
        fieldnames = list(asdict(debug_rows[0]).keys()) if debug_rows else list(DebugRow.__dataclass_fields__.keys())
        writer = SpreadsheetSafeWriter(csv.DictWriter(f, fieldnames=fieldnames))
        writer.writeheader()
        for row in debug_rows:
            writer.writerow(asdict(row))


def write_summary_txt(stats: dict[str, Any], path: Path) -> None:
    """Write a plain-English summary for non-technical users."""
    path.parent.mkdir(parents=True, exist_ok=True)
    lines = [
        "YouTube Licence Audit Summary",
        "=" * 30,
        f"Generated: {dt.datetime.now().strftime('%Y-%m-%d %H:%M:%S')}",
        f"Script version: {SCRIPT_VERSION}",
        "",
        "Main results",
        "------------",
        f"Input file: {stats.get('input_path')}",
        f"Output file: {stats.get('output_path')}",
        f"Debug CSV: {stats.get('debug_csv')}",
        "",
        "Counts",
        "------",
    ]
    ordered_keys = [
        "tables_seen",
        "tables_with_youtube_rows",
        "tables_processed",
        "headerless_tables_fixed",
        "detail_columns_removed",
        "rows_seen_in_processed_tables",
        "rows_with_youtube_urls",
        "rows_without_urls_but_not_blank",
        "rows_processed",
        "unique_youtube_ids",
        "api_creative_common_count",
        "api_standard_count",
        "api_unknown_or_missing_count",
        "description_cc_claim_count",
        "manual_review_count",
    ]
    for key in ordered_keys:
        if key in stats:
            lines.append(f"{key}: {stats[key]}")
    lines.extend([
        "",
        "Interpretation",
        "--------------",
        "The YouTube API licence and description-text scan are kept separate.",
        "If YouTube says Standard but the description mentions Creative Commons, the row is marked for manual review rather than silently treated as CC.",
        "If YouTube says Creative Commons but the description contains restrictive wording, the row is also marked for manual review.",
    ])
    path.write_text("\n".join(lines), encoding="utf-8")


def write_log_docx(stats: dict[str, Any], debug_rows: list[DebugRow], path: Path) -> None:
    """Write a human-readable DOCX log explaining the hidden audit details.

    The main Word source document only gets two new columns: License and Date
    Checked. This log document is where future teams can see what those values
    mean and inspect the detailed evidence behind them.
    """
    path.parent.mkdir(parents=True, exist_ok=True)
    doc = Document()

    doc.add_heading("YouTube Licence Audit Log", level=1)
    doc.add_paragraph(
        "This log explains the licence columns written into the main DOCX file "
        "and records the detailed evidence that is intentionally kept out of the "
        "main source table."
    )

    doc.add_heading("Columns appended to the main document", level=2)
    table = doc.add_table(rows=1, cols=3)
    table.style = "Table Grid"
    headers = ["Column", "Meaning", "How to use it"]
    for i, header in enumerate(headers):
        table.rows[0].cells[i].text = header

    for values in [
        (
            "License",
            "The final licence label for the video row. It is based on YouTube's official status.license field, with manual-review wording added if text signals conflict.",
            "Use this as the simple value in the research source table. Rows containing 'manual review' should be checked by a person before relying on them.",
        ),
        (
            "Date Checked",
            "The date on which this script checked the YouTube API and scanned the returned title/description metadata.",
            "Use this to show when the licence record was last verified, because YouTube metadata can change.",
        ),
    ]:
        cells = table.add_row().cells
        for i, value in enumerate(values):
            cells[i].text = value

    doc.add_heading("Detailed fields stored in the CSV/log, not in the main document", level=2)
    detail_table = doc.add_table(rows=1, cols=2)
    detail_table.style = "Table Grid"
    detail_table.rows[0].cells[0].text = "Detailed field"
    detail_table.rows[0].cells[1].text = "Meaning"
    details = [
        ("YouTube API License", "Human-readable version of YouTube status.license, such as Standard YouTube License or Creative Commons Attribution (CC BY)."),
        ("Raw YouTube API License", "The raw API value returned by YouTube: usually youtube or creativeCommon."),
        ("Description License Signals", "Regex/text matches found in the YouTube title, description, and optionally tags. These are evidence flags only; they do not override the API value."),
        ("License Notes", "Plain-English explanation of the decision and any conflict requiring manual review."),
        ("Needs manual review", "YES when the API result and description signals conflict, when the API result is missing, or when wording is ambiguous."),
        ("Script version", "The version of the script used to create the audit output."),
    ]
    for name, meaning in details:
        cells = detail_table.add_row().cells
        cells[0].text = name
        cells[1].text = meaning

    doc.add_heading("Run summary", level=2)
    summary_table = doc.add_table(rows=1, cols=2)
    summary_table.style = "Table Grid"
    summary_table.rows[0].cells[0].text = "Item"
    summary_table.rows[0].cells[1].text = "Value"
    summary_keys = [
        "input_path", "output_path", "debug_csv", "rows_processed",
        "unique_youtube_ids", "api_creative_common_count", "api_standard_count",
        "api_unknown_or_missing_count", "description_cc_claim_count",
        "manual_review_count", "tables_seen", "tables_with_youtube_rows",
        "tables_processed", "headerless_tables_fixed", "detail_columns_removed",
    ]
    for key in summary_keys:
        if key in stats:
            cells = summary_table.add_row().cells
            cells[0].text = key
            cells[1].text = str(stats.get(key, ""))

    doc.add_heading("How the licence decision is made", level=2)
    for paragraph in [
        "YouTube's platform licence is read from status.license. The raw value youtube means Standard YouTube License. The raw value creativeCommon means Creative Commons Attribution (CC BY).",
        "The script also scans the YouTube title and description for terms such as Creative Commons, CC BY, CC0, public domain, all rights reserved, do not re-upload, royalty-free, copyright-free, and fair use.",
        "Description wording is not treated as stronger than the YouTube platform flag. It is recorded as a signal and can trigger manual review when it conflicts with the API result.",
        "The original input document is not overwritten. The output document is a copy of the original with only License and Date Checked appended to relevant video tables.",
    ]:
        doc.add_paragraph(paragraph)

    manual_rows = [row for row in debug_rows if row.needs_manual_review == "YES"]
    doc.add_heading("Rows needing manual review", level=2)
    if not manual_rows:
        doc.add_paragraph("No manual-review rows were detected in this run.")
    else:
        doc.add_paragraph(
            "The table below lists rows where a human should inspect the YouTube page/API data before relying on the licence value."
        )
        review_table = doc.add_table(rows=1, cols=6)
        review_table.style = "Table Grid"
        for i, header in enumerate(["Table", "Row", "Video ID", "Final licence", "Signals", "Notes"]):
            review_table.rows[0].cells[i].text = header
        for row in manual_rows[:200]:
            cells = review_table.add_row().cells
            cells[0].text = str(row.table_number)
            cells[1].text = str(row.row_number)
            cells[2].text = row.video_id
            cells[3].text = row.final_license
            cells[4].text = row.description_signals[:700]
            cells[5].text = row.notes[:900]
        if len(manual_rows) > 200:
            doc.add_paragraph(f"Only the first 200 manual-review rows are shown here. See the debug CSV for all {len(manual_rows)} rows.")

    doc.add_heading("Output files", level=2)
    doc.add_paragraph(f"Main checked DOCX file: {stats.get('output_path', '')}")
    doc.add_paragraph(f"Detailed debug CSV: {stats.get('debug_csv', '')}")
    doc.add_paragraph(f"Generated: {dt.datetime.now().strftime('%Y-%m-%d %H:%M:%S')}")
    doc.add_paragraph(f"Script version: {SCRIPT_VERSION}")

    doc.save(str(path))


def process_document(args: argparse.Namespace) -> int:
    """Main document-processing routine."""
    input_path: Path = args.docx
    output_path: Path = args.output
    checked_date = args.checked_date or dt.datetime.now().astimezone().date().isoformat()

    if not input_path.exists():
        print(f"ERROR: file not found: {input_path}", file=sys.stderr)
        return 2
    if input_path.suffix.lower() != ".docx":
        print("ERROR: input file must be a .docx file", file=sys.stderr)
        return 2
    if not args.api_key:
        print("ERROR: missing YouTube API key. Set YOUTUBE_API_KEY or pass --api-key.", file=sys.stderr)
        return 2

    try:
        term_dictionary = load_term_dictionary(args.terms_json)
        compiled_terms = compile_terms(term_dictionary)
    except Exception as exc:
        print(f"ERROR: could not load/compile licence terms: {exc}", file=sys.stderr)
        return 2

    print(f"Opening: {input_path}")
    document = open_docx_document(input_path)
    work_items, stats = collect_work_items(
        document=document,
        process_all_tables=not args.only_headered_tables,
        insert_headers_for_headerless_tables=not args.no_insert_headerless_headers,
        link_header_names=args.link_header_name,
        max_cell_chars=args.max_cell_chars,
    )

    all_video_ids: list[str] = []
    for item in work_items:
        for url in item.urls:
            video_id = extract_youtube_video_id(url)
            if video_id:
                all_video_ids.append(video_id)
    unique_video_ids = unique_preserving_order(all_video_ids)

    print(f"Found {len(work_items)} video rows and {len(unique_video_ids)} unique YouTube IDs.")
    print("Checking YouTube API...")
    try:
        metadata = fetch_youtube_metadata(args.api_key, unique_video_ids, args.timeout, args.sleep_seconds)
    except Exception as exc:
        print(f"ERROR: YouTube API check failed: {exc}", file=sys.stderr)
        return 2

    debug_rows: list[DebugRow] = []
    api_cc = api_standard = api_unknown = desc_cc_count = manual_count = 0

    for item in work_items:
        row_assessments: list[Assessment] = []
        for url in item.urls:
            video_id = extract_youtube_video_id(url)
            if not video_id:
                continue
            assessment, debug = assess_video(
                video_id=video_id,
                url=url,
                metadata=metadata,
                compiled_terms=compiled_terms,
                scan_title=not args.no_scan_title,
                scan_tags=args.scan_tags,
            )
            row_assessments.append(assessment)

            if assessment.raw_api_license == "creativeCommon":
                api_cc += 1
            elif assessment.raw_api_license == "youtube":
                api_standard += 1
            else:
                api_unknown += 1
            if "strong_cc_signals" in assessment.description_signals or "public_domain_signals" in assessment.description_signals:
                desc_cc_count += 1
            if assessment.needs_manual_review == "YES":
                manual_count += 1

            debug_rows.append(
                DebugRow(
                    table_number=item.table_number,
                    row_number=item.row_number,
                    word_first_column_text=item.first_col_text,
                    video_id=video_id,
                    url=url,
                    youtube_title=debug.get("youtube_title", ""),
                    youtube_channel=debug.get("youtube_channel", ""),
                    youtube_published_at=debug.get("youtube_published_at", ""),
                    raw_api_license=assessment.raw_api_license,
                    api_license_label=assessment.api_license,
                    final_license=assessment.final_license,
                    description_signals=assessment.description_signals,
                    notes=assessment.notes,
                    needs_manual_review=assessment.needs_manual_review,
                    checked_date=checked_date,
                    script_version=SCRIPT_VERSION,
                )
            )

        combined = combine_assessments(row_assessments)
        cells = item.row.cells
        set_cell_text(cells[item.col_final], combined.final_license, args.max_cell_chars)
        set_cell_text(cells[item.col_date], checked_date, args.max_cell_chars)

    output_path.parent.mkdir(parents=True, exist_ok=True)
    document.save(str(output_path))

    if args.debug_csv:
        write_debug_csv(debug_rows, args.debug_csv)

    stats.update(
        {
            "input_path": str(input_path),
            "output_path": str(output_path),
            "debug_csv": str(args.debug_csv or ""),
            "rows_processed": len(work_items),
            "unique_youtube_ids": len(unique_video_ids),
            "api_creative_common_count": api_cc,
            "api_standard_count": api_standard,
            "api_unknown_or_missing_count": api_unknown,
            "description_cc_claim_count": desc_cc_count,
            "manual_review_count": manual_count,
        }
    )
    if args.summary_txt:
        write_summary_txt(stats, args.summary_txt)
    if args.log_docx:
        write_log_docx(stats, debug_rows, args.log_docx)

    print("\nDone.")
    print(f"Rows processed:              {len(work_items)}")
    print(f"Unique YouTube IDs:          {len(unique_video_ids)}")
    print(f"API Creative Commons videos: {api_cc}")
    print(f"API Standard videos:         {api_standard}")
    print(f"Manual-review rows:          {manual_count}")
    print(f"Saved DOCX file:             {output_path}")
    if args.debug_csv:
        print(f"Saved debug CSV:             {args.debug_csv}")
    if args.summary_txt:
        print(f"Saved summary:               {args.summary_txt}")
    if args.log_docx:
        print(f"Saved log document:          {args.log_docx}")

    return 0


def parse_args() -> argparse.Namespace:
    parser = argparse.ArgumentParser(
        description="Audit YouTube licences in a DOCX file and add only License + Date Checked columns."
    )
    parser.add_argument("docx", type=Path, help="Input .docx file")
    parser.add_argument("--output", "-o", type=Path, required=True, help="Output .docx file")
    parser.add_argument("--api-key", default=os.getenv("YOUTUBE_API_KEY"), help="YouTube Data API key")
    parser.add_argument("--checked-date", default=dt.datetime.now().astimezone().date().isoformat())
    parser.add_argument("--debug-csv", type=Path, help="Optional detailed CSV output")
    parser.add_argument("--summary-txt", type=Path, help="Optional plain-English summary TXT")
    parser.add_argument("--log-docx", type=Path, help="Optional DOCX log explaining the licence columns and detailed audit evidence")
    parser.add_argument("--terms-json", type=Path, help="Optional licence term dictionary JSON")
    parser.add_argument(
        "--only-headered-tables",
        action="store_true",
        help="Only process tables whose first header cell is Link/URL. Default processes any table with YouTube links in the first column.",
    )
    parser.add_argument(
        "--no-insert-headerless-headers",
        action="store_true",
        help="Do not insert a header row in headerless tables. Not recommended for normal team use.",
    )
    parser.add_argument(
        "--link-header-name",
        action="append",
        default=["Link", "Links", "URL", "Video URL"],
        help="Allowed first-column header name. Can be repeated.",
    )
    parser.add_argument("--no-scan-title", action="store_true", help="Do not scan YouTube titles for licence terms")
    parser.add_argument("--scan-tags", action="store_true", help="Also scan YouTube tags where returned by the API")
    parser.add_argument("--timeout", type=int, default=60, help="YouTube API timeout in seconds")
    parser.add_argument("--sleep-seconds", type=float, default=0.0, help="Delay between API batches, seconds")
    parser.add_argument("--max-cell-chars", type=int, default=650, help="Maximum text length written into Word cells")
    return parser.parse_args()


def main() -> int:
    args = parse_args()
    return process_document(args)


if __name__ == "__main__":
    raise SystemExit(main())
