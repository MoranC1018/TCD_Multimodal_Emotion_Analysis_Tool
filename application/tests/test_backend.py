from __future__ import annotations

import csv
import hashlib
import json
import os
import subprocess
import sys
import tempfile
import unittest
from datetime import datetime, timezone
from pathlib import Path
from unittest.mock import patch

from docx import Document

from analysis.combined_summary import AUDIO_METRICS, AUDIO_REQUIRED_METRICS, VIDEO_METRICS
from analysis.profile import AnalysisProfile, ManualGroup, ProfileMember
from application import backend


class ProcurementUiBackendTests(unittest.TestCase):
    def test_nearby_metadata_rejects_oversized_control_json(self) -> None:
        with tempfile.TemporaryDirectory() as temp_dir:
            video = Path(temp_dir) / "clip.mp4"
            video.write_bytes(b"video")
            video.with_suffix(".json").write_text(
                '{"padding":"' + ("x" * (1024 * 1024)) + '"}',
                encoding="utf-8",
            )

            with self.assertRaisesRegex(ValueError, "metadata JSON exceeds"):
                backend.read_nearby_metadata(video)

    def test_catalog_scan_preserves_source_order_metadata_pooled_label_and_youtube_language(self) -> None:
        with tempfile.TemporaryDirectory() as temp_dir:
            root = Path(temp_dir)
            local_video = root / "local.mp4"
            local_video.write_bytes(b"synthetic video")
            source = root / "sources.csv"
            with source.open("w", encoding="utf-8-sig", newline="") as handle:
                writer = csv.writer(handle)
                writer.writerow(["Link", "Speaker", "Country", "Language", "Upload Date"])
                writer.writerows(
                    [
                        ["local.mp4", "", "Ireland", "Irish", "ignored"],
                        ["https://youtu.be/abcdefghijk", "Speaker A", "Canada", "French", "ignored"],
                        ["https://youtu.be/abcdefghijk", "Speaker A", "Japan", "", "ignored"],
                    ]
                )

            with (
                patch.object(backend, "load_youtube_api_key", return_value="key"),
                patch.object(
                    backend,
                    "fetch_youtube_api_metadata",
                    return_value={
                        "abcdefghijk": {
                            "title": "API title",
                            "duration_seconds": 42,
                            "youtube_language": "fr-CA",
                        }
                    },
                ),
            ):
                result = backend.scan_input_source(source, duration_reader=lambda _path: 12.0)

        self.assertEqual(result.source_kind, "catalog")
        self.assertEqual(result.catalog_format, "csv")
        self.assertRegex(result.catalog_sha256, r"^[0-9a-f]{64}$")
        self.assertEqual(result.metadata_headers, ["Country", "Language"])
        self.assertEqual([item.id for item in result.sources], ["source-0001", "source-0002", "source-0003"])
        self.assertEqual([item.source_kind for item in result.sources], ["file", "youtube", "youtube"])
        self.assertEqual(result.sources[0].speaker, "Pooled (no speaker)")
        self.assertEqual(result.sources[0].metadata, {"Country": "Ireland", "Language": "Irish"})
        self.assertEqual(result.sources[1].metadata["Language"], "French")
        self.assertEqual(result.sources[1].youtube_language, "fr-CA")
        self.assertEqual(result.sources[1].source_id, "source-0002")
        self.assertEqual(result.sources[2].source_id, "source-0003")

    def test_youtube_language_uses_audio_then_default_then_blank(self) -> None:
        item = backend.VideoItem(
            id="source-0001",
            title="https://youtu.be/abcdefghijk",
            speaker="Pooled (no speaker)",
            source_path="https://youtu.be/abcdefghijk",
            source_kind="youtube",
            youtube_url="https://www.youtube.com/watch?v=abcdefghijk",
            video_id="abcdefghijk",
        )
        cases = [
            ({"defaultAudioLanguage": "ga", "defaultLanguage": "en"}, "ga"),
            ({"defaultAudioLanguage": "", "defaultLanguage": "en"}, "en"),
            ({}, ""),
        ]
        for metadata, expected in cases:
            with self.subTest(metadata=metadata):
                updated = backend.apply_youtube_metadata([item], {"abcdefghijk": metadata})
                self.assertEqual(updated[0].youtube_language, expected)

    def test_catalog_command_for_every_mode_binds_digest_and_selected_source_ids(self) -> None:
        for mode in ("standard", "full", "manual", "clean-speaker-beta"):
            request = backend.RunRequest(
                mode=mode,
                source_path=Path(r"C:\input\sources.csv"),
                output_root=Path(r"C:\output\run"),
                selected_ids=["source-0003", "source-0001"],
                catalog_sha256="a" * 64,
                segment_manifest=Path(r"C:\segments\selection.json") if mode == "manual" else None,
                segment_manifest_sha256="b" * 64 if mode == "manual" else "",
                segment_expected_source=r"C:\input\sources.csv" if mode == "manual" else "",
            )

            command = backend.build_run_command(
                request,
                repo_root=Path(r"C:\repo"),
                python_executable=Path(r"C:\Python312\python.exe"),
            )

            with self.subTest(mode=mode):
                self.assertEqual(command[:3], [r"C:\Python312\python.exe", "-m", "procurement.catalog_runner"])
                self.assertEqual(command[command.index("--catalog-sha256") + 1], "a" * 64)
                self.assertEqual(command[command.index("--mode") + 1], mode)
                source_ids = [command[index + 1] for index, value in enumerate(command) if value == "--source-id"]
                self.assertEqual(source_ids, ["source-0003", "source-0001"])
                run_root = Path(command[command.index("--run-root") + 1])
                self.assertEqual(run_root.parent, Path(r"C:\output\run"))
                self.assertNotEqual(run_root, Path(r"C:\output\run"))

    def test_long_catalog_stem_preserves_unique_suffix_for_every_fresh_run(self) -> None:
        source = Path(r"C:\input") / (("very-long-catalog-name-" * 12) + ".csv")
        request = backend.RunRequest(
            mode="standard",
            source_path=source,
            output_root=Path(r"C:\output"),
            selected_ids=["source-0001"],
            catalog_sha256="a" * 64,
        )

        first = backend.build_run_command(request, repo_root=Path(r"C:\repo"))
        second = backend.build_run_command(request, repo_root=Path(r"C:\repo"))
        first_root = Path(first[first.index("--run-root") + 1])
        second_root = Path(second[second.index("--run-root") + 1])

        self.assertEqual(first_root.parent, Path(r"C:\output"))
        self.assertLessEqual(len(first_root.name), 120)
        self.assertNotEqual(first_root, second_root)
        self.assertIn("_catalog_standard_", first_root.name)

    def test_local_docx_cannot_downgrade_to_legacy_run_without_catalog_authorization(self) -> None:
        with self.assertRaisesRegex(ValueError, "launcher-validated catalog SHA"):
            backend.build_run_command(
                backend.RunRequest(
                    mode="standard",
                    source_path=Path(r"C:\input\sources.docx"),
                    output_root=Path(r"C:\output"),
                    selected_speakers=["Speaker A"],
                ),
                repo_root=Path(r"C:\repo"),
            )

    def test_catalog_command_preserves_mode_specific_options(self) -> None:
        request = backend.RunRequest(
            mode="clean-speaker-beta",
            source_path=Path(r"C:\input\sources.csv"),
            output_root=Path(r"C:\output"),
            selected_ids=["source-0002"],
            catalog_sha256="a" * 64,
            percentage=0.25,
            max_segment_seconds=45,
            beta_output_mode="percentage",
            beta_min_clean_seconds=12.0,
            beta_gap_seconds=2.5,
            beta_identity_stills=18,
            beta_scan_fps=1.5,
            beta_validation_fps=2.5,
            beta_face_confidence=0.72,
            beta_speaker_confidence=0.68,
            beta_worker_count=3,
            beta_device="cpu",
            beta_keep_debug=True,
            beta_resource_guard_percent=12.0,
            beta_resource_poll_seconds=3.0,
            beta_resource_guard_timeout_seconds=120.0,
            beta_parallel_detector_streams=True,
            beta_reference_audio=Path(r"C:\profiles\speaker.wav"),
            beta_only_video_ids=["abcdefghijk"],
            beta_random_one=True,
            beta_random_seed="repeatable",
            beta_isolated_video_processes=True,
            beta_skip_first_videos=12,
            beta_skip_completed_outputs=True,
            beta_video_cooldown_seconds=45.0,
            beta_max_download_height=480,
            beta_max_affinity_cores=2,
            beta_native_threads=1,
            beta_cpu_throttle_high_percent=96.0,
            beta_cpu_throttle_low_percent=91.0,
            beta_ram_throttle_high_percent=97.0,
            beta_ram_throttle_low_percent=92.0,
        )

        command = backend.build_run_command(
            request,
            repo_root=Path(r"C:\repo"),
            python_executable=Path(r"C:\Python312\python.exe"),
        )

        expected_pairs = {
            "--percentage": "0.25",
            "--max-segment-seconds": "45",
            "--output-mode": "percentage",
            "--min-clean-seconds": "12.0",
            "--gap-seconds": "2.5",
            "--identity-stills": "18",
            "--scan-fps": "1.5",
            "--validation-fps": "2.5",
            "--face-confidence": "0.72",
            "--speaker-confidence": "0.68",
            "--workers": "3",
            "--device": "cpu",
            "--resource-guard-percent": "12.0",
            "--resource-poll-seconds": "3.0",
            "--resource-guard-timeout-seconds": "120.0",
            "--reference-audio": str(Path(r"C:\profiles\speaker.wav").resolve()),
            "--only-video-id": "abcdefghijk",
            "--random-seed": "repeatable",
            "--skip-first-videos": "12",
            "--video-cooldown-seconds": "45.0",
            "--max-download-height": "480",
            "--max-affinity-cores": "2",
            "--native-threads": "1",
            "--cpu-throttle-high-percent": "96.0",
            "--cpu-throttle-low-percent": "91.0",
            "--ram-throttle-high-percent": "97.0",
            "--ram-throttle-low-percent": "92.0",
        }
        for option, expected in expected_pairs.items():
            with self.subTest(option=option):
                self.assertEqual(command[command.index(option) + 1], expected)
        for flag in (
            "--keep-debug",
            "--parallel-detectors",
            "--random-one",
            "--isolated-video-processes",
            "--skip-completed-outputs",
        ):
            self.assertIn(flag, command)

    def test_manual_catalog_command_preserves_validated_focus_manifest_binding(self) -> None:
        request = backend.RunRequest(
            mode="manual",
            source_path=Path(r"C:\input\sources.docx"),
            output_root=Path(r"C:\output"),
            selected_ids=["source-0001"],
            catalog_sha256="a" * 64,
            segment_manifest=Path(r"C:\segments\selection.json"),
            segment_manifest_sha256="b" * 64,
            segment_expected_source=r"C:\input\sources.docx",
        )

        command = backend.build_run_command(
            request,
            repo_root=Path(r"C:\repo"),
            python_executable=Path(r"C:\Python312\python.exe"),
        )

        self.assertEqual(command[command.index("--manifest-sha256") + 1], "b" * 64)
        self.assertEqual(command[command.index("--expected-source") + 1], r"C:\input\sources.docx")
        self.assertEqual(
            command[command.index("--segments-json") + 1],
            str(Path(r"C:\segments\selection.json").resolve()),
        )

    def test_analysis_speaker_discovery_reads_imported_reports_without_writing(self) -> None:
        with tempfile.TemporaryDirectory() as temp_dir:
            root = Path(temp_dir)
            imported = root / "imported"
            self._write_sectioned_report(
                imported / "emotion" / "Speaker A" / "combined" / "other_findings" / "descriptive_statistics.csv",
                AUDIO_METRICS,
            )
            before = self._tree_hash(root)

            result = backend.discover_analysis_speakers(
                (backend.AnalysisModalityRunRequest("audio", "import", imported),)
            )

            self.assertEqual(
                result,
                {
                    "speakers": [
                        {"key": "speakera", "name": "Speaker A", "availableIn": ["audio"]},
                    ],
                    "warnings": [],
                },
            )
            self.assertEqual(self._tree_hash(root), before)

    def test_analysis_speaker_discovery_reads_imported_text_results_without_writing(self) -> None:
        with tempfile.TemporaryDirectory() as temp_dir:
            root = Path(temp_dir)
            imported = root / "text-results"
            self._write_text_summary(
                imported / "text_output" / "multimodal" / "speaker_level_summary.csv"
            )
            before = self._tree_hash(root)

            result = backend.discover_analysis_speakers(
                (backend.AnalysisModalityRunRequest("text", "import", imported),)
            )

            self.assertEqual(
                result,
                {
                    "speakers": [
                        {"key": "speakerc", "name": "Speaker C", "availableIn": ["text"]},
                    ],
                    "warnings": [],
                },
            )
            self.assertEqual(self._tree_hash(root), before)

    def test_analysis_speaker_discovery_keeps_arbitrary_imported_speakers(self) -> None:
        with tempfile.TemporaryDirectory() as temp_dir:
            imported = Path(temp_dir) / "imported"
            self._write_sectioned_report(
                imported / "emotion" / "Speaker A" / "combined" / "other_findings" / "descriptive_statistics.csv",
                AUDIO_METRICS,
            )
            self._write_sectioned_report(
                imported / "emotion" / "Unknown Person" / "combined" / "other_findings" / "descriptive_statistics.csv",
                AUDIO_METRICS,
            )

            result = backend.discover_analysis_speakers(
                (backend.AnalysisModalityRunRequest("audio", "import", imported),)
            )

        self.assertEqual(
            result["speakers"],
            [
                {"key": "speakera", "name": "Speaker A", "availableIn": ["audio"]},
                {"key": "unknownperson", "name": "Unknown Person", "availableIn": ["audio"]},
            ],
        )
        self.assertEqual(result["warnings"], [])

    def test_analysis_speaker_discovery_accepts_legacy_audio_baseline_metrics(self) -> None:
        with tempfile.TemporaryDirectory() as temp_dir:
            imported = Path(temp_dir) / "imported"
            self._write_sectioned_report(
                imported / "emotion" / "Speaker A" / "combined" / "other_findings" / "descriptive_statistics.csv",
                AUDIO_REQUIRED_METRICS,
            )

            result = backend.discover_analysis_speakers(
                (backend.AnalysisModalityRunRequest("audio", "import", imported),)
            )

        self.assertEqual(result["speakers"][0]["key"], "speakera")
        self.assertEqual(result["warnings"], [])

    def test_analysis_speaker_discovery_unions_fresh_imotions_and_audio_in_canonical_order(self) -> None:
        with tempfile.TemporaryDirectory() as temp_dir:
            root = Path(temp_dir)
            imotions = root / "imotions"
            audio = root / "audio"
            self._write_imotions_csv(imotions / "Speaker B" / "sample.csv")
            self._write_audio_csv(audio / "Speaker A" / "speech" / "audio_analysis.csv", "Speaker A")
            self._write_audio_csv(audio / "Speaker B" / "speech" / "audio_analysis.csv", "Speaker B")
            before = self._tree_hash(root)

            result = backend.discover_analysis_speakers(
                (
                    backend.AnalysisModalityRunRequest("audio", "run", audio),
                    backend.AnalysisModalityRunRequest("imotions", "run", imotions),
                )
            )
            self.assertEqual(self._tree_hash(root), before)

        self.assertEqual(
            result,
            {
                "speakers": [
                    {"key": "speakera", "name": "Speaker A", "availableIn": ["audio"]},
                    {"key": "speakerb", "name": "Speaker B", "availableIn": ["imotions", "audio"]},
                ],
                "warnings": [],
            },
        )

    def test_analysis_speaker_discovery_accepts_arbitrary_fresh_labels(self) -> None:
        with tempfile.TemporaryDirectory() as temp_dir:
            source = Path(temp_dir) / "audio"
            self._write_audio_csv(source / "Unknown Person" / "speech" / "audio_analysis.csv", "Unknown Person")

            result = backend.discover_analysis_speakers(
                (backend.AnalysisModalityRunRequest("audio", "run", source),)
            )

        self.assertEqual(
            result["speakers"],
            [{"key": "unknownperson", "name": "Unknown Person", "availableIn": ["audio"]}],
        )
        self.assertEqual(result["warnings"], [])

    def test_analysis_speaker_discovery_rejects_empty_or_unusable_fresh_sources(self) -> None:
        with tempfile.TemporaryDirectory() as temp_dir:
            root = Path(temp_dir)
            empty_imotions = root / "empty-imotions"
            empty_audio = root / "empty-audio"
            header_only_imotions = root / "header-only-imotions"
            malformed_audio = root / "malformed-audio"
            empty_audio_report = root / "empty-audio-report"
            empty_imotions.mkdir()
            empty_audio.mkdir()
            self._write_header_only_imotions_csv(header_only_imotions / "Speaker B" / "header-only.csv")
            malformed = malformed_audio / "Speaker B" / "speech" / "audio_analysis.csv"
            malformed.parent.mkdir(parents=True)
            malformed.write_text("not,a,valid,audio,report\n", encoding="utf-8")
            empty_report = empty_audio_report / "Speaker B" / "speech" / "audio_analysis.csv"
            empty_report.parent.mkdir(parents=True)
            empty_report.touch()

            invalid_modalities = (
                backend.AnalysisModalityRunRequest("imotions", "run", empty_imotions),
                backend.AnalysisModalityRunRequest("audio", "run", empty_audio),
                backend.AnalysisModalityRunRequest("imotions", "run", header_only_imotions),
                backend.AnalysisModalityRunRequest("audio", "run", malformed_audio),
                backend.AnalysisModalityRunRequest("audio", "run", empty_audio_report),
            )
            for modality in invalid_modalities:
                with self.subTest(modality=modality), self.assertRaisesRegex(ValueError, "usable|valid|missing"):
                    backend.discover_analysis_speakers((modality,))

    def test_analysis_speaker_discovery_rejects_sources_without_required_finite_emotions(self) -> None:
        with tempfile.TemporaryDirectory() as temp_dir:
            root = Path(temp_dir)
            metadata_only = root / "metadata-only"
            emotions_disabled = root / "emotions-disabled"
            nonnumeric_imotions = root / "nonnumeric-imotions"
            self._write_metadata_only_audio_csv(
                metadata_only / "Speaker B" / "speech" / "audio_analysis.csv",
                "Speaker B",
            )
            self._write_audio_csv(
                emotions_disabled / "Speaker B" / "speech" / "audio_analysis.csv",
                "Speaker B",
                emotion_value="",
            )
            self._write_imotions_csv(
                nonnumeric_imotions / "Speaker B" / "speech.csv",
                emotion_value="not-a-number",
            )

            invalid_modalities = (
                backend.AnalysisModalityRunRequest("audio", "run", metadata_only),
                backend.AnalysisModalityRunRequest("audio", "run", emotions_disabled),
                backend.AnalysisModalityRunRequest("imotions", "run", nonnumeric_imotions),
            )
            for modality in invalid_modalities:
                with self.subTest(modality=modality), self.assertRaisesRegex(
                    ValueError,
                    "required.*emotion|finite.*emotion|emotion.*metrics",
                ):
                    backend.discover_analysis_speakers((modality,))

    @staticmethod
    def _tree_hash(root: Path) -> str:
        digest = hashlib.sha256()
        for path in sorted(root.rglob("*")):
            digest.update(path.relative_to(root).as_posix().encode("utf-8"))
            if path.is_file():
                digest.update(path.read_bytes())
        return digest.hexdigest()

    @staticmethod
    def _write_sectioned_report(path: Path, metrics: tuple[str, ...]) -> None:
        path.parent.mkdir(parents=True, exist_ok=True)
        lines: list[str] = []
        for metric in metrics:
            lines.extend(
                [
                    metric,
                    "classification,core,category,emotion,unit,score",
                    "metric,001,002,003,004,005",
                    "count,10,10,10,10,10",
                    "missing,0,0,0,0,0",
                    "mean,1,2,3,4,5",
                    "stddev,1,1,1,1,1",
                    "kurtosis,0,0,0,0,0",
                    "",
                ]
            )
        path.write_text("\n".join(lines), encoding="utf-8")

    @staticmethod
    def _write_imotions_csv(path: Path, emotion_value: str = "10") -> None:
        path.parent.mkdir(parents=True, exist_ok=True)
        with path.open("w", encoding="utf-8", newline="") as handle:
            csv.writer(handle).writerows(
                [
                    ["#INFO"],
                    ["#METADATA"],
                    ["#Category", "Timestamp", *("FEA(Emotions)" for _ in VIDEO_METRICS)],
                    ["#Description", "Timestamp", *VIDEO_METRICS],
                    ["#Unit", "Millisecond", *("Index" for _ in VIDEO_METRICS)],
                    ["#Group", "", *("Emotion" for _ in VIDEO_METRICS)],
                    ["#Display name", "", *VIDEO_METRICS],
                    [
                        "#Channel identifier",
                        "Timestamp",
                        *(f"FEA_{metric.replace(' ', '_')}" for metric in VIDEO_METRICS),
                    ],
                    ["#DATA"],
                    ["Row", "Timestamp", *VIDEO_METRICS],
                    [1, 0, *(emotion_value for _ in VIDEO_METRICS)],
                ]
            )

    @staticmethod
    def _write_header_only_imotions_csv(path: Path) -> None:
        path.parent.mkdir(parents=True, exist_ok=True)
        with path.open("w", encoding="utf-8", newline="") as handle:
            csv.writer(handle).writerows(
                [
                    ["#DATA"],
                    ["Row", "Timestamp", "Anger"],
                ]
            )

    @staticmethod
    def _write_audio_csv(path: Path, speaker: str, emotion_value: str = "0.5") -> None:
        path.parent.mkdir(parents=True, exist_ok=True)
        with path.open("w", encoding="utf-8", newline="") as handle:
            csv.writer(handle).writerows(
                [
                    ["#INFO"],
                    ["#SpeakerName", speaker],
                    ["#VideoTitle", "Speech"],
                    ["#DATA"],
                    [
                        "WindowIndex", "StartSeconds", "SpeakerName", "Anger", "Contempt",
                        "Disgust", "Fear", "Happiness", "Sadness", "Surprise", "Neutral",
                        "Other", "Valence", "Arousal", "Dominance",
                    ],
                    ["1", "0", speaker, *(emotion_value for _ in AUDIO_METRICS)],
                ]
            )

    @staticmethod
    def _write_metadata_only_audio_csv(path: Path, speaker: str) -> None:
        path.parent.mkdir(parents=True, exist_ok=True)
        with path.open("w", encoding="utf-8", newline="") as handle:
            csv.writer(handle).writerows(
                [
                    ["#INFO"],
                    ["#SpeakerName", speaker],
                    ["#VideoTitle", "Speech"],
                    ["#DATA"],
                    ["WindowIndex", "StartSeconds", "SpeakerName"],
                    ["1", "0", speaker],
                ]
            )

    @staticmethod
    def _write_text_summary(path: Path) -> None:
        path.parent.mkdir(parents=True, exist_ok=True)
        with path.open("w", encoding="utf-8", newline="") as handle:
            csv.writer(handle).writerows(
                [
                    [
                        "Country",
                        "Speaker",
                        "Speaker ID",
                        "Videos",
                        "Valid segments",
                        "RockSteady terms",
                        "Positive valence",
                        "Negative valence",
                        "Arousal / Activation",
                        "Dominance / Power",
                        "Affiliation / Social orientation",
                    ],
                    ["Group 2", "Speaker C", "speaker_c", 5, 40, 1200, 0.25, 0.10, 0.15, 0.05, 0.20],
                ]
            )

    def test_folder_scan_preserves_speaker_groups_and_reads_sidecar_metadata(self) -> None:
        with tempfile.TemporaryDirectory() as temp_dir:
            root = Path(temp_dir)
            video = root / "Speaker B" / "April 2022.mp4"
            video.parent.mkdir(parents=True)
            video.write_bytes(b"not a real mp4")
            (video.parent / "extraction_metadata.json").write_text(
                json.dumps(
                    {
                        "title": "Interview April",
                        "url": "https://www.youtube.com/watch?v=abcdefghijk",
                        "license": "Standard YouTube License",
                        "upload_date": "2022-04-12",
                        "duration": 512,
                    }
                ),
                encoding="utf-8",
            )

            result = backend.scan_input_source(root, duration_reader=lambda _path: None)

        self.assertEqual(result.source_kind, "folder")
        self.assertEqual(len(result.groups), 1)
        group = result.groups[0]
        self.assertEqual(group.speaker, "Speaker B")
        self.assertEqual(group.videos[0].title, "Interview April")
        self.assertEqual(group.videos[0].duration_seconds, 512)
        self.assertEqual(group.videos[0].license, "Standard YouTube License")
        self.assertEqual(group.videos[0].youtube_url, "https://www.youtube.com/watch?v=abcdefghijk")

    def test_scan_merges_case_and_whitespace_variants_into_one_speaker_group(self) -> None:
        items = [
            backend.VideoItem("1", "One", "Speaker B", "one.mp4", "folder"),
            backend.VideoItem("2", "Two", "  speaker   b ", "two.mp4", "folder"),
        ]

        groups = backend.group_video_items_by_speaker(items)

        self.assertEqual(len(groups), 1)
        self.assertEqual(groups[0].speaker, "Speaker B")
        self.assertEqual([video.id for video in groups[0].videos], ["1", "2"])

    def test_single_video_scan_accepts_local_file_path(self) -> None:
        with tempfile.TemporaryDirectory() as temp_dir:
            root = Path(temp_dir)
            video = root / "Speaker D" / "speech.mp4"
            video.parent.mkdir(parents=True)
            video.write_bytes(b"not a real mp4")

            result = backend.scan_input_source(video, duration_reader=lambda _path: 123.0)

        self.assertEqual(result.source_kind, "file")
        self.assertEqual(result.groups[0].speaker, "Speaker D")
        item = result.groups[0].videos[0]
        self.assertEqual(item.source_kind, "file")
        self.assertEqual(item.source_path, str(video.resolve()))
        self.assertEqual(item.duration_seconds, 123.0)

    def test_docx_scan_reads_youtube_rows_grouped_by_speaker(self) -> None:
        with tempfile.TemporaryDirectory() as temp_dir:
            docx_path = Path(temp_dir) / "videos.docx"
            doc = Document()
            table = doc.add_table(rows=2, cols=5)
            table.rows[0].cells[0].text = "Link"
            table.rows[0].cells[1].text = "Speaker"
            table.rows[0].cells[2].text = "Length"
            table.rows[0].cells[3].text = "Date Uploaded"
            table.rows[0].cells[4].text = "License"
            table.rows[1].cells[0].text = "https://www.youtube.com/watch?v=abcdefghijk"
            table.rows[1].cells[1].text = "Speaker B"
            table.rows[1].cells[2].text = "00:07:16"
            table.rows[1].cells[3].text = "2022-04-12"
            table.rows[1].cells[4].text = "Creative Commons"
            doc.save(docx_path)

            result = backend.scan_input_source(docx_path)

        self.assertEqual(result.source_kind, "catalog")
        self.assertEqual(result.groups[0].speaker, "Speaker B")
        video = result.groups[0].videos[0]
        self.assertEqual(video.video_id, "abcdefghijk")
        self.assertIsNone(video.duration_seconds)
        self.assertEqual(video.metadata["Date Uploaded"], "2022-04-12")
        self.assertEqual(video.metadata["License"], "Creative Commons")
        self.assertNotIn("Length", video.metadata)

    def test_docx_scan_accepts_quoted_pasted_windows_path(self) -> None:
        with tempfile.TemporaryDirectory() as temp_dir:
            docx_path = Path(temp_dir) / "videos.docx"
            doc = Document()
            table = doc.add_table(rows=2, cols=2)
            table.rows[0].cells[0].text = "Link"
            table.rows[0].cells[1].text = "Speaker"
            table.rows[1].cells[0].text = "https://www.youtube.com/watch?v=abcdefghijk"
            table.rows[1].cells[1].text = "Speaker B"
            doc.save(docx_path)

            result = backend.scan_input_source(f'"{docx_path}"', enrich_youtube=False)

        self.assertEqual(result.source_kind, "catalog")
        self.assertEqual(result.source_path, str(docx_path.resolve()))
        self.assertEqual(result.groups[0].videos[0].video_id, "abcdefghijk")

    def test_clean_user_supplied_path_only_removes_wrapping_quotes(self) -> None:
        self.assertEqual(backend.clean_user_supplied_path('  "C:\\Temp\\videos.docx"  '), r"C:\Temp\videos.docx")
        self.assertEqual(backend.clean_user_supplied_path("C:\\Temp\\Bob's Videos"), r"C:\Temp\Bob's Videos")

    def test_scan_accepts_direct_youtube_url(self) -> None:
        result = backend.scan_input_source(
            "https://youtu.be/abcdefghijk?t=30",
            enrich_youtube=False,
        )

        self.assertEqual(result.source_kind, "youtube")
        self.assertEqual(result.source_path, "https://www.youtube.com/watch?v=abcdefghijk")
        video = result.groups[0].videos[0]
        self.assertEqual(video.video_id, "abcdefghijk")
        self.assertEqual(video.source_kind, "youtube")

    def test_direct_youtube_scan_uses_public_title_metadata_without_api_key(self) -> None:
        public_metadata = {
            "abcdefghijk": {
                "title": "A real YouTube title",
                "thumbnail_url": "https://example.test/thumb.jpg",
            }
        }
        with (
            patch.object(backend, "load_youtube_api_key", return_value=""),
            patch.object(backend, "fetch_youtube_oembed_batch", return_value=public_metadata),
            patch.object(
                backend,
                "fetch_youtube_ytdlp_metadata",
                return_value={"duration_seconds": 125.0, "upload_date": "2026-07-01"},
            ),
        ):
            result = backend.scan_input_source("https://www.youtube.com/watch?v=abcdefghijk")

        video = result.groups[0].videos[0]
        self.assertEqual(video.title, "A real YouTube title")
        self.assertEqual(video.thumbnail_url, "https://example.test/thumb.jpg")
        self.assertEqual(video.duration_seconds, 125.0)
        self.assertEqual(video.upload_date, "2026-07-01")

    def test_scan_rejects_empty_folder_with_supported_formats(self) -> None:
        with tempfile.TemporaryDirectory() as temp_dir:
            with self.assertRaisesRegex(ValueError, "No supported videos"):
                backend.scan_input_source(temp_dir)

    def test_scan_rejects_unsupported_single_file_with_supported_formats(self) -> None:
        with tempfile.TemporaryDirectory() as temp_dir:
            source = Path(temp_dir) / "videos.txt"
            source.write_text("not a video", encoding="utf-8")

            with self.assertRaisesRegex(ValueError, "Unsupported input file type"):
                backend.scan_input_source(source)

    def test_scan_rejects_docx_without_youtube_links(self) -> None:
        with tempfile.TemporaryDirectory() as temp_dir:
            source = Path(temp_dir) / "videos.docx"
            document = Document()
            table = document.add_table(rows=2, cols=2)
            table.rows[0].cells[0].text = "Link"
            table.rows[0].cells[1].text = "Speaker"
            table.rows[1].cells[0].text = "No link supplied"
            table.rows[1].cells[1].text = "Speaker"
            document.save(source)

            with self.assertRaisesRegex(FileNotFoundError, "Catalog local source"):
                backend.scan_input_source(source, enrich_youtube=False)

    def test_run_request_rejects_percentage_outside_zero_to_one(self) -> None:
        request = backend.RunRequest(
            mode="standard",
            source_path=Path("videos"),
            output_root=Path("outputs"),
            percentage=1.5,
        )

        with self.assertRaisesRegex(ValueError, "Sample percentage"):
            backend.build_run_command(request, repo_root=Path("repo"))

    def test_clean_display_title_uses_short_youtube_fallback_for_url_titles(self) -> None:
        title = backend.clean_display_title(
            "https://www.youtube.com/watch?v=abcdefghijk",
            "https://www.youtube.com/watch?v=abcdefghijk",
        )

        self.assertEqual(title, "Title unavailable [abcdefghijk]")

    def test_clean_display_title_uses_short_fallback_for_youtu_be_titles(self) -> None:
        title = backend.clean_display_title(
            "https://youtu.be/abcdefghijk",
            "https://youtu.be/abcdefghijk",
        )

        self.assertEqual(title, "Title unavailable [abcdefghijk]")

    def test_unavailable_title_remains_eligible_for_metadata_fallback(self) -> None:
        self.assertTrue(backend.title_looks_like_youtube_reference("Title unavailable [abcdefghijk]"))

    def test_youtube_iso8601_duration_parser(self) -> None:
        self.assertEqual(backend.parse_youtube_iso8601_duration("PT1H2M3S"), 3723)
        self.assertEqual(backend.parse_youtube_iso8601_duration("PT7M16S"), 436)

    def test_youtube_metadata_replaces_url_title_and_unknowns(self) -> None:
        item = backend.VideoItem(
            id="docx:abcdefghijk",
            title="https://www.youtube.com/watch?v=abcdefghijk",
            speaker="Speaker",
            source_path="videos.docx",
            source_kind="docx",
            youtube_url="https://www.youtube.com/watch?v=abcdefghijk",
            video_id="abcdefghijk",
            duration_seconds=None,
            upload_date="",
            license="Unknown",
        )

        updated = backend.apply_youtube_metadata(
            [item],
            {
                "abcdefghijk": {
                    "title": "Real YouTube Title",
                    "duration_seconds": 436,
                    "upload_date": "2022-04-12",
                    "license": "Standard YouTube License",
                    "thumbnail_url": "https://example.test/thumb.jpg",
                }
            },
        )

        self.assertEqual(updated[0].title, "Real YouTube Title")
        self.assertEqual(updated[0].duration_seconds, 436)
        self.assertEqual(updated[0].upload_date, "2022-04-12")
        self.assertEqual(updated[0].license, "Standard YouTube License")
        self.assertEqual(updated[0].thumbnail_url, "https://example.test/thumb.jpg")

    def test_scan_json_includes_youtube_thumbnail_url(self) -> None:
        item = backend.VideoItem(
            id="docx:abcdefghijk",
            title="Video",
            speaker="Speaker",
            source_path="videos.docx",
            source_kind="docx",
            youtube_url="https://www.youtube.com/watch?v=abcdefghijk",
            video_id="abcdefghijk",
        )
        result = backend.ScanResult(
            source_path="videos.docx",
            source_kind="docx",
            groups=[backend.SpeakerGroup(speaker="Speaker", videos=[item])],
        )

        payload = backend.scan_result_to_json(result)

        video = payload["groups"][0]["videos"][0]
        self.assertEqual(video["thumbnail_url"], "https://i.ytimg.com/vi/abcdefghijk/hqdefault.jpg")

    def test_imotions_transcode_command_enforces_canonical_media(self) -> None:
        command = backend.build_imotions_transcode_command(
            Path(r"C:\input\clip.mov"),
            Path(r"C:\output\stitched_imotions.mp4"),
            start_seconds=1.25,
            duration_seconds=4.5,
        )

        self.assertEqual(command[command.index("-c:v") + 1], "libx264")
        self.assertEqual(command[command.index("-pix_fmt") + 1], "yuv420p")
        self.assertEqual(command[command.index("-r") + 1], "30")
        self.assertEqual(command[command.index("-ar") + 1], "48000")
        self.assertEqual(command[command.index("-ac") + 1], "2")
        self.assertEqual(command[command.index("-avoid_negative_ts") + 1], "make_zero")

    def test_imotions_concat_command_regenerates_audio_timestamps(self) -> None:
        command = backend.build_imotions_concat_command(
            Path(r"C:\output\segments.txt"),
            Path(r"C:\output\stitched_imotions.mp4"),
        )

        self.assertEqual(command[command.index("-c:v") + 1], "libx264")
        self.assertEqual(command[command.index("-r") + 1], "30")
        self.assertEqual(command[command.index("-fps_mode") + 1], "cfr")
        self.assertEqual(command[command.index("-c:a") + 1], "aac")
        self.assertEqual(command[command.index("-fflags") + 1], "+genpts")
        self.assertEqual(command[command.index("-avoid_negative_ts") + 1], "make_zero")

    def test_build_standard_docx_catalog_command_uses_source_ids_and_digest(self) -> None:
        repo_root = Path(r"C:\repo")
        command = backend.build_run_command(
            backend.RunRequest(
                mode="standard",
                source_path=Path(r"C:\input\videos.docx"),
                output_root=Path(r"C:\output"),
                percentage=0.25,
                max_segment_seconds=45,
                selected_ids=["source-0002", "source-0005"],
                catalog_sha256="a" * 64,
            ),
            repo_root=repo_root,
            python_executable=Path(r"C:\Python312\python.exe"),
        )

        self.assertEqual(command[:3], [r"C:\Python312\python.exe", "-m", "procurement.catalog_runner"])
        self.assertIn(r"C:\input\videos.docx", command)
        self.assertEqual(command[command.index("--percentage") + 1], "0.25")
        self.assertEqual(command[command.index("--max-segment-seconds") + 1], "45")
        self.assertEqual(
            [command[index + 1] for index, value in enumerate(command) if value == "--source-id"],
            ["source-0002", "source-0005"],
        )

    def test_prepare_docx_source_for_run_uses_local_cache_copy(self) -> None:
        with tempfile.TemporaryDirectory() as temp_dir:
            root = Path(temp_dir)
            source = root / "OneDrive Source.docx"
            doc = Document()
            table = doc.add_table(rows=1, cols=1)
            table.rows[0].cells[0].text = "https://www.youtube.com/watch?v=abcdefghijk"
            doc.save(source)

            prepared = backend.prepare_source_for_run(source, root / "repo")
            backend.run_docx_extractions.open_docx_document(prepared)

            self.assertIn("_local", prepared.parts)
            self.assertEqual(prepared.name.split("_", 1)[0], "OneDrive")

    def test_prepare_direct_youtube_source_materialises_pipeline_docx(self) -> None:
        with tempfile.TemporaryDirectory() as temp_dir:
            repo_root = Path(temp_dir)

            prepared = backend.prepare_source_for_run(
                "https://www.youtube.com/watch?v=abcdefghijk",
                repo_root,
            )
            result = backend.scan_input_source(prepared, enrich_youtube=False)

        self.assertEqual(prepared.name, "youtube_abcdefghijk.docx")
        self.assertEqual(result.groups[0].videos[0].video_id, "abcdefghijk")

    def test_normalise_settings_keeps_youtube_cookie_browser_choice_local(self) -> None:
        settings = backend.normalise_ui_settings(
            {
                "youtubeApiKey": " key ",
                "huggingFaceToken": " hf ",
                "youtubeCookiesBrowser": "Edge",
            }
        )

        self.assertEqual(settings["youtubeCookiesBrowser"], "edge")

    def test_normalise_settings_drops_unknown_cookie_browser_choice(self) -> None:
        settings = backend.normalise_ui_settings({"youtubeCookiesBrowser": "suspicious-browser"})

        self.assertEqual(settings["youtubeCookiesBrowser"], "")

    def test_build_full_docx_catalog_command_preserves_full_mode(self) -> None:
        command = backend.build_run_command(
            backend.RunRequest(
                mode="full",
                source_path=Path(r"C:\input\videos.docx"),
                output_root=Path(r"C:\output"),
                selected_ids=["source-0001"],
                catalog_sha256="a" * 64,
            ),
            repo_root=Path(r"C:\repo"),
            python_executable=Path(r"C:\Python312\python.exe"),
        )

        self.assertEqual(command[:3], [r"C:\Python312\python.exe", "-m", "procurement.catalog_runner"])
        self.assertIn(r"C:\input\videos.docx", command)
        self.assertEqual(command[command.index("--mode") + 1], "full")
        self.assertEqual(command[command.index("--source-id") + 1], "source-0001")

    def test_build_local_standard_command_forwards_selected_speaker_groups(self) -> None:
        command = backend.build_run_command(
            backend.RunRequest(
                mode="standard",
                source_path=Path(r"C:\input\videos"),
                output_root=Path(r"C:\output"),
                selected_speakers=["Speaker A"],
            ),
            repo_root=Path(r"C:\repo"),
            python_executable=Path(r"C:\Python312\python.exe"),
        )

        self.assertEqual(command[command.index("--speaker") + 1], "Speaker A")

    def test_manual_mode_writes_segment_manifest_command(self) -> None:
        command = backend.build_run_command(
            backend.RunRequest(
                mode="manual",
                source_path=Path(r"C:\videos"),
                output_root=Path(r"C:\output"),
                segment_manifest=Path(r"C:\segments\selection.json"),
                segment_manifest_sha256="a" * 64,
                segment_expected_source=r"C:\videos",
            ),
            repo_root=Path(r"C:\repo"),
            python_executable=Path(r"C:\Python312\python.exe"),
        )

        self.assertEqual(command[:3], [r"C:\Python312\python.exe", "-m", "application.manual_segments"])
        self.assertIn(r"C:\segments\selection.json", command)
        self.assertIn("a" * 64, command)
        self.assertIn(r"C:\videos", command)

    def test_clean_speaker_beta_command_maps_all_ui_options(self) -> None:
        command = backend.build_run_command(
            backend.RunRequest(
                mode="clean-speaker-beta",
                source_path=Path(r"C:\videos"),
                output_root=Path(r"C:\output"),
                percentage=0.25,
                max_segment_seconds=45,
                beta_output_mode="percentage",
                beta_min_clean_seconds=12.0,
                beta_gap_seconds=2.5,
                beta_identity_stills=18,
                beta_scan_fps=1.5,
                beta_validation_fps=2.5,
                beta_max_download_height=480,
                beta_face_confidence=0.72,
                beta_speaker_confidence=0.68,
                beta_worker_count=3,
                beta_device="cpu",
                beta_keep_debug=True,
                beta_resource_guard_percent=12.0,
                beta_resource_poll_seconds=3.0,
                beta_resource_guard_timeout_seconds=120.0,
                beta_parallel_detector_streams=True,
                beta_reference_audio=Path(r"C:\profiles\speaker.wav"),
                beta_only_video_ids=["abc123", "xyz789"],
                beta_random_one=True,
                beta_random_seed="repeatable",
                beta_isolated_video_processes=True,
                beta_skip_first_videos=12,
                beta_skip_completed_outputs=True,
                beta_video_cooldown_seconds=45.0,
                beta_max_affinity_cores=2,
                beta_native_threads=1,
                beta_cpu_throttle_high_percent=96.0,
                beta_cpu_throttle_low_percent=91.0,
                beta_ram_throttle_high_percent=97.0,
                beta_ram_throttle_low_percent=92.0,
                selected_speakers=["Speaker A"],
            ),
            repo_root=Path(r"C:\repo"),
            python_executable=Path(r"C:\Python312\python.exe"),
        )

        self.assertEqual(command[:3], [r"C:\Python312\python.exe", "-m", "procurement.procurement_beta.cli"])
        self.assertIn("--output-mode", command)
        self.assertIn("percentage", command)
        self.assertIn("--min-clean-seconds", command)
        self.assertIn("12.0", command)
        self.assertIn("--gap-seconds", command)
        self.assertIn("2.5", command)
        self.assertIn("--identity-stills", command)
        self.assertIn("18", command)
        self.assertIn("--scan-fps", command)
        self.assertIn("1.5", command)
        self.assertIn("--validation-fps", command)
        self.assertIn("2.5", command)
        self.assertIn("--max-download-height", command)
        self.assertIn("480", command)
        self.assertIn("--face-confidence", command)
        self.assertIn("0.72", command)
        self.assertIn("--speaker-confidence", command)
        self.assertIn("0.68", command)
        self.assertIn("--workers", command)
        self.assertIn("3", command)
        self.assertIn("--device", command)
        self.assertIn("cpu", command)
        self.assertIn("--keep-debug", command)
        self.assertIn("--resource-guard-percent", command)
        self.assertIn("12.0", command)
        self.assertIn("--resource-poll-seconds", command)
        self.assertIn("3.0", command)
        self.assertIn("--resource-guard-timeout-seconds", command)
        self.assertIn("120.0", command)
        self.assertIn("--parallel-detectors", command)
        self.assertIn("--reference-audio", command)
        self.assertIn(r"C:\profiles\speaker.wav", command)
        self.assertEqual(command.count("--only-video-id"), 2)
        self.assertIn("abc123", command)
        self.assertIn("xyz789", command)
        self.assertIn("--random-one", command)
        self.assertIn("--random-seed", command)
        self.assertIn("repeatable", command)
        self.assertIn("--isolated-video-processes", command)
        self.assertIn("--skip-first-videos", command)
        self.assertIn("12", command)
        self.assertIn("--skip-completed-outputs", command)
        self.assertIn("--video-cooldown-seconds", command)
        self.assertIn("45.0", command)
        self.assertIn("--max-affinity-cores", command)
        self.assertIn("--native-threads", command)
        self.assertIn("--cpu-throttle-high-percent", command)
        self.assertIn("96.0", command)
        self.assertIn("--cpu-throttle-low-percent", command)
        self.assertIn("91.0", command)
        self.assertIn("--ram-throttle-high-percent", command)
        self.assertIn("97.0", command)
        self.assertIn("--ram-throttle-low-percent", command)
        self.assertIn("92.0", command)
        self.assertEqual(command[command.index("--speaker") + 1], "Speaker A")

    def test_clean_speaker_beta_defaults_to_one_worker_without_hard_ui_cap(self) -> None:
        command = backend.build_run_command(
            backend.RunRequest(
                mode="clean-speaker-beta",
                source_path=Path(r"C:\videos"),
                output_root=Path(r"C:\output"),
            ),
            repo_root=Path(r"C:\repo"),
            python_executable=Path(r"C:\Python312\python.exe"),
        )

        worker_index = command.index("--workers")
        self.assertEqual(command[worker_index + 1], "1")
        height_index = command.index("--max-download-height")
        self.assertEqual(command[height_index + 1], "720")
        guard_index = command.index("--resource-guard-percent")
        self.assertEqual(command[guard_index + 1], "15.0")
        timeout_index = command.index("--resource-guard-timeout-seconds")
        self.assertEqual(command[timeout_index + 1], "900.0")
        self.assertNotIn("--parallel-detectors", command)
        self.assertIn("--isolated-video-processes", command)
        self.assertIn("--skip-completed-outputs", command)
        cooldown_index = command.index("--video-cooldown-seconds")
        self.assertEqual(command[cooldown_index + 1], "60.0")
        affinity_index = command.index("--max-affinity-cores")
        self.assertEqual(command[affinity_index + 1], "2")
        native_threads_index = command.index("--native-threads")
        self.assertEqual(command[native_threads_index + 1], "1")

    def test_audio_batch_command_maps_ui_options_to_audio_pipeline(self) -> None:
        command = backend.build_audio_command(
            backend.AudioRunRequest(
                mode="batch",
                source_path=Path(r"C:\downloads"),
                output_root=Path(r"C:\audio-out"),
                window_seconds=12.5,
                stride_seconds=4.0,
                opensmile_feature_set="compare16",
                include_emotions=False,
                device="cpu",
                keep_temp_audio=True,
                debug=True,
                stop_on_error=True,
                selected_source_ids=("source-0003", "source-0001"),
                catalog_sha256="a" * 64,
            ),
            repo_root=Path(r"C:\repo"),
            python_executable=Path(r"C:\Python312\python.exe"),
        )

        self.assertEqual(command[:3], [r"C:\Python312\python.exe", r"C:\repo\processing\audio_analysis\run_audio_analysis.py", "batch"])
        self.assertIn(r"C:\downloads", command)
        self.assertIn(r"C:\audio-out", command)
        self.assertIn("--skip-emotion-models", command)
        self.assertIn("--keep-temp-audio", command)
        self.assertIn("--debug", command)
        self.assertIn("--stop-on-error", command)
        self.assertIn("compare16", command)
        self.assertEqual(
            [command[index + 1] for index, value in enumerate(command) if value == "--source-id"],
            ["source-0003", "source-0001"],
        )
        self.assertEqual(command[command.index("--catalog-sha256") + 1], "a" * 64)

    def test_audio_single_command_keeps_emotions_enabled_by_default(self) -> None:
        command = backend.build_audio_command(
            backend.AudioRunRequest(
                mode="single",
                source_path=Path(r"C:\video\stitched_imotions.mp4"),
                output_root=Path(r"C:\audio-out"),
            ),
            repo_root=Path(r"C:\repo"),
            python_executable=Path(r"C:\Python312\python.exe"),
        )

        self.assertEqual(command[2], "single")
        self.assertNotIn("--skip-emotion-models", command)
        self.assertNotIn("--stop-on-error", command)

    def test_audio_command_rejects_invalid_advanced_options(self) -> None:
        invalid_requests = [
            backend.AudioRunRequest(
                mode="batch",
                source_path=Path(r"C:\downloads"),
                output_root=Path(r"C:\audio-out"),
                window_seconds=0.1,
            ),
            backend.AudioRunRequest(
                mode="batch",
                source_path=Path(r"C:\downloads"),
                output_root=Path(r"C:\audio-out"),
                opensmile_feature_set="unknown",
            ),
            backend.AudioRunRequest(
                mode="batch",
                source_path=Path(r"C:\downloads"),
                output_root=Path(r"C:\audio-out"),
                device="quantum",
            ),
        ]

        for request in invalid_requests:
            with self.subTest(request=request):
                with self.assertRaises(ValueError):
                    backend.build_audio_command(request, repo_root=Path(r"C:\repo"))

    def test_audio_command_rejects_invalid_duplicate_or_single_mode_source_ids(self) -> None:
        invalid_requests = (
            backend.AudioRunRequest(
                mode="batch",
                source_path=Path(r"C:\downloads"),
                output_root=Path(r"C:\audio-out"),
                selected_source_ids=("not-a-source-id",),
            ),
            backend.AudioRunRequest(
                mode="batch",
                source_path=Path(r"C:\downloads"),
                output_root=Path(r"C:\audio-out"),
                selected_source_ids=("source-0001", "source-0001"),
            ),
            backend.AudioRunRequest(
                mode="single",
                source_path=Path(r"C:\video.mp4"),
                output_root=Path(r"C:\audio-out"),
                selected_source_ids=("source-0001",),
                catalog_sha256="a" * 64,
            ),
            backend.AudioRunRequest(
                mode="batch",
                source_path=Path(r"C:\downloads"),
                output_root=Path(r"C:\audio-out"),
                selected_source_ids=("source-0001",),
                catalog_sha256="wrong",
            ),
        )

        for request in invalid_requests:
            with self.subTest(request=request), self.assertRaises(ValueError):
                backend.build_audio_command(request, repo_root=Path(r"C:\repo"))

    def test_eula_file_defaults_to_false_when_missing(self) -> None:
        with tempfile.TemporaryDirectory() as temp_dir:
            repo_root = Path(temp_dir)

            state = backend.load_eula_state(repo_root)
            eula_text = backend.eula_path(repo_root).read_text(encoding="utf-8")

        self.assertFalse(state["termsAccepted"])
        self.assertIn("terms_accepted=false", eula_text)

    def test_eula_true_gets_acceptance_timestamp_data(self) -> None:
        with tempfile.TemporaryDirectory() as temp_dir:
            repo_root = Path(temp_dir)
            backend.eula_path(repo_root).parent.mkdir(parents=True)
            backend.eula_path(repo_root).write_text("terms_accepted=true\n", encoding="utf-8")

            state = backend.load_eula_state(
                repo_root,
                now=datetime(2026, 6, 22, 16, 30, tzinfo=timezone.utc),
            )
            eula_text = backend.eula_path(repo_root).read_text(encoding="utf-8")

        self.assertTrue(state["termsAccepted"])
        self.assertEqual(state["acceptedAt"], "2026-06-22T16:30:00Z")
        self.assertIn("# data: accepted_at=2026-06-22T16:30:00Z", eula_text)

    def test_eula_revoke_writes_false_and_clears_timestamp(self) -> None:
        with tempfile.TemporaryDirectory() as temp_dir:
            repo_root = Path(temp_dir)

            backend.write_eula_state(repo_root, True, accepted_at="2026-06-22T16:30:00Z")
            state = backend.write_eula_state(repo_root, False)
            eula_text = backend.eula_path(repo_root).read_text(encoding="utf-8")

        self.assertFalse(state["termsAccepted"])
        self.assertIn("# data: accepted_at=", eula_text)
        self.assertIn("terms_accepted=false", eula_text)

    def test_ui_settings_store_secrets_outside_primary_and_backup_json(self) -> None:
        with tempfile.TemporaryDirectory() as temp_dir:
            repo_root = Path(temp_dir)
            credential_root = repo_root / "protected-credentials"
            with patch.dict(os.environ, {"MEA_CREDENTIAL_STORE_ROOT": str(credential_root)}):
                saved = backend.save_ui_settings(
                    repo_root,
                    {
                        "youtubeApiKey": "  'abc123'  ",
                        "unexpected": "ignored",
                    },
                )
                backend.save_ui_settings(repo_root, {"maxCpuCores": 2})
                loaded = backend.load_ui_settings(repo_root)
                stored_key = backend.load_youtube_api_key(settings_path=backend.ui_settings_path(repo_root))
            primary = backend.ui_settings_path(repo_root).read_text(encoding="utf-8")
            backup = backend.settings_backup_path(backend.ui_settings_path(repo_root)).read_text(encoding="utf-8")

        self.assertNotIn("youtubeApiKey", saved)
        self.assertEqual(loaded, saved | {"maxCpuCores": 2})
        self.assertNotIn("unexpected", loaded)
        self.assertEqual(stored_key, "abc123")
        self.assertNotIn("abc123", primary + backup)
        self.assertNotIn("youtubeApiKey", primary + backup)

    def test_public_ui_settings_masks_secrets_and_never_returns_raw_values(self) -> None:
        with tempfile.TemporaryDirectory() as temp_dir:
            repo_root = Path(temp_dir)
            with patch.dict(os.environ, {"MEA_CREDENTIAL_STORE_ROOT": str(repo_root / "credentials")}):
                backend.save_ui_settings(
                    repo_root,
                    {
                        "youtubeApiKey": "youtube-secret-1234",
                        "huggingFaceToken": "huggingface-secret-5678",
                    },
                )

                with patch.dict(
                    "os.environ",
                    {
                        "YOUTUBE_API_KEY": "",
                        "HF_TOKEN": "",
                        "HUGGINGFACE_TOKEN": "",
                        "HUGGING_FACE_HUB_TOKEN": "",
                        "MEA_CREDENTIAL_STORE_ROOT": str(repo_root / "credentials"),
                    },
                ):
                    public = backend.public_ui_settings(repo_root)

        self.assertNotIn("youtubeApiKey", public)
        self.assertNotIn("huggingFaceToken", public)
        self.assertTrue(public["youtubeApiKeyConfigured"])
        self.assertEqual(public["youtubeApiKeyMasked"], "********1234")
        self.assertTrue(public["huggingFaceTokenConfigured"])
        self.assertEqual(public["huggingFaceTokenMasked"], "********5678")

    def test_legacy_primary_and_backup_secrets_migrate_then_are_scrubbed(self) -> None:
        with tempfile.TemporaryDirectory() as temp_dir:
            repo_root = Path(temp_dir)
            settings_path = backend.ui_settings_path(repo_root)
            settings_path.parent.mkdir(parents=True)
            settings_path.write_text(
                json.dumps({"youtubeApiKey": "legacy-youtube", "maxCpuCores": 3}),
                encoding="utf-8",
            )
            backup_path = backend.settings_backup_path(settings_path)
            backup_path.write_text(
                json.dumps({"huggingFaceToken": "legacy-huggingface", "maxCpuCores": 2}),
                encoding="utf-8",
            )

            with patch.dict(
                os.environ,
                {
                    "YOUTUBE_API_KEY": "",
                    "HF_TOKEN": "",
                    "HUGGINGFACE_TOKEN": "",
                    "HUGGING_FACE_HUB_TOKEN": "",
                    "MEA_CREDENTIAL_STORE_ROOT": str(repo_root / "credentials"),
                },
            ):
                loaded = backend.load_ui_settings(repo_root)
                youtube_key = backend.load_youtube_api_key(settings_path=settings_path)
                hf_token = backend.load_huggingface_token(settings_path=settings_path)

            persisted = settings_path.read_text(encoding="utf-8") + backup_path.read_text(encoding="utf-8")

        self.assertEqual(loaded["maxCpuCores"], 3)
        self.assertEqual(youtube_key, "legacy-youtube")
        self.assertEqual(hf_token, "legacy-huggingface")
        for forbidden in ("legacy-youtube", "legacy-huggingface", "youtubeApiKey", "huggingFaceToken"):
            self.assertNotIn(forbidden, persisted)

    def test_oversized_legacy_primary_and_backup_are_purged_to_safe_defaults(self) -> None:
        with tempfile.TemporaryDirectory() as temp_dir:
            repo_root = Path(temp_dir)
            settings_path = backend.ui_settings_path(repo_root)
            settings_path.parent.mkdir(parents=True)
            backup_path = backend.settings_backup_path(settings_path)
            settings_path.write_text(
                json.dumps({"youtubeApiKey": "oversized-youtube-secret", "padding": "x" * 70_000}),
                encoding="utf-8",
            )
            backup_path.write_text(
                json.dumps({"huggingFaceToken": "oversized-hf-secret", "padding": "x" * 70_000}),
                encoding="utf-8",
            )

            loaded = backend.load_ui_settings(repo_root)
            warning = backend.public_ui_settings(repo_root)["settingsWarning"]
            persisted = settings_path.read_bytes() + backup_path.read_bytes()

        self.assertEqual(loaded, backend.DEFAULT_UI_SETTINGS)
        self.assertIn("purged", str(warning).lower())
        for forbidden in (b"oversized-youtube-secret", b"oversized-hf-secret", b"youtubeApiKey", b"huggingFaceToken"):
            self.assertNotIn(forbidden, persisted)

    def test_invalid_utf8_legacy_settings_are_purged_without_aborting_startup(self) -> None:
        with tempfile.TemporaryDirectory() as temp_dir:
            repo_root = Path(temp_dir)
            settings_path = backend.ui_settings_path(repo_root)
            settings_path.parent.mkdir(parents=True)
            backup_path = backend.settings_backup_path(settings_path)
            settings_path.write_bytes(b'\xff{"youtubeApiKey":"invalid-primary-secret"}')
            backup_path.write_bytes(b'\xfe{"huggingFaceToken":"invalid-backup-secret"}')

            loaded = backend.load_ui_settings(repo_root)
            warning = backend.public_ui_settings(repo_root)["settingsWarning"]
            persisted = settings_path.read_bytes() + backup_path.read_bytes()

        self.assertEqual(loaded, backend.DEFAULT_UI_SETTINGS)
        self.assertIn("purged", str(warning).lower())
        for forbidden in (b"invalid-primary-secret", b"invalid-backup-secret", b"youtubeApiKey", b"huggingFaceToken"):
            self.assertNotIn(forbidden, persisted)

    def test_blank_secret_update_preserves_existing_value_and_explicit_clear_removes_it(self) -> None:
        with tempfile.TemporaryDirectory() as temp_dir:
            repo_root = Path(temp_dir)
            with patch.dict(os.environ, {"MEA_CREDENTIAL_STORE_ROOT": str(repo_root / "credentials")}):
                backend.save_ui_settings(repo_root, {"youtubeApiKey": "keep-me"})
                backend.save_ui_settings(repo_root, {"maxCpuCores": 2})
                backend.save_ui_settings(repo_root, {"youtubeApiKey": ""})
                preserved = backend.load_youtube_api_key(settings_path=backend.ui_settings_path(repo_root))
                backend.save_ui_settings(repo_root, {"clearYouTubeApiKey": True})
                cleared = backend.load_youtube_api_key(settings_path=backend.ui_settings_path(repo_root))
            persisted = backend.ui_settings_path(repo_root).read_text(encoding="utf-8")
            persisted += backend.settings_backup_path(backend.ui_settings_path(repo_root)).read_text(encoding="utf-8")

        self.assertEqual(preserved, "keep-me")
        self.assertEqual(cleared, "")
        self.assertNotIn("keep-me", persisted)

    def test_resource_settings_are_normalised_to_release_bounds(self) -> None:
        settings = backend.normalise_ui_settings(
            {
                "maxCpuPercent": 3,
                "maxCpuCores": 999,
                "maxGpuPercent": 400,
                "ramLimitMode": "GB",
                "maxRamPercent": 99,
                "maxRamGb": 0,
                "nativeThreads": 0,
                "resourcePollSeconds": 0,
            }
        )

        self.assertEqual(settings["maxCpuPercent"], 10.0)
        self.assertEqual(settings["maxCpuCores"], 256)
        self.assertEqual(settings["maxGpuPercent"], 100.0)
        self.assertEqual(settings["ramLimitMode"], "gb")
        self.assertEqual(settings["maxRamPercent"], 95.0)
        self.assertEqual(settings["maxRamGb"], 1.0)
        self.assertEqual(settings["nativeThreads"], 1)
        self.assertEqual(settings["resourcePollSeconds"], 0.5)

    def test_global_resource_settings_drive_clean_speaker_safeguards(self) -> None:
        settings = backend.clean_speaker_resource_settings(
            {
                "resourceLimitsEnabled": True,
                "maxCpuPercent": 82,
                "maxCpuCores": 6,
                "maxGpuPercent": 91,
                "ramLimitMode": "percent",
                "maxRamPercent": 88,
                "nativeThreads": 3,
                "resourcePollSeconds": 4,
            }
        )

        self.assertEqual(settings["resource_guard_percent"], 18.0)
        self.assertEqual(settings["max_affinity_cores"], 6)
        self.assertEqual(settings["native_threads"], 3)
        self.assertEqual(settings["cpu_high_percent"], 82.0)
        self.assertEqual(settings["cpu_low_percent"], 77.0)
        self.assertEqual(settings["ram_high_percent"], 88.0)
        self.assertEqual(settings["resource_poll_seconds"], 4.0)

    def test_disabling_global_resource_limits_disables_clean_speaker_guards(self) -> None:
        settings = backend.clean_speaker_resource_settings(
            {
                "resourceLimitsEnabled": False,
                "maxCpuCores": 2,
                "nativeThreads": 1,
            }
        )

        self.assertEqual(settings["resource_guard_percent"], 0.0)
        self.assertEqual(settings["max_affinity_cores"], 0)
        self.assertEqual(settings["cpu_high_percent"], 100.0)
        self.assertEqual(settings["ram_high_percent"], 100.0)

    def test_clean_speaker_command_preserves_valid_zero_options(self) -> None:
        command = backend.build_run_command(
            backend.RunRequest(
                mode="clean-speaker-beta",
                source_path=Path(r"C:\videos"),
                output_root=Path(r"C:\output"),
                beta_gap_seconds=0,
                beta_resource_guard_percent=0,
                beta_resource_guard_timeout_seconds=0,
                beta_video_cooldown_seconds=0,
                beta_max_affinity_cores=0,
            ),
            repo_root=Path(r"C:\repo"),
            python_executable=Path(r"C:\Python312\python.exe"),
        )

        for option in (
            "--gap-seconds",
            "--resource-guard-percent",
            "--resource-guard-timeout-seconds",
            "--video-cooldown-seconds",
            "--max-affinity-cores",
        ):
            with self.subTest(option=option):
                index = command.index(option)
                self.assertEqual(float(command[index + 1]), 0.0)

    def test_clean_speaker_validation_rejects_nonfinite_and_inverted_thresholds(self) -> None:
        invalid_requests = [
            backend.RunRequest(
                mode="clean-speaker-beta",
                source_path=Path(r"C:\videos"),
                output_root=Path(r"C:\output"),
                beta_gap_seconds=float("nan"),
            ),
            backend.RunRequest(
                mode="clean-speaker-beta",
                source_path=Path(r"C:\videos"),
                output_root=Path(r"C:\output"),
                beta_cpu_throttle_high_percent=80,
                beta_cpu_throttle_low_percent=90,
            ),
        ]

        for request in invalid_requests:
            with self.subTest(request=request), self.assertRaises(ValueError):
                backend.build_run_command(request, repo_root=Path(r"C:\repo"))

    def test_youtube_api_key_can_come_from_ui_settings(self) -> None:
        with tempfile.TemporaryDirectory() as temp_dir:
            repo_root = Path(temp_dir)
            with patch.dict(os.environ, {"MEA_CREDENTIAL_STORE_ROOT": str(repo_root / "credentials")}):
                backend.save_ui_settings(repo_root, {"youtubeApiKey": "from-settings"})
                missing_config = repo_root / "missing.env"

                with patch.dict("os.environ", {"YOUTUBE_API_KEY": "", "MEA_CREDENTIAL_STORE_ROOT": str(repo_root / "credentials")}):
                    key = backend.load_youtube_api_key(config_path=missing_config, settings_path=backend.ui_settings_path(repo_root))

        self.assertEqual(key, "from-settings")

    def test_huggingface_token_can_come_from_ui_settings(self) -> None:
        with tempfile.TemporaryDirectory() as temp_dir:
            repo_root = Path(temp_dir)
            with patch.dict(os.environ, {"MEA_CREDENTIAL_STORE_ROOT": str(repo_root / "credentials")}):
                backend.save_ui_settings(repo_root, {"huggingFaceToken": " hf_token "})

                with patch.dict("os.environ", {"HF_TOKEN": "", "MEA_CREDENTIAL_STORE_ROOT": str(repo_root / "credentials")}):
                    token = backend.load_huggingface_token(settings_path=backend.ui_settings_path(repo_root))

        self.assertEqual(token, "hf_token")

    def test_youtube_api_key_has_no_plaintext_config_fallback(self) -> None:
        with tempfile.TemporaryDirectory() as temp_dir:
            root = Path(temp_dir)
            config = root / "config.env"
            config.write_text("YOUTUBE_API_KEY=plaintext-secret\n", encoding="utf-8")
            with patch.dict(
                os.environ,
                {"YOUTUBE_API_KEY": "", "MEA_CREDENTIAL_STORE_ROOT": str(root / "credentials")},
            ):
                key = backend.load_youtube_api_key(
                    config_path=config,
                    settings_path=root / "missing-settings.json",
                )

        self.assertEqual(key, "")

    def test_build_imotions_analysis_command_maps_ui_options_to_analysis_cli(self) -> None:
        command = backend.build_analysis_command(
            backend.AnalysisRunRequest(
                mode="imotions",
                source_path=Path(r"C:\imotions-export"),
                output_root=Path(r"C:\reports"),
                write_graphs=False,
                include_logscale=True,
                include_landmarks=True,
                include_timing=True,
                exclude_geometry=True,
            ),
            repo_root=Path(r"C:\repo"),
            python_executable=Path(r"C:\Python312\python.exe"),
        )

        self.assertEqual(command[:3], [r"C:\Python312\python.exe", "-m", "analysis.imotions"])
        self.assertIn(r"C:\imotions-export", command)
        self.assertIn("--output-root", command)
        self.assertIn(r"C:\reports", command)
        self.assertIn("--no-graphs", command)
        self.assertIn("--logscale", command)
        self.assertIn("--include-landmarks", command)
        self.assertIn("--include-timing", command)
        self.assertIn("--exclude-geometry", command)

    def test_build_audio_analysis_command_uses_analysis_audio(self) -> None:
        command = backend.build_analysis_command(
            backend.AnalysisRunRequest(
                mode="audio",
                source_path=Path(r"C:\audio-output"),
                output_root=Path(r"C:\reports"),
                include_logscale=True,
            ),
            repo_root=Path(r"C:\repo"),
            python_executable=Path(r"C:\Python312\python.exe"),
        )

        self.assertEqual(command[:3], [r"C:\Python312\python.exe", "-m", "analysis.audio"])
        self.assertIn(r"C:\audio-output", command)
        self.assertIn("--logscale", command)

    def test_build_analysis_workflow_command_serializes_mixed_modalities_and_groups(self) -> None:
        command = backend.build_analysis_workflow_command(
            backend.AnalysisWorkflowRunRequest(
                output_root=Path(r"C:\reports"),
                modalities=(
                    backend.AnalysisModalityRunRequest("imotions", "run", Path(r"C:\videos")),
                    backend.AnalysisModalityRunRequest("audio", "import", Path(r"C:\audio-reports")),
                    backend.AnalysisModalityRunRequest("text", "import", Path(r"C:\text-results")),
                ),
                speaker_groups=(
                    backend.AnalysisSpeakerGroupRunRequest(
                        group_id="group-1",
                        name="Group 1",
                        speaker_ids=("speaker_b", "speaker_a"),
                    ),
                ),
                default_reference=0.0,
                reference_overrides={"zeta": 2.0, "alpha": 1.0},
                include_construct_comparison=False,
                include_probability_sheets=False,
                confidence_level=0.90,
                headline_policy="equal",
                write_graphs=False,
                include_logscale=True,
                include_landmarks=True,
                include_timing=True,
                exclude_geometry=True,
            ),
            repo_root=Path(r"C:\repo"),
            python_executable=Path(r"C:\Python312\python.exe"),
        )

        self.assertEqual(command[:3], [r"C:\Python312\python.exe", "-m", "analysis.workflow"])
        self.assertEqual(command[command.index("--imotions-source") + 1], r"C:\videos")
        self.assertEqual(command[command.index("--imotions-method") + 1], "run")
        self.assertEqual(command[command.index("--audio-source") + 1], r"C:\audio-reports")
        self.assertEqual(command[command.index("--audio-method") + 1], "import")
        self.assertEqual(command[command.index("--text-source") + 1], r"C:\text-results")
        self.assertEqual(command[command.index("--text-method") + 1], "import")
        self.assertEqual(
            json.loads(command[command.index("--speaker-groups-json") + 1]),
            [{"id": "group-1", "name": "Group 1", "speakerKeys": ["speakerb", "speakera"]}],
        )
        self.assertEqual(
            command[command.index("--reference-overrides-json") + 1],
            '{"alpha": 1.0, "zeta": 2.0}',
        )
        self.assertIn("--no-graphs", command)
        self.assertIn("--logscale", command)
        self.assertIn("--include-landmarks", command)
        self.assertIn("--include-timing", command)
        self.assertIn("--exclude-geometry", command)
        self.assertIn("--no-construct-comparison", command)
        self.assertIn("--no-probability-sheets", command)
        self.assertEqual(command[command.index("--confidence-level") + 1], "0.9")
        self.assertEqual(command[command.index("--headline-policy") + 1], "equal")
        self.assertNotIn("--request-file", command)

    def test_build_analysis_workflow_command_rejects_text_and_duplicate_modalities(self) -> None:
        invalid_modalities = (
            (backend.AnalysisModalityRunRequest("text", "run", Path(r"C:\text")),),
            (
                backend.AnalysisModalityRunRequest("audio", "run", Path(r"C:\audio-one")),
                backend.AnalysisModalityRunRequest("audio", "import", Path(r"C:\audio-two")),
            ),
        )

        for modalities in invalid_modalities:
            with self.subTest(modalities=modalities), self.assertRaises(ValueError):
                backend.build_analysis_workflow_command(
                    backend.AnalysisWorkflowRunRequest(
                        output_root=Path(r"C:\reports"),
                        modalities=modalities,
                    ),
                    repo_root=Path(r"C:\repo"),
                )

        command = backend.build_analysis_workflow_command(
            backend.AnalysisWorkflowRunRequest(
                output_root=Path(r"C:\reports"),
                modalities=(backend.AnalysisModalityRunRequest("audio", "run", Path(r"C:\audio")),),
                speaker_groups=(
                    backend.AnalysisSpeakerGroupRunRequest(
                        group_id="group-1",
                        name="Group 1",
                        speaker_ids=("Researcher Alpha",),
                    ),
                ),
            ),
            repo_root=Path(r"C:\repo"),
        )
        self.assertIn("researcheralpha", command[command.index("--speaker-groups-json") + 1])

    def test_build_analysis_workflow_command_requires_groups_only_for_combined_output(self) -> None:
        base_request = backend.AnalysisWorkflowRunRequest(
            output_root=Path(r"C:\reports"),
            modalities=(backend.AnalysisModalityRunRequest("audio", "run", Path(r"C:\audio")),),
            speaker_groups=(),
            write_combined_workbook=True,
        )

        with self.assertRaisesRegex(ValueError, "speaker group"):
            backend.build_analysis_workflow_command(base_request, repo_root=Path(r"C:\repo"))

        command = backend.build_analysis_workflow_command(
            backend.AnalysisWorkflowRunRequest(
                output_root=base_request.output_root,
                modalities=base_request.modalities,
                speaker_groups=(),
                write_combined_workbook=False,
            ),
            repo_root=Path(r"C:\repo"),
        )
        self.assertIn("--no-combined-workbook", command)
        self.assertEqual(json.loads(command[command.index("--speaker-groups-json") + 1]), [])

    @unittest.skipUnless(sys.platform == "win32", "Windows junction regression")
    def test_analysis_workflow_command_preserves_output_junction_for_child_rejection(self) -> None:
        with tempfile.TemporaryDirectory() as temp_dir:
            root = Path(temp_dir)
            text_source = root / "text-results"
            text_source.mkdir()
            redirect = root / "outside"
            redirect.mkdir()
            parent_junction = root / "selected-output-parent"
            completed = subprocess.run(
                ["cmd", "/c", "mklink", "/J", str(parent_junction), str(redirect)],
                check=False,
                capture_output=True,
                text=True,
            )
            if completed.returncode != 0:
                self.skipTest(f"Could not create test junction: {completed.stderr.strip()}")
            selected_output = parent_junction / "reports"
            command = backend.build_analysis_workflow_command(
                backend.AnalysisWorkflowRunRequest(
                    output_root=selected_output,
                    modalities=(
                        backend.AnalysisModalityRunRequest("text", "import", text_source),
                    ),
                    write_combined_workbook=False,
                    include_probability_sheets=False,
                ),
                repo_root=Path(__file__).resolve().parents[2],
                python_executable=Path(sys.executable),
            )

            self.assertEqual(
                command[command.index("--output-root") + 1],
                str(Path(os.path.abspath(selected_output))),
            )
            child = subprocess.run(
                command,
                cwd=Path(__file__).resolve().parents[2],
                check=False,
                capture_output=True,
                text=True,
                timeout=30,
            )

            self.assertNotEqual(child.returncode, 0)
            self.assertRegex(child.stderr + child.stdout, "reparse|junction")
            self.assertEqual(tuple(redirect.iterdir()), ())

    def test_analysis_profile_context_and_command_have_no_source_or_group_caps(self) -> None:
        with tempfile.TemporaryDirectory() as temp_dir:
            root = Path(temp_dir)
            run = root / "run"
            selected = run / "processing" / "audio"
            selected.mkdir(parents=True)
            ordinary_text = root / "ordinary-text-results"
            ordinary_text.mkdir()
            sources = []
            for index in range(1, 15):
                speaker = f"Researcher {index:02d}"
                sources.append(
                    {
                        "source_id": f"source-{index:04d}",
                        "speaker": speaker,
                        "speaker_display": speaker,
                        "selected": True,
                        "system_metadata": {"title": f"Interview {index:02d}"},
                        "user_metadata": {
                            "Country": "Ireland" if index % 2 else "Japan",
                            "Language": "English",
                        },
                        "output_mapping": {"video_directory": str(run / speaker / f"source-{index:04d}")},
                    }
                )
            sources.append(
                {
                    "source_id": "source-0015",
                    "speaker": "Not selected",
                    "speaker_display": "Not selected",
                    "selected": False,
                    "system_metadata": {"title": "Not selected"},
                    "user_metadata": {"Country": "France", "Language": "French"},
                    "output_mapping": {"video_directory": str(run / "Not selected" / "source-0015")},
                }
            )
            manifest = run / "source_manifest.json"
            manifest.write_text(
                json.dumps(
                    {
                        "format_version": 1,
                        "catalog": {"metadata_headers": ["Country", "Language"]},
                        "sources": sources,
                    }
                ),
                encoding="utf-8",
            )
            (run / "source_metadata.csv").write_text(
                "SourceID,Country,Language\n"
                + "\n".join(
                    (
                        "source-0015,France,French"
                        if index == 15
                        else f"source-{index:04d},{'Ireland' if index % 2 else 'Japan'},English"
                    )
                    for index in range(1, 16)
                )
                + "\n",
                encoding="utf-8",
            )

            context = backend.discover_analysis_profile_context(
                (
                    backend.AnalysisModalityRunRequest("audio", "run", selected),
                    backend.AnalysisModalityRunRequest("text", "import", ordinary_text),
                )
            )
            digest = hashlib.sha256(manifest.read_bytes()).hexdigest()
            profile = AnalysisProfile(
                manifest,
                digest,
                sort_fields=("Country", "Language"),
                automatic_group_field="Country",
                manual_groups=tuple(
                    ManualGroup(
                        f"group-{index}",
                        f"Group {index}",
                        (ProfileMember("source", f"source-{index:04d}"),),
                    )
                    for index in range(1, 6)
                ),
            )
            command = backend.build_analysis_workflow_command(
                backend.AnalysisWorkflowRunRequest(
                    output_root=root / "reports",
                    modalities=(
                        backend.AnalysisModalityRunRequest("audio", "run", selected),
                        backend.AnalysisModalityRunRequest("text", "import", ordinary_text),
                    ),
                    analysis_profile=profile,
                ),
                repo_root=root,
                python_executable=Path(r"C:\Python312\python.exe"),
            )

        self.assertEqual(context["sourceManifest"], str(manifest.resolve()))
        self.assertEqual(context["sourceManifestSha256"], digest)
        self.assertEqual([field["name"] for field in context["metadataFields"]], ["Country", "Language"])
        self.assertEqual(context["metadataFields"][0]["values"], ["Ireland", "Japan"])
        self.assertEqual(len(context["sources"]), 14)
        self.assertEqual(len(context["speakers"]), 14)
        self.assertNotIn("--speaker-groups-json", command)
        self.assertEqual(command[command.index("--text-source") + 1], str(ordinary_text.resolve()))
        profile_payload = json.loads(command[command.index("--analysis-profile-json") + 1])
        self.assertEqual(len(profile_payload["manual_groups"]), 5)
        self.assertEqual(profile_payload["automatic_group_field"], "Country")

    def test_analysis_profile_command_rejects_an_unrelated_modality_manifest(self) -> None:
        with tempfile.TemporaryDirectory() as temp_dir:
            root = Path(temp_dir)

            def write_run(name: str) -> Path:
                run = root / name
                selected = run / "processing" / "audio"
                selected.mkdir(parents=True)
                manifest = run / "source_manifest.json"
                manifest.write_text(
                    json.dumps(
                        {
                            "format_version": 1,
                            "catalog": {"metadata_headers": ["Country"]},
                            "sources": [
                                {
                                    "source_id": "source-0001",
                                    "speaker": "Researcher Alpha",
                                    "speaker_display": "Researcher Alpha",
                                    "selected": True,
                                    "system_metadata": {"title": "Interview"},
                                    "user_metadata": {"Country": "Ireland"},
                                    "output_mapping": {
                                        "video_directory": str(run / "source-0001")
                                    },
                                }
                            ],
                        }
                    ),
                    encoding="utf-8",
                )
                (run / "source_metadata.csv").write_text(
                    "SourceID,Country\nsource-0001,Ireland\n",
                    encoding="utf-8",
                )
                return selected

            profile_source = write_run("profile-run")
            selected_source = write_run("selected-run")
            profile_manifest = profile_source.parents[1] / "source_manifest.json"
            profile = AnalysisProfile(
                profile_manifest,
                hashlib.sha256(profile_manifest.read_bytes()).hexdigest(),
            )

            with self.assertRaisesRegex(ValueError, "associated with the selected modality"):
                backend.build_analysis_workflow_command(
                    backend.AnalysisWorkflowRunRequest(
                        output_root=root / "output",
                        modalities=(
                            backend.AnalysisModalityRunRequest(
                                "audio", "run", selected_source
                            ),
                        ),
                        analysis_profile=profile,
                    ),
                    repo_root=root,
                )

    def test_analysis_profile_command_rejects_text_speaker_split_by_manual_or_automatic_groups(
        self,
    ) -> None:
        with tempfile.TemporaryDirectory() as temp_dir:
            root = Path(temp_dir)
            run = root / "run"
            audio = run / "analysis-input" / "audio"
            text = root / "ordinary-text-results"
            audio.mkdir(parents=True)
            text.mkdir()
            sources = [
                {
                    "source_id": f"source-{index:04d}",
                    "speaker": "Researcher Alpha",
                    "speaker_display": "Researcher Alpha",
                    "selected": True,
                    "system_metadata": {"title": f"Interview {index}"},
                    "user_metadata": {"Country": country},
                    "output_mapping": {
                        "video_directory": str(run / "Researcher Alpha" / f"source-{index:04d}")
                    },
                }
                for index, country in ((1, "Ireland"), (2, "Japan"))
            ]
            manifest = run / "source_manifest.json"
            manifest.write_text(
                json.dumps(
                    {
                        "format_version": 1,
                        "catalog": {"metadata_headers": ["Country"]},
                        "sources": sources,
                    }
                ),
                encoding="utf-8",
            )
            (run / "source_metadata.csv").write_text(
                "SourceID,Country\nsource-0001,Ireland\nsource-0002,Japan\n",
                encoding="utf-8",
            )
            digest = hashlib.sha256(manifest.read_bytes()).hexdigest()
            profiles = (
                AnalysisProfile(manifest, digest, automatic_group_field="Country"),
                AnalysisProfile(
                    manifest,
                    digest,
                    manual_groups=(
                        ManualGroup(
                            "first",
                            "First interview",
                            (ProfileMember("source", "source-0001"),),
                        ),
                    ),
                ),
            )
            modalities = (
                backend.AnalysisModalityRunRequest("audio", "import", audio),
                backend.AnalysisModalityRunRequest("text", "import", text),
            )

            for profile in profiles:
                with self.subTest(profile=profile), self.assertRaisesRegex(
                    ValueError,
                    "Text is speaker-level|same output group",
                ):
                    backend.build_analysis_workflow_command(
                        backend.AnalysisWorkflowRunRequest(
                            output_root=root / "output",
                            modalities=modalities,
                            analysis_profile=profile,
                        ),
                        repo_root=root,
                    )


if __name__ == "__main__":
    unittest.main()
